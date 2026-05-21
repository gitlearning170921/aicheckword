import os, re
from docx import Document

d = r'C:\Users\yuwell\Downloads\_extract_ctp'
files = [os.path.join(d, n) for n in os.listdir(d) if n.endswith('.docx')]
# pattern for ventilator model: letters then digits or Auto / ST etc.
pat = re.compile(r'YH[\- ]?\w+')
for f in files:
    print('FILE', os.path.basename(f))
    try:
        doc = Document(f)
    except Exception as e:
        print(' open fail', e)
        continue
    texts = []
    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    texts.append(c.text)
    full = '\n'.join(texts)
    models = sorted(set(m.upper() for m in pat.findall(full)))
    print(' total chars:', len(full))
    print(' models found:', len(models))
    print(' models:', models[:80])
    print()

# also dump 600 chars sample
sample = files[0]
doc = Document(sample)
out = []
for t in doc.tables[:2]:
    for r in t.rows[:10]:
        out.append(' | '.join((c.text or '').strip() for c in r.cells))
print('SAMPLE TABLE rows:\n' + '\n'.join(out)[:1500])
