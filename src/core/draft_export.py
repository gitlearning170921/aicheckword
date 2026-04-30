"""
文档初稿生成：输出产物导出（同格式）+ 修订记录写入。

说明：
- 目标：在“基础文件（Base）”上生成后，输出一个与基础文件同后缀的可下载文件。
- 默认以“保留基础文件格式”为最高优先级：不重建整份文档内容，避免样式/版式被打乱。
- 当前策略：复制基础文件 → 在原文件末尾追加“修订记录” + “自动生成内容（附录）”。
  这样可以最大限度保持注册递交所需的原模板格式；正文的精确段落/表格级修改依赖
  可定位的结构化 patch（由模型对照参考与基础差异生成），可按模板规则扩展。
"""

from __future__ import annotations

import datetime as _dt
import copy
import json
import random
import re
import shutil
import tempfile
import uuid
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, Optional, Tuple, Any, List

# 测试用例类表格：支持可配置的编号规则（优先来自 patch/meta），用于：
# - 已存在编号：整行相似度≥阈值则原地更新，否则分配新编号插入
# - 新插入：按同规则下“同模块前缀”的最大序号 +1 分配，避免与已有编号冲突
_TC_ROW_SIMILARITY_UPDATE_THRESHOLD = 0.9

_DEFAULT_TC_ID_RULES = [
    {
        "name": "dash_prefix_num",
        "regex": r"^([A-Za-z]+\d*)-(\d+)$",  # 例：GN3-21
        "prefix_group": 1,
        "number_group": 2,
        "render": "{prefix}-{num}",
    }
]


def _compile_tc_id_rules(raw_rules: Any) -> List[Dict[str, Any]]:
    rules: List[Dict[str, Any]] = []
    if isinstance(raw_rules, list):
        for r in raw_rules:
            if not isinstance(r, dict):
                continue
            rx = (r.get("regex") or "").strip()
            if not rx:
                continue
            try:
                compiled = re.compile(rx)
            except Exception:
                continue
            rules.append(
                {
                    "name": (r.get("name") or "rule").strip() or "rule",
                    "regex": rx,
                    "compiled": compiled,
                    "prefix_group": int(r.get("prefix_group") or 1),
                    "number_group": int(r.get("number_group") or 2),
                    "render": (r.get("render") or "{prefix}-{num}").strip() or "{prefix}-{num}",
                }
            )
    if rules:
        return rules
    # fallback：默认规则
    out: List[Dict[str, Any]] = []
    for r in _DEFAULT_TC_ID_RULES:
        try:
            out.append({**r, "compiled": re.compile(r["regex"])})
        except Exception:
            continue
    return out


def _parse_tc_id_first_cell(s: str, tc_rules: Optional[List[Dict[str, Any]]] = None) -> Optional[Tuple[str, int, Dict[str, Any]]]:
    txt = (s or "").strip()
    if not txt:
        return None
    rules = tc_rules or _compile_tc_id_rules(None)
    for r in rules:
        try:
            m = r["compiled"].match(txt)
        except Exception:
            m = None
        if not m:
            continue
        try:
            prefix = (m.group(int(r.get("prefix_group") or 1)) or "").strip()
            num = int(m.group(int(r.get("number_group") or 2)))
        except Exception:
            continue
        if prefix:
            return prefix, num, r
    return None


def _render_tc_id(prefix: str, num: int, rule: Dict[str, Any]) -> str:
    try:
        tmpl = (rule.get("render") or "{prefix}-{num}").strip() or "{prefix}-{num}"
        return tmpl.format(prefix=prefix, num=num)
    except Exception:
        return f"{prefix}-{num}"


def _word_cell_text_as_revised(cell) -> str:
    """按「接受修订后」的有效文本读取：跳过 w:del / w:moveFrom 内文本，保留 w:ins 与普通 run。

    用于编号扫描：若编号曾被修订，cell.text / paragraph.text 可能仍含删除线内容或漏读插入，
    导致最大流水号低估；此处以 OOXML 祖先节点为准近似「修订后」结果。
    """
    try:
        from docx.oxml.ns import qn

        tc = cell._tc  # type: ignore[attr-defined]
    except Exception:
        return ""
    try:
        pieces: List[str] = []
        for t in tc.iter(qn("w:t")):
            txt = t.text
            if not txt:
                continue
            el = t
            skip = False
            while el is not None:
                tag = el.tag
                if tag == qn("w:del"):
                    skip = True
                    break
                if tag == qn("w:moveFrom"):
                    skip = True
                    break
                el = el.getparent()
            if skip:
                continue
            pieces.append(txt)
        return "".join(pieces).strip()
    except Exception:
        return ""


def _word_cell_text_best(cell) -> str:
    """尽量读出单元格可见文本：优先「修订后」近似；再段落拼接；最后 cell.text。"""
    r = _word_cell_text_as_revised(cell)
    if (r or "").strip():
        return (r or "").strip()
    try:
        paras = getattr(cell, "paragraphs", None) or []
        if paras:
            t = "\n".join((p.text or "") for p in paras).strip()
            if t:
                return t
    except Exception:
        pass
    try:
        return (cell.text or "").strip()
    except Exception:
        return ""


def _table_max_tc_suffix_for_prefix_word(
    tbl,
    prefix: str,
    tc_rules: Optional[List[Dict[str, Any]]] = None,
    *,
    max_cols: int = 8,
    max_rows: int = 8000,
) -> int:
    """同前缀下流水号最大值：扫描每行前若干列（编号未必在第 1 列，合并格时尤甚）。"""
    mx = 0
    try:
        rows = tbl.rows
        lim = min(len(rows), max_rows)
        for row in rows[:lim]:
            if not row.cells:
                continue
            ncells = len(row.cells)
            for ci in range(min(max_cols, ncells)):
                txt = _word_cell_text_best(row.cells[ci])
                p = _parse_tc_id_first_cell(txt, tc_rules)
                if p and p[0] == prefix:
                    mx = max(mx, p[1])
    except Exception:
        pass
    return mx


def _table_find_row_idx_by_id_any_cell(
    tbl, id_text: str, *, max_cols: int = 8, max_rows: int = 8000
) -> int:
    t = (id_text or "").strip()
    if not t:
        return -1
    try:
        rows = tbl.rows
        lim = min(len(rows), max_rows)
        for i, row in enumerate(rows[:lim]):
            if not row.cells:
                continue
            ncells = len(row.cells)
            for ci in range(min(max_cols, ncells)):
                if _word_cell_text_best(row.cells[ci]).strip() == t:
                    return i
    except Exception:
        pass
    return -1


def _table_tc_id_exists_word(tbl, id_text: str, **kwargs) -> bool:
    return _table_find_row_idx_by_id_any_cell(tbl, id_text, **kwargs) >= 0


def _next_unique_tc_id_word(
    tbl,
    prefix: str,
    rule: Dict[str, Any],
    tc_rules: Optional[List[Dict[str, Any]]],
    start_num: int,
    *,
    max_try: int = 5000,
) -> str:
    """从 start_num 起渲染编号，若表中已存在则递增，直到不冲突（防重复）。"""
    num = int(start_num)
    for _ in range(max_try):
        cand = _render_tc_id(prefix, num, rule)
        if not _table_tc_id_exists_word(tbl, cand):
            return cand
        num += 1
    return _render_tc_id(prefix, int(start_num), rule)


def _word_row_joined(row, col_n: int) -> str:
    parts: List[str] = []
    try:
        for ci in range(min(col_n, len(row.cells))):
            parts.append(_word_cell_text_best(row.cells[ci]))
    except Exception:
        pass
    return "\t".join(parts)


def _table_max_tc_suffix_for_prefix_xlsx(
    ws,
    prefix: str,
    tc_rules: Optional[List[Dict[str, Any]]] = None,
    *,
    scan_cols: int = 8,
    max_row: int = 8000,
) -> int:
    mx = 0
    try:
        last_r = min(ws.max_row or 0, max_row)
        for r in range(1, last_r + 1):
            for c in range(1, scan_cols + 1):
                v = ws.cell(row=r, column=c).value
                s = "" if v is None else str(v)
                p = _parse_tc_id_first_cell(s, tc_rules)
                if p and p[0] == prefix:
                    mx = max(mx, p[1])
    except Exception:
        pass
    return mx


def _xlsx_find_row_by_id_any_col(
    ws, id_text: str, *, scan_cols: int = 8, max_scan: int = 8000
) -> int:
    t = (id_text or "").strip()
    lim = min((ws.max_row or 0) + 1, max_scan)
    try:
        for r in range(1, lim):
            for c in range(1, scan_cols + 1):
                v = ws.cell(row=r, column=c).value
                s = "" if v is None else str(v)
                if s.strip() == t:
                    return r
    except Exception:
        pass
    return -1


def _xlsx_tc_id_exists(ws, id_text: str, **kwargs) -> bool:
    return _xlsx_find_row_by_id_any_col(ws, id_text, **kwargs) >= 0


def _next_unique_tc_id_xlsx(
    ws,
    prefix: str,
    rule: Dict[str, Any],
    tc_rules: Optional[List[Dict[str, Any]]],
    start_num: int,
    *,
    max_try: int = 5000,
) -> str:
    num = int(start_num)
    for _ in range(max_try):
        cand = _render_tc_id(prefix, num, rule)
        if not _xlsx_tc_id_exists(ws, cand):
            return cand
        num += 1
    return _render_tc_id(prefix, int(start_num), rule)


def _xlsx_row_joined(ws, r: int, col_n: int) -> str:
    parts: List[str] = []
    try:
        for c in range(1, col_n + 1):
            v = ws.cell(row=r, column=c).value
            parts.append("" if v is None else str(v))
    except Exception:
        pass
    return "\t".join(parts)


def _xlsx_cell_row_col(cell_ref: str) -> Tuple[int, int]:
    """'D6' -> (row, col_index)。"""
    from openpyxl.utils.cell import coordinate_from_string
    from openpyxl.utils import column_index_from_string

    col_letter, row = coordinate_from_string(str(cell_ref).strip().upper())
    return int(row), int(column_index_from_string(col_letter))


def _xlsx_last_nonblank_col(ws, r: int, max_c: int = 60) -> int:
    last = 0
    try:
        for c in range(1, min(max_c, int(ws.max_column or 0) + 1, 200)):
            v = ws.cell(row=r, column=c).value
            if v is not None and str(v).strip():
                last = c
    except Exception:
        pass
    return last


def _xlsx_first_cell_risk_id_token(s: str) -> Optional[str]:
    """识别如 HC37 / GN3-21 等首列风险或编号单元格（用于去重更新）。"""
    t = (s or "").strip()
    if not t:
        return None
    # 允许 HC37Violation… 粘连：编号后紧跟英文大写也算边界
    m = re.match(r"^([A-Z]{2,5}\d{1,4})(?=[A-Z]|\b|_|-|\s|$)", t)
    return m.group(1) if m else None


def _xlsx_split_row_values_for_insert(raw: str, col_hint: int) -> List[str]:
    """insert_table_row：优先 \\t 分列；无制表符时尝试 | 或粘连英文边界拆分，避免整行写入 A 列。"""
    s = (raw or "").strip()
    if not s:
        return []
    if "\t" in s:
        return [p.strip() for p in s.split("\t")]
    if "|" in s:
        parts = [p.strip() for p in s.split("|")]
        if len(parts) >= 2:
            return parts
    if col_hint < 2 or len(s) < 50:
        return [s]
    m = re.match(r"^([A-Z]{2,5}\d{1,4})(.+)$", s)
    if not m:
        return [s]
    parts: List[str] = [m.group(1)]
    rest = (m.group(2) or "").strip()
    if not rest:
        return [s]
    chunks = re.split(r"(?<=[a-z0-9\)\.\”\"])(?=[A-Z][a-zA-Z])", rest)
    for ch in chunks:
        ch = ch.strip()
        if ch:
            parts.append(ch)
    if len(parts) < 2:
        return [s]
    if len(parts) > col_hint and col_hint >= 2:
        head = parts[: col_hint - 1]
        tail = " ".join(parts[col_hint - 1 :])
        parts = head + [tail]
    return parts


def _word_find_similar_row_idx(tbl, *, joined: str, col_n: int, start: int, end: int) -> int:
    """在 Word 表格中查找与 joined 高度相似的行，用于去重/改为原地更新。"""
    try:
        aa = (joined or "").strip()
        if not aa:
            return -1
        n = len(getattr(tbl, "rows", []) or [])
        s = max(0, min(int(start), n))
        e = max(0, min(int(end), n))
        if e <= s:
            return -1
        best_i = -1
        best_sim = 0.0
        for i in range(s, e):
            try:
                bb = (_word_row_joined(tbl.rows[i], col_n) or "").strip()
                if not bb:
                    continue
                if bb == aa:
                    return i
                sim = float(SequenceMatcher(None, aa, bb).ratio())
                if sim > best_sim:
                    best_sim = sim
                    best_i = i
            except Exception:
                continue
        # 用更严格阈值：无编号场景只做“强相似”去重，避免误伤
        return best_i if best_sim >= 0.95 else -1
    except Exception:
        return -1


def _word_clone_row_after(tbl, anchor_row):
    """
    克隆 anchor_row 的 XML 并插入到其后，尽量继承原行的单元格样式/结构（含复选框等）。
    失败则返回 None，由上层走 add_row 回退。
    """
    try:
        new_tr = copy.deepcopy(anchor_row._tr)
        anchor_row._tr.addnext(new_tr)
        try:
            # python-docx 私有类型：尽量兼容，失败则由上层回退
            from docx.table import _Row  # type: ignore

            return _Row(new_tr, tbl)
        except Exception:
            # 退化：刷新 rows 后取 anchor_row 的下一个
            try:
                rows = list(getattr(tbl, "rows", []) or [])
                idx = rows.index(anchor_row)
                if idx >= 0 and idx + 1 < len(rows):
                    return rows[idx + 1]
            except Exception:
                pass
        return None
    except Exception:
        return None


def _word_unique_cells(row) -> List[Any]:
    """返回一行中按底层 tc 去重后的 cell 列表（合并单元格在 row.cells 中可能重复出现）。"""
    out: List[Any] = []
    seen: set = set()
    try:
        for c in list(getattr(row, "cells", []) or []):
            try:
                tc = getattr(c, "_tc", None)
                key = id(tc) if tc is not None else id(c)
            except Exception:
                key = id(c)
            if key in seen:
                continue
            seen.add(key)
            out.append(c)
    except Exception:
        return list(getattr(row, "cells", []) or [])
    return out


def _word_best_template_row_for_insert(tbl, anchor_row_idx: int):
    """
    选择“更像数据行”的模板行用于克隆，避免克隆表头（表头常有合并单元格导致写入挤到同一格）。
    优先：锚点下一行；否则在前几行中选 unique_cells 数最多的行。
    """
    try:
        rows = list(getattr(tbl, "rows", []) or [])
        if not rows:
            return None
        if 0 <= int(anchor_row_idx) < len(rows) - 1:
            cand = rows[int(anchor_row_idx) + 1]
            if len(_word_unique_cells(cand)) >= 2:
                return cand
        # 扫描前几行，选列数最多的（通常是数据行而非合并表头）
        best = rows[min(max(int(anchor_row_idx), 0), len(rows) - 1)]
        best_n = len(_word_unique_cells(best))
        for r in rows[: min(6, len(rows))]:
            n = len(_word_unique_cells(r))
            if n > best_n:
                best, best_n = r, n
        return best
    except Exception:
        return None


def _word_pick_value_cell_hits(
    hits: List[Tuple[Any, int, int, Any]], anchor: str
) -> List[Tuple[Any, int, int, Any]]:
    """
    Word 表格：多命中时尽量选择“值单元格”而非“字段名/标签”单元格。
    规则：
    - 优先过滤掉 text == anchor 的单元格（典型标签格）
    - 若命中表头行（第 1 行字段名）：优先落到下一行同列的值格（避免覆盖表头）
    - 同一行若仍多格命中：取该行最右侧单元格
    """
    a = (anchor or "").strip()
    if not hits:
        return hits
    pool: List[Tuple[Any, int, int, Any]] = []
    for tbl, r_idx, c_idx, c in hits:
        try:
            if a and (c.text or "").strip() == a:
                continue
        except Exception:
            pass
        pool.append((tbl, r_idx, c_idx, c))
    if not pool:
        pool = list(hits)

    # 表头命中：把落点移到“下一行同列”（常见修订记录/统计表格）
    shifted: List[Tuple[Any, int, int, Any]] = []
    for tbl, r_idx, c_idx, c in pool:
        try:
            if int(r_idx) == 0 and len(getattr(tbl, "rows", []) or []) >= 2:
                row1 = tbl.rows[1]
                ci = int(c_idx)
                ci = max(0, min(ci, len(row1.cells) - 1))
                shifted.append((tbl, 1, ci, row1.cells[ci]))
            else:
                shifted.append((tbl, r_idx, c_idx, c))
        except Exception:
            shifted.append((tbl, r_idx, c_idx, c))
    pool = shifted
    try:
        # 以 (table, row_idx) 分组，取最右侧列
        by_row: Dict[Tuple[int, int], List[Tuple[int, Tuple[Any, int, int, Any]]]] = {}
        for ent in pool:
            tbl, r_idx, c_idx, c = ent
            key = (id(tbl), int(r_idx))
            by_row.setdefault(key, []).append((int(c_idx), ent))
        out: List[Tuple[Any, int, int, Any]] = []
        for _k, grp in by_row.items():
            grp.sort(key=lambda x: x[0], reverse=True)
            out.append(grp[0][1])
        return out
    except Exception:
        return pool


def _xlsx_pick_value_cell_hits(
    hits: List[Tuple[Any, Any, str]], anchor: str
) -> List[Tuple[Any, Any, str]]:
    """多命中时去掉「仅字段名单元格」，并优先保留同行最右侧（常见标签在左、值在右）。"""
    a = (anchor or "").strip()
    if not hits:
        return hits
    filtered: List[Tuple[Any, Any, str]] = []
    for ws, cell, before in hits:
        b = str(before or "").strip()
        if a and b == a:
            continue
        filtered.append((ws, cell, before))
    pool = filtered if filtered else list(hits)
    try:
        from collections import defaultdict

        by_row: Dict[Tuple[Any, int], List[Tuple[Any, Any, str]]] = defaultdict(list)
        for ws, cell, before in pool:
            by_row[(ws, int(cell.row))].append((ws, cell, before))
        out: List[Tuple[Any, Any, str]] = []
        for _k, grp in by_row.items():
            grp.sort(key=lambda x: int(x[1].column), reverse=True)
            out.append(grp[0])
        out.sort(key=lambda x: (x[0].title, x[1].row, -int(x[1].column)))
        return out
    except Exception:
        return pool


def _now_str() -> str:
    return _dt.datetime.now().strftime("%Y-%m-%d %H:%M")


def _new_revision_id() -> str:
    return str(random.randint(100000, 99999999))


def _enable_track_revisions(doc) -> None:
    """在 settings 中打开「跟踪修订」，便于 Word 以修订视图显示插入/删除。"""
    try:
        el = doc.settings.element
    except Exception:
        return
    try:
        from docx.oxml.ns import qn

        for child in list(el):
            if child.tag == qn("w:trackRevisions"):
                return
        from docx.oxml import OxmlElement

        el.append(OxmlElement("w:trackRevisions"))
    except Exception:
        pass


def _enable_update_fields_on_open(doc) -> None:
    """
    打开文档时自动更新域（Fields），用于目录（TOC）、页码、交叉引用等在页数变化后自动重算。
    注：python-docx 无法在服务端真正“计算并更新”这些域，只能设置该开关由 Word/WPS 在打开时更新。
    """
    try:
        el = doc.settings.element
    except Exception:
        return
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for child in list(el):
            if child.tag == qn("w:updateFields"):
                try:
                    child.set(qn("w:val"), "true")
                except Exception:
                    pass
                return
        uf = OxmlElement("w:updateFields")
        try:
            uf.set(qn("w:val"), "true")
        except Exception:
            pass
        el.append(uf)
    except Exception:
        pass


def _replace_paragraph_with_track_changes(
    p, before: str, after: str, author: str = "aicheckword", *, run_rpr=None
) -> None:
    """
    将段落改为「删除旧全文 + 插入新全文」的修订结构（w:del / w:ins），保留段落样式对象 p 不变。
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    iso = _dt.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    b = before or ""
    a = after or ""

    # 清空原有子节点（runs 等），但尽量保留段落属性 pPr（对齐/间距/样式等），避免表格/正文格式跑偏
    p_el = p._p
    ppr = None
    try:
        for child in list(p_el):
            try:
                if child.tag == qn("w:pPr"):
                    ppr = child
                    break
            except Exception:
                continue
    except Exception:
        ppr = None
    for child in list(p_el):
        try:
            if ppr is not None and child is ppr:
                continue
        except Exception:
            pass
        p_el.remove(child)
    # 确保 pPr 在最前（Word 通常要求 pPr 先于内容）
    try:
        if ppr is not None:
            try:
                p_el.remove(ppr)
            except Exception:
                pass
            p_el.insert(0, ppr)
    except Exception:
        pass

    # w:del
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), _new_revision_id())
    del_el.set(qn("w:author"), author)
    del_el.set(qn("w:date"), iso)
    r0 = OxmlElement("w:r")
    try:
        if run_rpr is not None:
            r0.append(copy.deepcopy(run_rpr))
    except Exception:
        pass
    dt = OxmlElement("w:delText")
    try:
        dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    except Exception:
        pass
    dt.text = b
    r0.append(dt)
    del_el.append(r0)
    p_el.append(del_el)

    # w:ins
    ins_el = OxmlElement("w:ins")
    ins_el.set(qn("w:id"), _new_revision_id())
    ins_el.set(qn("w:author"), author)
    ins_el.set(qn("w:date"), iso)
    r1 = OxmlElement("w:r")
    try:
        if run_rpr is not None:
            r1.append(copy.deepcopy(run_rpr))
    except Exception:
        pass
    t = OxmlElement("w:t")
    try:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    except Exception:
        pass
    t.text = a
    r1.append(t)
    ins_el.append(r1)
    p_el.append(ins_el)


def _delete_paragraph_with_track_changes(p, before: str, author: str = "aicheckword") -> None:
    """仅删除（w:del），不插入新文本。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    iso = _dt.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    b = before or ""

    p_el = p._p
    for child in list(p_el):
        p_el.remove(child)

    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), _new_revision_id())
    del_el.set(qn("w:author"), author)
    del_el.set(qn("w:date"), iso)
    r0 = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    try:
        dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    except Exception:
        pass
    dt.text = b
    r0.append(dt)
    del_el.append(r0)
    p_el.append(del_el)


def _replace_table_cell_with_track_changes(cell, before: str, after: str, author: str = "aicheckword") -> None:
    """
    在表格单元格内写入修订：删除旧文本 + 插入新文本（尽量保留单元格本身结构，不重建表格）。
    """
    # python-docx 对 cell.text 的赋值会重建段落并丢失 runs；这里直接在第一个段落写入 w:del/w:ins
    try:
        paras = list(getattr(cell, "paragraphs", []) or [])
    except Exception:
        paras = []
    if not paras:
        try:
            p = cell.add_paragraph("")
            paras = [p]
        except Exception:
            return

    p0 = paras[0]
    # 捕获基础字体/字号等 run 属性，保证插入与基础文档一致（如 Times New Roman）
    run_rpr = None
    try:
        if p0.runs:
            run_rpr = p0.runs[0]._r.rPr  # type: ignore[attr-defined]
    except Exception:
        run_rpr = None
    # 清理多余段落，避免旧内容残留
    try:
        tc = cell._tc  # type: ignore[attr-defined]
        for p in paras[1:]:
            try:
                tc.remove(p._p)
            except Exception:
                continue
    except Exception:
        pass

    _replace_paragraph_with_track_changes(p0, before or "", after or "", author=author, run_rpr=run_rpr)


def _insert_paragraph_with_track_changes(
    p, text: str, author: str = "aicheckword", *, run_rpr=None
) -> None:
    """
    将段落内容写为“插入修订”。
    部分 Word 版本对“仅 w:ins、无 w:del”的插入在修订气泡/窗格中展示不稳定；
    这里用“零宽字符占位删除 + 全文插入”形成成对修订，便于与正文替换类修订一致展示。
    """
    if not (text or "").strip():
        try:
            p_el = p._p
            for child in list(p_el):
                p_el.remove(child)
        except Exception:
            pass
        return
    _replace_paragraph_with_track_changes(p, "\u200b", text or "", author=author, run_rpr=run_rpr)


def _render_modules_diagram_png(*, title: str, modules: List[str], out_path: str) -> str:
    """
    生成一个“正式文档风格”的模块框图 PNG（不追求与原图一模一样，但保证清晰、可审阅、一致性强）。
    仅依赖 Pillow。
    """
    from PIL import Image, ImageDraw, ImageFont

    mods = [m.strip() for m in (modules or []) if str(m or "").strip()]
    if not mods:
        mods = ["(No modules)"]

    # 基础参数
    W = 1400
    margin = 60
    title_h = 90
    box_h = 90
    gap_y = 28
    cols = 2 if len(mods) >= 6 else 1
    col_gap = 60
    rows = (len(mods) + cols - 1) // cols
    H = margin + title_h + rows * box_h + (rows - 1) * gap_y + margin
    img = Image.new("RGB", (W, H), "white")
    dr = ImageDraw.Draw(img)

    # 字体：优先用系统字体（Windows 常见）
    def _load_font(size: int):
        for fp in [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
            r"C:\Windows\Fonts\msyh.ttc",  # 微软雅黑（中英文）
        ]:
            try:
                return ImageFont.truetype(fp, size=size)
            except Exception:
                continue
        return ImageFont.load_default()

    font_title = _load_font(32)
    font_box = _load_font(26)

    # 标题
    t = (title or "").strip() or "Logical Structure / Module Overview"
    dr.text((margin, margin), t, fill=(0, 0, 0), font=font_title)
    y0 = margin + title_h

    # 计算列宽
    box_w = (W - margin * 2 - (col_gap if cols == 2 else 0)) // cols

    # 画框
    for idx, m in enumerate(mods):
        r = idx // cols
        c = idx % cols
        x1 = margin + c * (box_w + (col_gap if cols == 2 else 0))
        y1 = y0 + r * (box_h + gap_y)
        x2 = x1 + box_w
        y2 = y1 + box_h
        dr.rounded_rectangle([x1, y1, x2, y2], radius=14, outline=(0, 0, 0), width=3, fill=(255, 255, 255))
        # 文本居中，过长则自动换行（简化实现：按宽度切）
        text = str(m)
        # 简单断行：按字符宽度估算
        max_w = box_w - 30
        lines = []
        cur = ""
        for ch in text:
            test = (cur + ch).strip()
            tw = dr.textlength(test, font=font_box)
            if tw <= max_w or not cur:
                cur = test
            else:
                lines.append(cur)
                cur = ch
        if cur:
            lines.append(cur)
        if len(lines) > 2:
            lines = lines[:2]
            # 末行加省略
            lines[-1] = (lines[-1][: max(0, len(lines[-1]) - 1)] + "…") if lines[-1] else "…"

        total_h = len(lines) * 32
        ty = y1 + (box_h - total_h) / 2
        for li in lines:
            tw = dr.textlength(li, font=font_box)
            tx = x1 + (box_w - tw) / 2
            dr.text((tx, ty), li, fill=(0, 0, 0), font=font_box)
            ty += 32

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path, format="PNG")
    return out_path


def _docx_replace_or_insert_image_after_paragraph(doc, *, anchor_p_idx: int, image_path: str, width_inches: float = 6.5) -> bool:
    """在指定段落之后替换最近的图片段落，找不到则插入一张图片段落。"""
    from docx.shared import Inches

    # 在 anchor 后找一个“图片段落”
    try:
        for j in range(anchor_p_idx + 1, min(anchor_p_idx + 12, len(doc.paragraphs))):
            p2 = doc.paragraphs[j]
            xml = p2._p.xml
            if "w:drawing" in xml or "w:pict" in xml:
                # 清空该段落并插入新图
                for r in list(p2.runs):
                    try:
                        r.text = ""
                    except Exception:
                        pass
                try:
                    for child in list(p2._p):
                        p2._p.remove(child)
                except Exception:
                    pass
                run = p2.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                return True
    except Exception:
        pass

    # 插入新段落图片
    try:
        p = doc.paragraphs[anchor_p_idx]
        new_p = p.insert_paragraph_after("")  # type: ignore[attr-defined]
    except Exception:
        # fallback：在末尾加
        new_p = doc.add_paragraph("")
    try:
        run = new_p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        return True
    except Exception:
        return False

def _rev_table_rows(meta: Dict) -> Tuple[list, str]:
    """
    返回 (rows, version_tag)
    rows: list[list[str]] 用于写入修订记录表格/工作表
    """
    ts = _now_str()
    version_tag = (meta.get("version_tag") or "").strip() or _dt.datetime.now().strftime("draft-%Y%m%d%H%M%S")
    rows = [
        ["日期", "版本", "修订内容", "生成方式/来源"],
        [
            ts,
            version_tag,
            (meta.get("change_summary") or "基于基础文件按规则补写/修订生成").strip(),
            (meta.get("generated_by") or "aicheckword 文档初稿生成").strip(),
        ],
    ]
    return rows, version_tag


def _revision_entry_fields(meta: Dict) -> Dict[str, str]:
    """
    修订记录行语义字段（与具体表头列顺序无关，由表头映射决定落列）。
    """
    m = meta or {}
    version_tag = (m.get("version_tag") or "").strip() or _dt.datetime.now().strftime("draft-%Y%m%d%H%M%S")
    # 日期：支持模板常见格式（Excel 多用 2024.10.24，Word 多用 2024/10/24）
    date_fmt = (m.get("revision_date_format") or "slash").strip().lower()
    now = _dt.datetime.now()
    if date_fmt in ("dot", "excel", "xlsx"):
        date_str = now.strftime("%Y.%m.%d")
    else:
        date_str = now.strftime("%Y/%m/%d")
    return {
        "version": version_tag,
        "date": date_str,
        "author": (m.get("revision_author") or m.get("author") or "aicheckword").strip(),
        "change_no": (m.get("revision_change_no") or m.get("change_no") or "NA").strip(),
        "change_content": (m.get("change_summary") or "基于基础文件按规则补写/修订生成").strip(),
        "source": (m.get("generated_by") or "aicheckword 文档初稿生成").strip(),
    }


def _compute_next_version(prev: str) -> str:
    """
    将既有版本号升级到下一版：
    - A/1 -> A/2
    - v1 -> v2
    - 01 -> 02（保留位数）
    - 1.2 -> 1.3（简单最后段递增）
    无法解析时返回原值。
    """
    s = (prev or "").strip()
    if not s:
        return s
    import re

    m = re.fullmatch(r"([A-Za-z]+)\s*/\s*(\d+)", s)
    if m:
        return f"{m.group(1)}/{int(m.group(2)) + 1}"
    m = re.fullmatch(r"([Vv])\s*(\d+)", s)
    if m:
        return f"{m.group(1)}{int(m.group(2)) + 1}"
    m = re.fullmatch(r"(\d+)", s)
    if m:
        w = len(m.group(1))
        return str(int(m.group(1)) + 1).zfill(w)
    m = re.fullmatch(r"(\d+(?:\.\d+)+)", s)
    if m:
        parts = s.split(".")
        try:
            parts[-1] = str(int(parts[-1]) + 1)
            return ".".join(parts)
        except Exception:
            return s
    return s


def _docx_sync_version_everywhere(doc, *, old: str, new: str) -> int:
    """
    在 docx 中同步版本号：封面/正文/页眉页脚（不追求精确字段定位，做“精确旧值替换”）。
    返回替换次数。
    """
    if not old or not new or old == new:
        return 0

    def _replace_in_paras(paras):
        n = 0
        for p in list(paras or []):
            try:
                txt = p.text or ""
            except Exception:
                continue
            if old not in txt:
                continue
            try:
                _replace_paragraph_with_track_changes(p, txt, txt.replace(old, new))
            except Exception:
                # fallback：不做修订，直接替换 runs
                try:
                    for r in list(p.runs):
                        if old in (r.text or ""):
                            r.text = (r.text or "").replace(old, new)
                    n += 1
                except Exception:
                    pass
            else:
                n += 1
        return n

    count = 0
    # 正文
    try:
        count += _replace_in_paras(getattr(doc, "paragraphs", []) or [])
    except Exception:
        pass
    # 页眉页脚
    try:
        for sec in list(getattr(doc, "sections", []) or []):
            for hf in [getattr(sec, "header", None), getattr(sec, "footer", None)]:
                if hf is None:
                    continue
                count += _replace_in_paras(getattr(hf, "paragraphs", []) or [])
    except Exception:
        pass
    return count


def _xlsx_sync_version_everywhere(wb, *, old: str, new: str) -> int:
    """
    在 xlsx 中同步版本号：扫描每个工作表的前若干行列，做精确旧值替换。
    返回替换次数。
    """
    if not old or not new or old == new:
        return 0
    cnt = 0
    try:
        for ws in list(getattr(wb, "worksheets", []) or []):
            max_r = min((ws.max_row or 0), 120) or 120
            max_c = min((ws.max_column or 0), 30) or 30
            for r in range(1, max_r + 1):
                for c in range(1, max_c + 1):
                    try:
                        v = ws.cell(row=r, column=c).value
                    except Exception:
                        continue
                    if v is None:
                        continue
                    s = str(v)
                    if old in s:
                        ws.cell(row=r, column=c, value=s.replace(old, new))
                        cnt += 1
    except Exception:
        pass
    return cnt


def _header_cell_norm(s: str) -> str:
    t = (s or "").replace("\n", " ").replace("\r", " ").strip().lower()
    while "  " in t:
        t = t.replace("  ", " ")
    return t


def _header_to_slot(header_text: str) -> Optional[str]:
    """
    将表头单元格文本映射到语义槽：version / date / author / change_no / change_content / source
    """
    h = _header_cell_norm(header_text)
    if not h:
        return None

    # 变更号 / CR（优先于泛泛的 change）
    if any(
        k in h
        for k in [
            "change no",
            "change no.",
            "change#",
            "change number",
            "cr no",
            "变更编号",
            "变更号",
            "变更单号",
            "ecr",
        ]
    ):
        return "change_no"
    # 日期
    if any(
        k in h
        for k in [
            "change date",
            "revision date",
            "rev date",
            "date",
            "日期",
            "变更日期",
            "修订日期",
            "发布日期",
        ]
    ) and "content" not in h:
        return "date"
    # 版本
    if any(k in h for k in ["version", "ver", "版本", "版次", "版 本"]):
        return "version"
    # 作者
    if any(
        k in h
        for k in [
            "author",
            "prepared",
            "responsible",
            "编制",
            "编制人",
            "修订人",
            "审核人",
            "作者",
            "责任人",
        ]
    ):
        return "author"
    # 变更内容 / 说明
    if any(
        k in h
        for k in [
            "change content",
            "revision content",
            "description of change",
            "description",
            "修订内容",
            "变更内容",
            "变更说明",
            "说明",
            "摘要",
            "remarks",
        ]
    ):
        return "change_content"
    # 来源 / 生成方式
    if any(k in h for k in ["source", "生成方式", "来源", "generated", "tool"]):
        return "source"
    # 宽泛匹配（放最后，避免误判）
    if h == "change" or h.startswith("change ") and "content" not in h and "date" not in h:
        return "change_no"
    if "content" in h and "change" in h:
        return "change_content"
    return None


def _build_header_slot_to_col(headers: List[str]) -> Dict[str, int]:
    """slot -> 0-based column index（每个槽只取第一次命中）。"""
    out: Dict[str, int] = {}
    for i, raw in enumerate(headers):
        slot = _header_to_slot(str(raw or ""))
        if slot and slot not in out:
            out[slot] = i
    return out


def _revision_values_by_header(headers: List[str], fields: Dict[str, str]) -> List[str]:
    """按表头列数生成一行值，未识别列留空。"""
    n = len(headers)
    row = [""] * n
    slot_map = _build_header_slot_to_col(headers)
    for slot, col in slot_map.items():
        if 0 <= col < n:
            row[col] = str(fields.get(slot) or "")
    # 若表头完全未映射（老模板），按列数退回简单顺序：version,date,author,change_no,change_content 或 4 列旧顺序
    if not slot_map and n > 0:
        legacy = [
            fields.get("date", ""),
            fields.get("version", ""),
            fields.get("change_content", ""),
            fields.get("source", ""),
        ]
        for i in range(min(n, len(legacy))):
            row[i] = legacy[i]
    return row


def _norm_ws(s: str) -> str:
    return " ".join(((s or "").replace("\u00a0", " ")).split()).strip().lower()


def _para_text_similarity(a: str, b: str) -> float:
    aa = _norm_ws(a)
    bb = _norm_ws(b)
    if not aa or not bb:
        return 0.0
    try:
        return float(SequenceMatcher(None, aa, bb).ratio())
    except Exception:
        return 0.0


def _looks_like_duplicate_followup_paragraph(new_text: str, candidate_text: str) -> bool:
    """判断 candidate 是否已承载与 new_text 高度重复的内容（用于避免 insert 造成“叠两段相似正文”）。"""
    nt = _norm_ws(new_text)
    ct = _norm_ws(candidate_text)
    if not nt or not ct:
        return False
    if nt == ct:
        return True
    if _para_text_similarity(nt, ct) >= 0.82:
        return True
    # 一方显著包含另一方（常见于：原段落更长，新段落是其子集/改写摘要）
    if len(nt) >= 40 and nt in ct:
        return True
    if len(ct) >= 40 and ct in nt:
        return True
    # token 覆盖率（对英文长句更稳）
    try:
        toks = [x for x in re.split(r"[^0-9a-zA-Z\u4e00-\u9fff]+", nt) if len(x) >= 4]
        if len(toks) >= 6:
            hit = sum(1 for t in toks if t.lower() in ct)
            if hit / max(1, len(toks)) >= 0.78:
                return True
    except Exception:
        pass
    return False


def _paragraph_effective_run_rpr(p) -> Any:
    """取段落中第一个带 rPr 的 run 属性副本，便于修订插入/替换时继承字体字号等。"""
    try:
        for r in getattr(p, "runs", None) or []:
            try:
                rp = r._r.rPr  # type: ignore[attr-defined]
                if rp is not None:
                    return copy.deepcopy(rp)
            except Exception:
                continue
    except Exception:
        pass
    return None


def _max_bracket_ref_index_in_text(t: str) -> int:
    mx = 0
    try:
        for m in re.finditer(r"(?:^|\n)\s*\[(\d+)\]\s*", t or ""):
            try:
                mx = max(mx, int(m.group(1)))
            except Exception:
                continue
    except Exception:
        pass
    return mx


def _max_bracket_ref_index_in_doc(doc) -> int:
    mx = 0
    try:
        for p in doc.paragraphs:
            mx = max(mx, _max_bracket_ref_index_in_text(p.text or ""))
    except Exception:
        pass
    return mx


def _looks_like_reference_entry(s: str) -> bool:
    sl = (s or "").lower()
    if not sl.strip():
        return False
    return bool(
        re.search(
            r"\b(iec|iso|en\s*iso|ieee|fda|ansi|aami|eu|mdr|ivdr|gxp|ich|usp|ep|jp|cn|gb|yy|yyt)\b",
            sl,
        )
    )


def _split_definitions_and_standards_blob(seg: str) -> List[str]:
    """将单段内多条「缩略语定义 / [n] 引用 / 法规长串」拆成多条段落文本。"""
    s = (seg or "").strip()
    if not s:
        return []
    if len(re.findall(r"\[\d+\]\s*", s)) > 1:
        bits = re.split(r"(?=\[\d+\]\s*)", s)
        return [b.strip() for b in bits if b.strip()]
    if len(re.findall(r"\brefers to\b", s, flags=re.I)) > 1:
        bits = re.split(
            r'(?i)(?<=[.!?;,\"”\)])\s+(?=[\w\[\(][\w\./\-\s]{0,48}\s+refers to\b)',
            s,
        )
        bits = [b.strip() for b in bits if b.strip()]
        if len(bits) > 1:
            return bits
    if len(s) > 400 and s.count(";") >= 2 and re.search(r"\b(IEC|ISO|EN\s*ISO|FDA|ANSI|AAMI)\b", s, re.I):
        bits = re.split(r";\s+(?=[A-Z(\[])", s)
        bits = [b.strip() for b in bits if b.strip()]
        if len(bits) > 1:
            return bits
    return [s]


def _expand_docx_insert_text_to_paragraphs(text: str) -> List[str]:
    raw = (text or "").strip()
    if not raw:
        return []
    out: List[str] = []
    for line in raw.splitlines():
        line = line.strip()
        if not line:
            continue
        out.extend(_split_definitions_and_standards_blob(line))
    return out if out else _split_definitions_and_standards_blob(raw)


def _insert_needs_bracket_renumber(chunks: List[str], heading_ctx: str) -> bool:
    """是否在插入/替换后需要对 [n] 参考文献样式续号（避免误伤「缩略语定义」章节）。"""
    h = (heading_ctx or "").lower()
    ref_like = sum(1 for c in chunks if _looks_like_reference_entry(c))
    if any(k in h for k in ("reference", "参考文献", "引用文件", "normative")):
        if ref_like >= max(1, (len(chunks) + 1) // 2):
            return True
        if any(re.match(r"^\s*\[\d+\]\s*", (c or "").strip()) for c in chunks):
            return True
        return False
    if len(chunks) >= 2 and ref_like >= 2:
        return True
    if len(chunks) >= 2 and any(re.match(r"^\s*\[\d+\]\s*", (c or "").strip()) for c in chunks):
        return True
    return False


def _normalize_reference_chunks(chunks: List[str], doc) -> List[str]:
    """去掉原 [n] 前缀后按文档当前最大序号续编。"""
    max_doc = _max_bracket_ref_index_in_doc(doc)
    bodies: List[str] = []
    for c in chunks:
        c0 = (c or "").strip()
        if not c0:
            continue
        m = re.match(r"^\s*\[(\d+)\]\s*(.*)$", c0, flags=re.S)
        bodies.append((m.group(2) if m else c0).strip())
    out: List[str] = []
    n = max_doc
    for b in bodies:
        if not b:
            continue
        n += 1
        out.append(f"[{n}] {b}")
    return out


def _copy_paragraph_layout_and_style(src, dst) -> None:
    """插入段尽量与锚点段同样式、缩进与行距（python-docx 可写属性）。"""
    try:
        dst.style = src.style
    except Exception:
        pass
    try:
        s_pf = getattr(src, "paragraph_format", None)
        d_pf = getattr(dst, "paragraph_format", None)
        if not s_pf or not d_pf:
            return
        for attr in (
            "alignment",
            "left_indent",
            "right_indent",
            "first_line_indent",
            "space_before",
            "space_after",
            "line_spacing",
            "line_spacing_rule",
            "keep_together",
            "keep_with_next",
            "page_break_before",
            "widow_control",
        ):
            try:
                setattr(d_pf, attr, getattr(s_pf, attr))
            except Exception:
                continue
    except Exception:
        pass


def _docx_revision_table_header_score(hdr_text: str) -> int:
    """根据表头文本判断更像“修订记录表”还是“封面签字/审批表”。"""
    h = (hdr_text or "").strip().lower()
    if not h:
        return -999

    # 签字/审批流（与“变更/日期”等泛词高度撞车）
    neg = [
        "author",
        "reviewer",
        "approver",
        "编制",
        "审核",
        "批准",
        "签字",
        "签名",
        "会签",
        "审定",
    ]
    has_neg = any(k in h for k in neg)

    # “修订记录”语义核心：必须至少命中其一，否则不应把仅含 version/date 的表当修订表
    core = [
        "revision",
        "history",
        "record",
        "changelog",
        "change log",
        "change-record",
        "change record",
        "change history",
        "document history",
        "修订",
        "变更",
    ]
    has_core = any(k in h for k in core)

    # 辅助字段（单独出现不足以认定）
    support = [
        "version",
        "rev",
        "日期",
        "date",
        "说明",
        "内容",
        "原因",
        "summary",
        "description",
    ]
    sup_hits = sum(1 for k in support if k in h)

    if has_neg and not has_core:
        return -1000
    score = 0
    if has_core:
        score += 10
    score += sup_hits * 2
    if has_neg and has_core:
        # 少数模板会在修订表里写“编制/审核”，但不应像封面那样强惩罚
        score -= 3
    return score


def _docx_score_revision_table_candidate(tbl) -> int:
    """给候选“修订记录表”打分；用于避免误中封面签字/审批表等。"""
    try:
        if not getattr(tbl, "rows", None) or len(tbl.rows) < 1:
            return -999
        hdr_cells = tbl.rows[0].cells
        hdr_text = " | ".join([(c.text or "").strip() for c in hdr_cells])
        sc = _docx_revision_table_header_score(hdr_text)
        if sc <= -1000:
            return sc
        # 列数过少更像签字块（修订表通常≥4列，但不绝对）
        try:
            if len(hdr_cells) <= 3:
                sc -= 3
        except Exception:
            pass
        return sc
    except Exception:
        return -999


def _docx_find_revision_history_table(doc) -> Any:
    """
    在 docx 中寻找“修订记录/Revision History/Change Record”表格。
    规则：
    - 优先找表头行包含典型列名关键字的表（中文/英文均可）
    """
    # 严格定位：优先寻找“修订记录/Revision History/Change Record”等标题段落之后紧跟的第一张表
    # 标题关键字（中英文都可能出现）
    title_kws = [
        "修订记录",
        "变更记录",
        "版本变更记录",
        "变更说明",
        "修订说明",
        "变更履历",
        "修订履历",
        "changes",
        "change log",
        "changelog",
        "document history",
        "document revision",
        "revision record",
        "revision history",
        "change record",
        "change history",
        "change-record",
    ]
    try:
        from docx.oxml.ns import qn
        from docx.table import Table
    except Exception:
        qn = None
        Table = None

    try:
        body = doc.element.body  # type: ignore[attr-defined]
        children = list(body.iterchildren())
        for i, el in enumerate(children):
            try:
                if qn and el.tag != qn("w:p"):
                    continue
            except Exception:
                # 兜底：用字符串判断
                if not str(getattr(el, "tag", "")).endswith("}p"):
                    continue
            # 段落文本
            try:
                p_text = "".join([t.text or "" for t in el.iter() if getattr(t, "text", None) is not None]).strip()
            except Exception:
                p_text = ""
            if not p_text:
                continue
            if any(k in p_text.lower() for k in [x.lower() for x in title_kws]):
                # 往后找第一张“看起来像修订记录表”的表（避免标题后紧跟封面签字表）
                for j in range(i + 1, min(i + 60, len(children))):
                    el2 = children[j]
                    try:
                        if qn and el2.tag != qn("w:tbl"):
                            continue
                    except Exception:
                        if not str(getattr(el2, "tag", "")).endswith("}tbl"):
                            continue
                    if Table is None:
                        return None
                    cand_tbl = Table(el2, doc)
                    try:
                        if not cand_tbl.rows:
                            continue
                        hdr_text0 = " | ".join([(c.text or "").strip() for c in cand_tbl.rows[0].cells])
                        if _docx_revision_table_header_score(hdr_text0) >= 9:
                            return cand_tbl
                    except Exception:
                        continue
    except Exception:
        pass

    # 兜底：按“修订表特征”打分选择（避免仅靠“变更/日期”等泛词误中封面签字表）
    try:
        best = None
        best_score = -10_000
        for tbl in list(getattr(doc, "tables", []) or []):
            sc = _docx_score_revision_table_candidate(tbl)
            if sc > best_score:
                best_score = sc
                best = tbl
        # 阈值：需要足够像“修订记录表”，否则宁可不写（由上层提示）
        if best is not None and best_score >= 9:
            return best
    except Exception:
        pass
    return None


def _xlsx_append_revision_row(wb, meta: Dict) -> bool:
    """
    在 Excel 的“修订记录”工作表里追加一行（不新建工作表）。
    严格定位：优先找包含“修订记录/Revision History/Change Record”的标题单元格，下方的表头行再追加。
    写入时按表头语义映射列（中英文），避免列顺序不同导致“插乱”。
    """
    # Excel 模板常见日期为 2024.09.26；Word 常见为 2024/10/24
    fields = _revision_entry_fields({**(meta or {}), "revision_date_format": "dot"})
    rows, _ = _rev_table_rows(meta or {})
    if not rows or len(rows) < 2:
        return False
    title_kws = [
        "修订记录",
        "变更记录",
        "版本变更记录",
        "revision record",
        "revision history",
        "change record",
        "change history",
        "change-record",
    ]
    header_kws = [
        "version",
        "date",
        "author",
        "change",
        "content",
        "revision",
        "日期",
        "版本",
        "编制",
        "变更",
        "内容",
        "说明",
    ]

    def _find_revision_sheet():
        # 1) 先按 sheet 名称猜测（中英文）
        name_kws = [
            "修订记录",
            "变更记录",
            "revision record",
            "revision history",
            "change record",
            "change history",
        ]
        try:
            for sn in list(getattr(wb, "sheetnames", []) or []):
                s0 = (sn or "").strip().lower()
                if not s0:
                    continue
                if any(k in s0 for k in [x.lower() for x in name_kws]):
                    return wb[sn]
        except Exception:
            pass
        # 2) 再扫每个 sheet 的前若干单元格找标题
        try:
            for ws0 in list(getattr(wb, "worksheets", []) or []):
                max_r = min((ws0.max_row or 0), 80) or 80
                max_c = min((ws0.max_column or 0), 20) or 20
                for r in range(1, max_r + 1):
                    for c in range(1, max_c + 1):
                        v = ws0.cell(row=r, column=c).value
                        s = "" if v is None else str(v).strip()
                        if not s:
                            continue
                        if any(k in s.lower() for k in [x.lower() for x in title_kws]):
                            return ws0
        except Exception:
            pass
        return None

    ws = _find_revision_sheet()
    if ws is None:
        return False

    title_row = None
    title_col = None
    try:
        for r in range(1, min((ws.max_row or 0) + 1, 200)):
            for c in range(1, min((ws.max_column or 0) + 1, 30)):
                v = ws.cell(row=r, column=c).value
                s = "" if v is None else str(v).strip()
                if not s:
                    continue
                if any(k in s.lower() for k in [x.lower() for x in title_kws]):
                    title_row, title_col = r, c
                    break
            if title_row:
                break
    except Exception:
        title_row = None

    header_row = None
    if title_row:
        # 标题下方 1~8 行内找表头
        for r in range(title_row + 1, min(title_row + 9, (ws.max_row or 0) + 1)):
            row_text = " | ".join(
                [str(ws.cell(row=r, column=c).value or "").strip() for c in range(1, min((ws.max_column or 0) + 1, 30))]
            ).lower()
            if any(k in row_text for k in [x.lower() for x in header_kws]):
                header_row = r
                break

    # 兜底：在前 30 行找表头
    if header_row is None:
        for r in range(1, min((ws.max_row or 0) + 1, 30)):
            row_text = " | ".join(
                [str(ws.cell(row=r, column=c).value or "").strip() for c in range(1, min((ws.max_column or 0) + 1, 30))]
            ).lower()
            if any(k in row_text for k in [x.lower() for x in header_kws]):
                header_row = r
                break

    # 若表为空：写默认表头 + 一行数据（无表头可映射时的兜底）
    if (ws.max_row or 0) < 1:
        for ci, v in enumerate(rows[0], 1):
            ws.cell(row=1, column=ci, value=str(v or ""))
        for ci, v in enumerate(rows[1], 1):
            ws.cell(row=2, column=ci, value=str(v or ""))
        return True

    max_c_scan = min((ws.max_column or 0) + 1, 40)
    if header_row is None:
        header_row = 1
    headers = [
        str(ws.cell(row=header_row, column=c).value or "").strip() for c in range(1, max_c_scan)
    ]
    # 去掉尾部连续空表头，保留中间空列（合并单元格可能导致空串）
    while headers and not (headers[-1] or "").strip():
        headers.pop()
    if not headers:
        headers = [str(x) for x in rows[0]]
    # 版本号：优先从修订记录表中读取上一版并升级（例如 A/1 -> A/2），而不是另写 draft-...
    # 兼容：部分模板表头被合并/留空导致 slot_map 识别失败时，仍尽量定位 “Version/版本” 列。
    old_ver = ""
    vcol: Optional[int] = None
    try:
        slot_map = _build_header_slot_to_col(headers)
        vcol0 = slot_map.get("version")
        if isinstance(vcol0, int):
            vcol = vcol0
        else:
            for i, raw in enumerate(headers):
                h = _norm_ws(str(raw or ""))
                if not h:
                    continue
                if any(k in h for k in ["version", "ver", "版本", "版次", "版 本"]):
                    vcol = i
                    break
        if isinstance(vcol, int):
            for rr in range(ws.max_row or 0, header_row, -1):
                s = str(ws.cell(row=rr, column=vcol + 1).value or "").strip()
                if s:
                    old_ver = s
                    break
    except Exception:
        old_ver = ""
    if old_ver:
        new_ver = _compute_next_version(old_ver)
        if new_ver and new_ver != old_ver:
            fields["version"] = new_ver
            try:
                (meta or {})["version_tag"] = new_ver
            except Exception:
                pass
    data_row = _revision_values_by_header(headers, fields)
    # 若表头映射失败（或 version 列未被识别），确保 version 至少写回到识别出的 vcol 位置
    try:
        if isinstance(vcol, int) and 0 <= vcol < len(data_row) and str(fields.get("version") or "").strip():
            if not str(data_row[vcol] or "").strip():
                data_row[vcol] = str(fields.get("version") or "")
    except Exception:
        pass

    # 计算追加行：从 header_row 起向下找最后一行有内容的记录（按表头列宽扫描）
    last = header_row
    ncols = max(len(headers), 1)
    for r in range(header_row + 1, (ws.max_row or 0) + 1):
        has_any = False
        for c in range(1, ncols + 1):
            v = ws.cell(row=r, column=c).value
            if str(v or "").strip():
                has_any = True
                break
        if has_any:
            last = r
    target_row = last + 1
    # 复制样式：优先继承上一条记录行（框线/字体/对齐等），保持与源文档一致
    try:
        src_style_row = last if last > header_row else None
        if src_style_row is None:
            # 若当前没有任何数据行，则尝试用表头下一行（有些模板预留空白样式行）
            cand = header_row + 1
            if cand <= (ws.max_row or 0):
                src_style_row = cand
        if isinstance(src_style_row, int) and src_style_row >= 1:
            try:
                ws.row_dimensions[target_row].height = ws.row_dimensions[src_style_row].height
            except Exception:
                pass
            for c in range(1, ncols + 1):
                s_cell = ws.cell(row=src_style_row, column=c)
                d_cell = ws.cell(row=target_row, column=c)
                # openpyxl style objects are immutable-ish; assign by copy
                try:
                    if getattr(s_cell, "has_style", False):
                        d_cell._style = copy.copy(s_cell._style)
                except Exception:
                    pass
                for attr in ("font", "border", "fill", "number_format", "protection", "alignment"):
                    try:
                        v = getattr(s_cell, attr, None)
                        if v is not None:
                            setattr(d_cell, attr, copy.copy(v))
                    except Exception:
                        pass
                try:
                    d_cell.comment = None
                except Exception:
                    pass
    except Exception:
        pass
    for ci, v in enumerate(data_row, 1):
        ws.cell(row=target_row, column=ci, value=str(v or ""))

    # 同步更新封面/页眉等版本信息（Excel 无页眉对象时，至少同步表格区）
    try:
        if old_ver and fields.get("version") and old_ver != fields.get("version"):
            _xlsx_sync_version_everywhere(wb, old=old_ver, new=str(fields.get("version") or ""))
    except Exception:
        pass
    return True


def _docx_append_revision_row(doc, meta: Dict, *, track_changes: bool = True) -> bool:
    """
    向基础文档已有的修订记录表追加一行（不新建章节/表）。
    按第一行表头语义映射列；可选对单元格写入跟踪修订（w:del/w:ins），便于在 Word 修订记录中看到插入。
    找不到修订记录表则返回 False。
    """
    if track_changes:
        _enable_track_revisions(doc)
    fields = _revision_entry_fields(meta or {})
    rows, _version_tag = _rev_table_rows(meta or {})
    if not rows or len(rows) < 2:
        return False
    tbl = _docx_find_revision_history_table(doc)
    if tbl is None:
        return False
    try:
        if not tbl.rows:
            return False
        # 有的模板修订表可能有“标题行/空行”在最上方，不一定 row0 就是表头。
        # 这里扫描前几行，选“可映射槽位最多”的那一行作为表头行。
        hdr_row_idx = 0
        headers: List[str] = []
        best_hits = -1
        scan_n = min(6, len(tbl.rows))
        for ri in range(scan_n):
            try:
                hs = [(c.text or "").strip() for c in tbl.rows[ri].cells]
            except Exception:
                continue
            slot_map0 = _build_header_slot_to_col(hs)
            hits = len(slot_map0)
            # 需要至少命中 version/date/change_content 之一，否则可能只是空白/分隔行
            if hits > best_hits and any(k in slot_map0 for k in ("version", "date", "change_content", "author")):
                best_hits = hits
                hdr_row_idx = ri
                headers = hs
        if not headers:
            hdr_row_idx = 0
            headers = [(c.text or "").strip() for c in tbl.rows[0].cells]
        slot_map = _build_header_slot_to_col(headers)
        # 日期格式：自动跟随现有修订记录列（例如 2026.04.13 / 2026/04/13）
        try:
            dcol = slot_map.get("date")
            if isinstance(dcol, int):
                sample = ""
                for rr in reversed(list(tbl.rows[hdr_row_idx + 1 :])):  # 跳过表头
                    try:
                        s = (rr.cells[dcol].text or "").strip()
                    except Exception:
                        s = ""
                    if s:
                        sample = s
                        break
                if sample:
                    if "." in sample and "/" not in sample:
                        fields["date"] = _dt.datetime.now().strftime("%Y.%m.%d")
                    elif "/" in sample:
                        fields["date"] = _dt.datetime.now().strftime("%Y/%m/%d")
        except Exception:
            pass

        # 版本号：优先从修订记录表中读取上一版并升级（例如 A/1 -> A/2），而不是另写 draft-...
        old_ver = ""
        try:
            vcol = slot_map.get("version")
            if isinstance(vcol, int):
                for rr in reversed(list(tbl.rows[hdr_row_idx + 1 :])):  # 跳过表头
                    try:
                        s = (rr.cells[vcol].text or "").strip()
                    except Exception:
                        s = ""
                    if s:
                        old_ver = s
                        break
        except Exception:
            old_ver = ""
        if old_ver:
            new_ver = _compute_next_version(old_ver)
            if new_ver and new_ver != old_ver:
                fields["version"] = new_ver
                try:
                    (meta or {})["version_tag"] = new_ver
                except Exception:
                    pass

        data_row = _revision_values_by_header(headers, fields)
        if not slot_map:
            # 表头无法映射时：Word 修订记录表常见顺序为 Version/Date/Author/Description（4列）或 +ChangeNo（5列）
            n = len(headers)
            if n == 4:
                data_row = [
                    str(fields.get("version") or ""),
                    str(fields.get("date") or ""),
                    str(fields.get("author") or ""),
                    str(fields.get("change_content") or ""),
                ]
            elif n == 5:
                data_row = [
                    str(fields.get("version") or ""),
                    str(fields.get("date") or ""),
                    str(fields.get("author") or ""),
                    str(fields.get("change_no") or ""),
                    str(fields.get("change_content") or ""),
                ]

        # 追加行：优先克隆最后一条数据行（或表头下一行）以继承框线/字体/合并结构
        last_data_idx = -1
        data_start = min(len(tbl.rows), hdr_row_idx + 1)
        for i in range(len(tbl.rows) - 1, data_start - 1, -1):
            try:
                if any((c.text or "").strip() for c in tbl.rows[i].cells):
                    last_data_idx = i
                    break
            except Exception:
                continue
        anchor_row = None
        if last_data_idx >= 0:
            anchor_row = tbl.rows[last_data_idx]
        elif hdr_row_idx + 1 < len(tbl.rows):
            anchor_row = tbl.rows[hdr_row_idx + 1]
        else:
            anchor_row = tbl.rows[hdr_row_idx]

        r = _word_clone_row_after(tbl, anchor_row) if anchor_row is not None else None
        if r is None:
            r = tbl.add_row()

        # 先清空（避免模板行自带占位符/旧内容），再按列写入
        try:
            for c in list(getattr(r, "cells", []) or []):
                try:
                    if track_changes:
                        _replace_table_cell_with_track_changes(c, (c.text or ""), "")
                    else:
                        c.text = ""
                except Exception:
                    pass
        except Exception:
            pass

        n_cell = min(len(r.cells), len(data_row))
        for ci in range(n_cell):
            try:
                cell = r.cells[ci]
                val = str(data_row[ci] or "")
                if track_changes and val.strip():
                    _replace_table_cell_with_track_changes(cell, "", val)
                else:
                    cell.text = val
            except Exception:
                continue
        # 同步更新封面/页眉页脚中的版本号（若检测到升级）
        try:
            if old_ver and fields.get("version") and old_ver != fields.get("version"):
                _docx_sync_version_everywhere(doc, old=old_ver, new=str(fields.get("version") or ""))
        except Exception:
            pass
        return True
    except Exception:
        return False


def sniff_word_processing_suffix(path: str | Path) -> str:
    """用于分支判断的 Word 后缀：.doc/.dot 若实为 OOXML（zip），按 .docx 处理；.docm 等走 OOXML 分支。"""
    p = Path(path)
    suf = p.suffix.lower()
    if suf in (".docx", ".docm", ".dotx", ".dotm"):
        return ".docx"
    if suf in (".doc", ".dot"):
        try:
            with open(p, "rb") as _bf:
                if _bf.read(4) == b"PK\x03\x04":
                    return ".docx"
        except Exception:
            pass
    return suf


def normalize_flat_txt_output_path(out: Path) -> Path:
    """
    将「纯文本降级输出」落盘路径规范为单一「.txt」扩展名。

    若误将 out_path 拼成「xxx.doc.txt」，pathlib 的 with_suffix('.txt') 不会改变路径
    （因为最后一个 suffix 已是 .txt），会写出畸形的 *.doc.txt 文件。
    此处反复剥掉 stem 末尾的 Office 类扩展名，直到得到稳定主名后再加 .txt。
    """
    parent = out.parent
    stem = out.stem
    while True:
        s = (stem or "").lower()
        hit = False
        for ext in (".docm", ".dotm", ".docx", ".dotx", ".doc", ".dot", ".rtf", ".wps"):
            if s.endswith(ext):
                stem = Path(stem).stem
                hit = True
                break
        if not hit:
            break
    safe = (stem or "output").strip() or "output"
    return parent / f"{safe}.txt"


def export_like_base(
    *,
    base_file_path: str,
    out_path: str,
    title: str,
    content_text: str,
    meta: Optional[Dict] = None,
) -> str:
    """
    按基础文件后缀导出同格式文件，并写入修订记录。
    返回 out_path。
    """
    meta = meta or {}
    base = Path(base_file_path)
    suffix = sniff_word_processing_suffix(base)
    out = Path(out_path)
    if suffix == ".docx" and out.suffix.lower() not in (".docx", ".docm", ".dotx", ".dotm"):
        out = out.with_suffix(".docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    rows, version_tag = _rev_table_rows(meta)

    if suffix == ".docx":
        from docx import Document

        # 以 base 为主：最大化保留原有格式/页眉页脚/表格/样式
        shutil.copyfile(str(base), str(out))
        doc = Document(str(out))
        _enable_update_fields_on_open(doc)

        # 修订记录：基础文档通常已自带修订记录章节/表。
        # 这里不新建“修订记录”章节，只在已有修订记录表中追加一行；找不到则跳过。
        try:
            _docx_append_revision_row(doc, meta)
        except Exception:
            pass

        doc.add_paragraph("")
        doc.add_heading("自动生成内容（附录）", level=1)
        doc.add_paragraph(f"标题：{title or out.stem}")
        doc.add_paragraph(f"版本：{version_tag}")
        doc.add_paragraph("")
        for line in (content_text or "").splitlines():
            # 附录按行写入，避免破坏原模板段落样式
            doc.add_paragraph(line)

        doc.save(str(out))
        return str(out)

    if suffix in (".xlsx", ".xls"):
        import openpyxl

        # 以 base 为主：复制原表格文件，避免格式/公式/样式丢失
        shutil.copyfile(str(base), str(out))
        wb = openpyxl.load_workbook(str(out))

        # 修订记录：不新建工作表。严格定位到“修订记录标题下的那张表”并追加一行；找不到则跳过。
        try:
            _xlsx_append_revision_row(wb, meta)
        except Exception:
            pass

        # 附录：写入自动生成内容（避免破坏原工作表）
        appendix_name = "自动生成内容"
        if appendix_name in wb.sheetnames:
            ws = wb[appendix_name]
            ws.delete_rows(1, ws.max_row or 1)
        else:
            ws = wb.create_sheet(appendix_name)
        ws.cell(row=1, column=1, value=title or out.stem)
        ws.cell(row=2, column=1, value=f"版本：{version_tag}")
        for i, line in enumerate((content_text or "").splitlines(), start=4):
            ws.cell(row=i, column=1, value=line)

        wb.save(str(out))
        return str(out)

    if suffix == ".pdf":
        # PDF 受限于“可编辑内容流/字体/版面”的复杂度：不重建，直接复制 base，
        # 并额外输出同名的 .draft.txt 供核对。
        shutil.copyfile(str(base), str(out))
        sidecar = out.with_suffix(out.suffix + ".draft.txt")
        sidecar.write_text((content_text or ""), encoding="utf-8")
        return str(out)

    # 其他格式：降级为 txt（仍保证可下载）。
    # 注意：不可用 out.suffix+".txt"；且当 out 误为「xxx.doc.txt」时 with_suffix('.txt') 无效，须 normalize。
    out_txt = normalize_flat_txt_output_path(out)
    out_txt.write_text((content_text or ""), encoding="utf-8")
    return str(out_txt)


def export_docx_inplace_patch(
    *,
    base_file_path: str,
    out_path: str,
    patch_json: str,
    meta: Optional[Dict] = None,
    track_changes: bool = True,
) -> Tuple[str, Dict[str, Any]]:
    """
    对 docx 按结构化 patch 做可定位修改（不整篇重建文档）：
    - 复制 base 到 out
    - 按 patch_json 中的 operations 定位并替换/插入/删除段落或表格单元格文本
    - 追加修订记录

    返回 (out_path, report)
    """
    meta = meta or {}
    base = Path(base_file_path)
    out = Path(out_path)
    if sniff_word_processing_suffix(base) == ".docx" and out.suffix.lower() not in (
        ".docx",
        ".docm",
        ".dotx",
        ".dotm",
    ):
        out = out.with_suffix(".docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    eff_base = base
    conv_tmp: Optional[Path] = None
    _sk_base = sniff_word_processing_suffix(base)
    if _sk_base != ".docx" and base.suffix.lower() in (".doc", ".dot"):
        conv_tmp = out.parent / f"_aicw_legacy_{uuid.uuid4().hex}.docx"
        from src.core.word_legacy_convert import convert_binary_word_to_docx

        _ok, _msg = convert_binary_word_to_docx(
            src_path=str(base.resolve()),
            dst_path=str(conv_tmp.resolve()),
        )
        if not (_ok and conv_tmp.is_file()):
            report: Dict[str, Any] = {"applied": [], "skipped": [], "errors": [], "changes": []}
            report["errors"].append(
                {"error": (_msg or "二进制 Word 需先转为 docx 才能就地修改：转换失败").strip()}
            )
            if conv_tmp.is_file():
                try:
                    conv_tmp.unlink()
                except OSError:
                    pass
            return str(out), report
        eff_base = conv_tmp

    from docx import Document

    try:
        shutil.copyfile(str(eff_base), str(out))
    finally:
        if conv_tmp is not None and conv_tmp.is_file() and conv_tmp.resolve() == eff_base.resolve():
            try:
                conv_tmp.unlink()
            except OSError:
                pass

    doc = Document(str(out))
    if track_changes:
        _enable_track_revisions(doc)
    _enable_update_fields_on_open(doc)

    report: Dict[str, Any] = {"applied": [], "skipped": [], "errors": [], "changes": []}

    try:
        patch_obj = json.loads(patch_json or "{}")
    except Exception as e:
        report["errors"].append({"error": f"patch_json 解析失败：{e}"})
        # 解析失败：不做任何正文替换，也不新增修订记录
        try:
            doc.save(str(out))
        except Exception:
            pass
        return str(out), report

    ops: List[Dict[str, Any]] = list(patch_obj.get("operations") or [])
    tc_rules = _compile_tc_id_rules((patch_obj or {}).get("tc_id_rules") or (meta or {}).get("tc_id_rules"))
    def _change_meta(op_obj: Dict[str, Any]) -> Dict[str, Any]:
        refs = op_obj.get("audit_point_refs")
        if not isinstance(refs, list):
            refs = []
        refs = [str(x).strip() for x in refs if str(x).strip()]
        out = {"audit_point_refs": refs}
        note = str(op_obj.get("audit_point_note") or "").strip()
        if note:
            out["audit_point_note"] = note
        return out

    def _para_heading_context(p_idx: int) -> str:
        # 粗定位“章节标题”：向上找最近的标题样式段落
        try:
            for j in range(p_idx, -1, -1):
                p = doc.paragraphs[j]
                name = ""
                try:
                    name = (p.style.name or "")
                except Exception:
                    name = ""
                if name and ("Heading" in name or "标题" in name):
                    return (p.text or "").strip()[:200]
        except Exception:
            pass
        return ""

    def _find_duplicate_followup_paragraph(p_idx: int, new_text: str, max_scan: int = 30):
        """在锚点段落后扫描，若已存在高度相似段落则返回该段落（避免重复插入）。"""
        try:
            n = len(doc.paragraphs)
            hi = min(n - 1, p_idx + max(1, int(max_scan)))
            for j in range(p_idx + 1, hi + 1):
                try:
                    p2 = doc.paragraphs[j]
                except Exception:
                    continue
                if _looks_like_duplicate_followup_paragraph(new_text, p2.text or ""):
                    return j, p2
        except Exception:
            return None
        return None

    def _replace_paragraph(p, new_text: str) -> None:
        # 保留段落样式，只替换文本 runs
        for r in list(p.runs):
            try:
                r.text = ""
            except Exception:
                pass
        if p.runs:
            p.runs[0].text = new_text or ""
        else:
            p.add_run(new_text or "")

    def _insert_paragraph_after(p, text: str) -> Any:
        """
        在段落 p 后插入一个新段落，返回新段落对象。
        python-docx 没有公开 API，这里用底层 XML 插入，尽量保留邻近结构与样式。
        """
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        p._p.addnext(new_p)
        new_para = Paragraph(new_p, p._parent)
        if text:
            new_para.add_run(text)
        return new_para

    def _find_paragraphs_containing(anchor: str):
        a = (anchor or "").strip()
        if not a:
            return []
        hits = []
        for idx, p in enumerate(doc.paragraphs):
            if a in (p.text or ""):
                hits.append((idx, p))
        return hits

    def _find_table_cells_containing(anchor: str):
        a = (anchor or "").strip()
        if not a:
            return []
        hits = []
        for t in doc.tables:
            for r_idx, row in enumerate(t.rows):
                for c_idx, cell in enumerate(row.cells):
                    if a in (cell.text or ""):
                        hits.append((t, r_idx, c_idx, cell))
        return hits

    def _find_table_rows_containing(anchor: str):
        a = (anchor or "").strip()
        if not a:
            return []
        hits = []
        for t in doc.tables:
            for r_idx, row in enumerate(t.rows):
                for cell in row.cells:
                    if a in (cell.text or ""):
                        hits.append((t, r_idx, row))
                        break
        return hits

    for op in ops:
        try:
            t = (op.get("type") or "").strip()
            anchor = (op.get("anchor") or "").strip()
            new_text = op.get("new_text")
            if new_text is None:
                new_text = ""
            new_text = str(new_text)
            require_unique = op.get("require_unique")
            if require_unique is None:
                require_unique = True
            require_unique = bool(require_unique)

            if not t or not anchor:
                report["skipped"].append({"op": op, "reason": "缺少 type 或 anchor"})
                continue

            if t == "replace_paragraph_contains":
                hits = _find_paragraphs_containing(anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中段落"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"段落命中不唯一：{len(hits)}"})
                    continue
                for p_idx, p in hits:
                    run_rpr = _paragraph_effective_run_rpr(p)
                    chunks = _expand_docx_insert_text_to_paragraphs(new_text)
                    heading_ctx = _para_heading_context(p_idx)
                    if _insert_needs_bracket_renumber(chunks, heading_ctx):
                        chunks = _normalize_reference_chunks(chunks, doc)
                    if not chunks:
                        report["skipped"].append({"op": op, "reason": "展开后 new_text 为空"})
                        continue
                    joined_after = "\n".join(chunks)
                    before = (p.text or "")
                    if (before or "").strip() == joined_after.strip():
                        report["skipped"].append({"op": op, "reason": "no-op：新旧文本一致（段落无需修改）"})
                        continue
                    first, rest = chunks[0], chunks[1:]
                    if track_changes:
                        _replace_paragraph_with_track_changes(p, before, first, run_rpr=run_rpr)
                    else:
                        _replace_paragraph(p, first)
                    last_p = p
                    for ch in rest:
                        new_p = _insert_paragraph_after(last_p, "" if track_changes else ch)
                        _copy_paragraph_layout_and_style(p, new_p)
                        if track_changes:
                            _insert_paragraph_with_track_changes(new_p, ch, run_rpr=run_rpr)
                        last_p = new_p
                    # 使用 track changes 时，python-docx 的 p.text 可能读不到 w:ins 文本，导致 after 为空
                    after = joined_after
                    report["changes"].append(
                        {
                            "type": t,
                            "anchor": anchor,
                            "heading": heading_ctx,
                            "paragraph_index": p_idx,
                            "before": before,
                            "after": after,
                            "paragraphs_written": len(chunks),
                            **_change_meta(op),
                        }
                    )
                report["applied"].append({"op": op, "hits": len(hits)})
                continue

            if t == "delete_paragraph_contains":
                hits = _find_paragraphs_containing(anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中段落（用于删除）"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"删除锚点命中不唯一：{len(hits)}"})
                    continue
                for p_idx, p in hits:
                    before = (p.text or "")
                    if track_changes:
                        _delete_paragraph_with_track_changes(p, before)
                    else:
                        _replace_paragraph(p, "")
                    after = "" if track_changes else (p.text or "")
                    report["changes"].append(
                        {
                            "type": t,
                            "anchor": anchor,
                            "heading": _para_heading_context(p_idx),
                            "paragraph_index": p_idx,
                            "before": before,
                            "after": after,
                            **_change_meta(op),
                        }
                    )
                report["applied"].append({"op": op, "hits": len(hits)})
                continue

            if t == "replace_table_cell_contains":
                hits = _find_table_cells_containing(anchor)
                hits = _word_pick_value_cell_hits(hits, anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中表格单元格"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"单元格命中不唯一：{len(hits)}"})
                    continue
                op_before = (op.get("before") or "").strip()
                for _tbl, _r_idx, _c_idx, c in hits:
                    before = (c.text or "")
                    before_s = str(before or "")
                    # 支持子串替换：避免多行单元格“整格覆盖”导致丢步骤
                    if op_before and op_before in before_s:
                        after_s = before_s.replace(op_before, new_text, 1)
                        if after_s == before_s:
                            report["skipped"].append(
                                {"op": op, "reason": "no-op：before 子串替换未改变单元格"}
                            )
                            continue
                        if track_changes:
                            _replace_table_cell_with_track_changes(c, before_s, after_s)
                        else:
                            c.text = after_s
                        after = after_s
                    else:
                        if before_s.strip() == (new_text or "").strip():
                            report["skipped"].append(
                                {"op": op, "reason": "no-op：新旧文本一致（单元格无需修改）"}
                            )
                            continue
                        if track_changes:
                            _replace_table_cell_with_track_changes(c, before_s, new_text)
                        else:
                            c.text = new_text
                        # 同理：track changes 下 cell.text 可能为空，改为使用 new_text 作为 after 展示
                        after = new_text if track_changes else (c.text or "")
                    report["changes"].append(
                        {
                            "type": t,
                            "anchor": anchor,
                            "before": before,
                            "after": after,
                            **_change_meta(op),
                        }
                    )
                report["applied"].append({"op": op, "hits": len(hits)})
                continue

            if t == "insert_paragraph_after_contains":
                hits = _find_paragraphs_containing(anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中段落（用于插入）"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"插入锚点命中不唯一：{len(hits)}"})
                    continue
                for p_idx, p in hits:
                    run_rpr = _paragraph_effective_run_rpr(p)
                    chunks = _expand_docx_insert_text_to_paragraphs(new_text)
                    heading_ctx = _para_heading_context(p_idx)
                    if _insert_needs_bracket_renumber(chunks, heading_ctx):
                        chunks = _normalize_reference_chunks(chunks, doc)
                    if not chunks:
                        report["skipped"].append({"op": op, "reason": "展开后 new_text 为空"})
                        continue
                    joined_after = "\n".join(chunks)
                    dup = _find_duplicate_followup_paragraph(p_idx, new_text, max_scan=35)
                    if dup is not None:
                        dup_idx, dup_p = dup
                        run_rpr_d = _paragraph_effective_run_rpr(dup_p) or run_rpr
                        before_dup = dup_p.text or ""
                        if (before_dup or "").strip() == joined_after.strip():
                            report["skipped"].append(
                                {
                                    "op": op,
                                    "reason": f"后续段落已含同类内容（去重）：paragraph_index={dup_idx}",
                                }
                            )
                            continue
                        first, rest = chunks[0], chunks[1:]
                        if track_changes:
                            _replace_paragraph_with_track_changes(
                                dup_p, before_dup, first, run_rpr=run_rpr_d
                            )
                        else:
                            _replace_paragraph(dup_p, first)
                        last_pd = dup_p
                        for ch in rest:
                            new_p = _insert_paragraph_after(last_pd, "" if track_changes else ch)
                            _copy_paragraph_layout_and_style(p, new_p)
                            if track_changes:
                                _insert_paragraph_with_track_changes(new_p, ch, run_rpr=run_rpr_d)
                            last_pd = new_p
                        report["changes"].append(
                            {
                                "type": "replace_paragraph_contains",
                                "anchor": anchor,
                                "heading": _para_heading_context(dup_idx),
                                "paragraph_index": dup_idx,
                                "before": before_dup,
                                "after": joined_after,
                                "paragraphs_written": len(chunks),
                                "note": "由 insert_paragraph_after_contains 去重改写后续段落",
                                **_change_meta(op),
                            }
                        )
                        continue

                    last_p = p
                    for ch in chunks:
                        new_p = _insert_paragraph_after(last_p, "" if track_changes else ch)
                        _copy_paragraph_layout_and_style(p, new_p)
                        if track_changes:
                            try:
                                _insert_paragraph_with_track_changes(new_p, ch, run_rpr=run_rpr)
                            except Exception:
                                pass
                        last_p = new_p
                    report["changes"].append(
                        {
                            "type": t,
                            "anchor": anchor,
                            "heading": heading_ctx,
                            "paragraph_index": p_idx,
                            "before": "",
                            "after": joined_after,
                            "paragraphs_inserted": len(chunks),
                            **_change_meta(op),
                        }
                    )
                report["applied"].append({"op": op, "hits": len(hits)})
                continue

            if t == "insert_table_row_after_contains":
                hits = _find_table_rows_containing(anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中表格行（用于插入）"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"插入表格行锚点命中不唯一：{len(hits)}"})
                    continue
                insert_ok = 0
                for tbl, r_idx, row in hits:
                    try:
                        # 列数：避免表头合并单元格导致“多列同一格”
                        tmpl_row = _word_best_template_row_for_insert(tbl, int(r_idx)) or row
                        col_n = max(len(_word_unique_cells(tmpl_row)), 1)
                        vals = [x for x in (new_text or "").split("\t")]
                        # Word：模型若未按 \\t 分列，会导致整行写入首格；这里按列数做尽力拆分兜底
                        if len(vals) == 1 and "\t" not in (new_text or ""):
                            vals2 = _xlsx_split_row_values_for_insert(new_text or "", col_n)
                            if len(vals2) > 1:
                                vals = vals2
                        if not vals or not any(str(x).strip() for x in vals):
                            report["skipped"].append(
                                {
                                    "op": op,
                                    "reason": "new_text 按制表符分列后无有效单元格内容（仅空白/制表符），未插入行",
                                }
                            )
                            continue

                        new_joined = "\t".join((vals[ci] if ci < len(vals) else "") for ci in range(col_n)).strip()
                        parsed = _parse_tc_id_first_cell(vals[0] or "", tc_rules)
                        if not parsed:
                            try:
                                if row.cells:
                                    parsed = _parse_tc_id_first_cell(row.cells[0].text or "", tc_rules)
                            except Exception:
                                parsed = None

                        # 无可解析编号时：在锚点附近窗口做强相似去重，避免重复插入“源文档已存在”的行
                        if not parsed:
                            ex2 = _word_find_similar_row_idx(
                                tbl,
                                joined=new_joined,
                                col_n=col_n,
                                start=max(0, int(r_idx) - 2),
                                end=min(len(getattr(tbl, "rows", []) or []), int(r_idx) + 12),
                            )
                            if ex2 >= 0:
                                ex_row2 = tbl.rows[ex2]
                                before_joined2 = _word_row_joined(ex_row2, col_n)
                                written_cells2: List[str] = []
                                ex_cells2 = _word_unique_cells(ex_row2)
                                for ci in range(min(col_n, len(ex_cells2))):
                                    v = vals[ci] if ci < len(vals) else ""
                                    before = (ex_cells2[ci].text or "") if ci < len(ex_cells2) else ""
                                    if track_changes:
                                        _replace_table_cell_with_track_changes(ex_cells2[ci], before, v)
                                    else:
                                        ex_cells2[ci].text = v
                                    written_cells2.append(v)
                                report["changes"].append(
                                    {
                                        "type": "update_table_row_in_place",
                                        "anchor": anchor,
                                        "table_row_index": ex2,
                                        "before": before_joined2,
                                        "after": "\t".join(written_cells2),
                                        "cells_preview": [str(x).strip() for x in written_cells2],
                                        "note": "检测到源文档附近已存在高度相似行，已原地更新（避免重复插入）",
                                        **_change_meta(op),
                                    }
                                )
                                insert_ok += 1
                                continue

                        # 首列可解析为用例编号（如 GN3-22）：按表内同前缀最大编号 +1 分配；已存在编号则按整行相似度决定更新或新行
                        if parsed:
                            prefix, _n, rule = parsed
                            max_n = _table_max_tc_suffix_for_prefix_word(tbl, prefix, tc_rules)
                            next_id = _next_unique_tc_id_word(
                                tbl, prefix, rule, tc_rules, max_n + 1
                            )
                            id_cell = (vals[0] or "").strip()
                            ex_idx = _table_find_row_idx_by_id_any_cell(tbl, id_cell)

                            if ex_idx >= 0:
                                ex_row = tbl.rows[ex_idx]
                                before_joined = _word_row_joined(ex_row, col_n)
                                sim = SequenceMatcher(
                                    None, new_joined, (before_joined or "").strip()
                                ).ratio()
                                if sim >= _TC_ROW_SIMILARITY_UPDATE_THRESHOLD:
                                    written_cells: List[str] = []
                                    ex_cells = _word_unique_cells(ex_row)
                                    for ci in range(min(col_n, len(ex_cells))):
                                        v = vals[ci] if ci < len(vals) else ""
                                        before = (ex_cells[ci].text or "") if ci < len(ex_cells) else ""
                                        if track_changes:
                                            _replace_table_cell_with_track_changes(
                                                ex_cells[ci], before, v
                                            )
                                        else:
                                            ex_cells[ci].text = v
                                        written_cells.append(v)
                                    joined = "\t".join(written_cells)
                                    report["changes"].append(
                                        {
                                            "type": "update_table_row_in_place",
                                            "anchor": anchor,
                                            "table_row_index": ex_idx,
                                            "tc_id": id_cell,
                                            "similarity": round(sim, 4),
                                            "before": before_joined,
                                            "after": joined,
                                            "cells_preview": [str(x).strip() for x in written_cells],
                                            "note": (
                                                f"首列编号已存在且整行相似度≥{_TC_ROW_SIMILARITY_UPDATE_THRESHOLD}，已原地更新该行"
                                            ),
                                            **_change_meta(op),
                                        }
                                    )
                                    insert_ok += 1
                                    continue
                                vals[0] = next_id
                            else:
                                vals[0] = next_id

                        # 插入新行：优先克隆锚点行（保留格式/结构），失败则回退 add_row
                        # 注意：不要克隆表头（可能合并单元格）；用“模板行（更像数据行）”来克隆
                        new_row = _word_clone_row_after(tbl, tmpl_row)
                        if new_row is None:
                            new_row = tbl.add_row()
                        new_cells = _word_unique_cells(new_row)
                        for ci in range(min(col_n, len(new_cells))):
                            v = vals[ci] if ci < len(vals) else ""
                            if track_changes:
                                _replace_table_cell_with_track_changes(new_cells[ci], "", v)
                            else:
                                new_cells[ci].text = v
                        # add_row 回退时需要把新行移动到锚点后；克隆插入已在 clone 内完成
                        try:
                            if getattr(new_row, "_tr", None) is not None and new_row is not None and (new_row is not row):
                                if new_row._tr.getparent() is not None and new_row._tr.getprevious() is not row._tr:
                                    row._tr.addnext(new_row._tr)
                        except Exception:
                            pass
                        written_cells = [(vals[ci] if ci < len(vals) else "") for ci in range(col_n)]
                        joined = "\t".join(written_cells)
                        ch: Dict[str, Any] = {
                            "type": t,
                            "anchor": anchor,
                            "table_row_index": r_idx,
                            "before": "",
                            "after": joined,
                            "cells_preview": [str(x).strip() for x in written_cells],
                        }
                        if parsed:
                            ch["tc_id_assigned"] = (vals[0] or "").strip()
                            ch["note"] = (
                                "首列已按表内同模块（前缀）最大编号 +1 分配，避免与已有编号冲突"
                            )
                        ch.update(_change_meta(op))
                        report["changes"].append(ch)
                        insert_ok += 1
                    except Exception as _ie:
                        report["errors"].append({"op": op, "error": f"插入表格行失败：{_ie}"})
                if insert_ok:
                    report["applied"].append({"op": op, "hits": insert_ok})
                continue

            if t == "replace_diagram_after_paragraph_contains":
                # 以 anchor 定位到“图题/图注/逻辑结构”附近段落，然后替换/插入一张模块框图
                hits = _find_paragraphs_containing(anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中段落（用于替换图示）"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"图示锚点命中不唯一：{len(hits)}"})
                    continue
                mods = [x.strip() for x in (new_text or "").replace("\n", ",").split(",") if x.strip()]
                if not mods:
                    report["skipped"].append({"op": op, "reason": "缺少模块列表 new_text（用于生成图示）"})
                    continue
                try:
                    tmp_png = Path(tempfile.mkdtemp(prefix="aicheckword_diagram_")) / "diagram.png"
                    _render_modules_diagram_png(title="Logical structure", modules=mods, out_path=str(tmp_png))
                    p_idx, _p = hits[0]
                    ok = _docx_replace_or_insert_image_after_paragraph(doc, anchor_p_idx=p_idx, image_path=str(tmp_png))
                    if not ok:
                        report["skipped"].append({"op": op, "reason": "图示替换/插入失败（未找到可写入位置）"})
                        continue
                    report["changes"].append(
                        {
                            "type": t,
                            "anchor": anchor,
                            "heading": _para_heading_context(p_idx),
                            "before": "",
                            "after": f"[diagram modules={len(mods)}]",
                            **_change_meta(op),
                        }
                    )
                    report["applied"].append({"op": op, "hits": 1})
                except Exception as e:
                    report["errors"].append({"op": op, "error": f"生成/插入图示失败：{e}"})
                continue

            report["skipped"].append({"op": op, "reason": f"不支持的 type：{t}"})
        except Exception as e:
            report["errors"].append({"op": op, "error": str(e)})

    # 修订记录：基础文档通常已自带修订记录章节/表。
    # 仅当本次确实写入了修改时，向基础文档已有修订记录表追加一行；不新建章节/表。
    try:
        if (len(report.get("applied") or []) > 0) or (len(report.get("changes") or []) > 0):
            _ok = _docx_append_revision_row(doc, meta, track_changes=bool(track_changes))
            if not _ok:
                report["skipped"].append({"reason": "未找到基础文档中的修订记录表，已跳过写入修订记录。"})
    except Exception as e:
        report["errors"].append({"error": f"写入修订记录失败：{e}"})

    doc.save(str(out))
    return str(out), report


def export_xlsx_inplace_patch(
    *,
    base_file_path: str,
    out_path: str,
    patch_json: str,
    meta: Optional[Dict] = None,
) -> Tuple[str, Dict[str, Any]]:
    """
    对 xlsx/xls 按结构化 patch 做可定位修改（openpyxl，不整表重建）：
    - 复制 base 到 out（保留格式/公式尽量不动）
    - 按 patch_json 定位并替换单元格、插入行
    - 写入“修订记录”工作表（复用 export_like_base 的表格格式）

    注：Excel 本身的“修订/跟踪更改”能力与 docx 不同，这里提供可审计的变更报告与修订记录表。
    """
    meta = meta or {}
    base = Path(base_file_path)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    report: Dict[str, Any] = {"applied": [], "skipped": [], "errors": [], "changes": []}
    try:
        patch_obj = json.loads(patch_json or "{}")
    except Exception as e:
        report["errors"].append({"error": f"patch_json 解析失败：{e}"})
        # 解析失败：仅复制基础文件，不新增修订记录
        shutil.copyfile(str(base), str(out))
        return str(out), report

    import openpyxl

    shutil.copyfile(str(base), str(out))
    wb = openpyxl.load_workbook(str(out))

    ops: List[Dict[str, Any]] = list((patch_obj or {}).get("operations") or [])
    tc_rules = _compile_tc_id_rules((patch_obj or {}).get("tc_id_rules") or (meta or {}).get("tc_id_rules"))
    def _change_meta(op_obj: Dict[str, Any]) -> Dict[str, Any]:
        refs = op_obj.get("audit_point_refs")
        if not isinstance(refs, list):
            refs = []
        refs = [str(x).strip() for x in refs if str(x).strip()]
        out = {"audit_point_refs": refs}
        note = str(op_obj.get("audit_point_note") or "").strip()
        if note:
            out["audit_point_note"] = note
        return out

    def _iter_cells():
        for ws in wb.worksheets:
            # 跳过我们自己写入的附录/修订表，避免二次匹配
            if ws.title in ("修订记录", "自动生成内容"):
                continue
            for row in ws.iter_rows():
                for cell in row:
                    yield ws, cell

    def _find_cells_containing(anchor: str):
        a = (anchor or "").strip()
        if not a:
            return []
        hits = []
        for ws, cell in _iter_cells():
            try:
                v = cell.value
            except Exception:
                v = None
            s = "" if v is None else str(v)
            if a in s:
                hits.append((ws, cell, s))
        return hits

    def _find_rows_containing(anchor: str):
        a = (anchor or "").strip()
        if not a:
            return []
        hits = []
        for ws in wb.worksheets:
            if ws.title in ("修订记录", "自动生成内容"):
                continue
            for r_idx, row in enumerate(ws.iter_rows(values_only=False), start=1):
                found = False
                for cell in row:
                    try:
                        v = cell.value
                    except Exception:
                        v = None
                    s = "" if v is None else str(v)
                    if a in s:
                        found = True
                        break
                if found:
                    hits.append((ws, r_idx))
        return hits

    for op in ops:
        try:
            t = (op.get("type") or "").strip()
            anchor = (op.get("anchor") or "").strip()
            new_text = op.get("new_text")
            if new_text is None:
                new_text = ""
            new_text = str(new_text)
            require_unique = op.get("require_unique")
            if require_unique is None:
                require_unique = True
            require_unique = bool(require_unique)

            if not t:
                report["skipped"].append({"op": op, "reason": "缺少 type"})
                continue
            sh0 = (op.get("sheet") or "").strip()
            cell0 = (op.get("cell") or "").strip()
            if t == "replace_table_cell_contains":
                if not anchor and not (sh0 and cell0):
                    report["skipped"].append(
                        {"op": op, "reason": "replace_table_cell_contains 需要 anchor，或同时提供 sheet+cell"}
                    )
                    continue
            elif not anchor:
                report["skipped"].append({"op": op, "reason": "缺少 anchor"})
                continue

            if t == "replace_table_cell_contains":
                hits: List[Tuple[Any, Any, str]] = []
                if sh0 and cell0:
                    try:
                        ws0 = wb[sh0]
                        ri, ci = _xlsx_cell_row_col(cell0)
                        c0 = ws0.cell(row=ri, column=ci)
                        bv = "" if c0.value is None else str(c0.value)
                        hits = [(ws0, c0, bv)]
                    except Exception as e:
                        report["errors"].append({"op": op, "error": f"sheet/cell 定位失败：{e}"})
                        continue
                else:
                    hits = _find_cells_containing(anchor)
                    hits = _xlsx_pick_value_cell_hits(hits, anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中单元格"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"单元格命中不唯一：{len(hits)}"})
                    continue
                op_before = (op.get("before") or "").strip()
                for ws, cell, before in hits:
                    before_s = str(before or "")
                    if op_before and op_before in before_s:
                        after_s = before_s.replace(op_before, new_text, 1)
                        if after_s == before_s:
                            report["skipped"].append(
                                {"op": op, "reason": "no-op：before 子串替换未改变单元格"}
                            )
                            continue
                        cell.value = after_s
                        rep_after = after_s
                    else:
                        if before_s.strip() == (new_text or "").strip():
                            report["skipped"].append(
                                {"op": op, "reason": "no-op：新旧文本一致（单元格无需修改）"}
                            )
                            continue
                        cell.value = new_text
                        rep_after = new_text
                    report["changes"].append(
                        {
                            "type": t,
                            "anchor": anchor,
                            "sheet": ws.title,
                            "cell": cell.coordinate,
                            "before": before,
                            "after": rep_after,
                            **_change_meta(op),
                        }
                    )
                report["applied"].append({"op": op, "hits": len(hits)})
                continue

            if t == "insert_table_row_after_contains":
                hits = _find_rows_containing(anchor)
                if not hits:
                    report["skipped"].append({"op": op, "reason": "未命中表格行（用于插入）"})
                    continue
                if require_unique and len(hits) != 1:
                    report["skipped"].append({"op": op, "reason": f"插入行锚点命中不唯一：{len(hits)}"})
                    continue
                ws, r_idx = hits[0]
                vals = [x for x in (new_text or "").split("\t")]
                col_hint = max(_xlsx_last_nonblank_col(ws, r_idx), 1)
                if len(vals) == 1 and "\t" not in (new_text or ""):
                    vals2 = _xlsx_split_row_values_for_insert(new_text or "", col_hint)
                    if len(vals2) > 1:
                        vals = vals2
                if not vals or not any(str(x).strip() for x in vals):
                    report["skipped"].append(
                        {
                            "op": op,
                            "reason": "new_text 按制表符分列后无有效单元格内容（仅空白/制表符），未插入行",
                        }
                    )
                    continue
                try:
                    col_n = min(80, max(len(vals), col_hint, int(ws.max_column or 0), 1))
                    new_joined = "\t".join(
                        (vals[ci] if ci < len(vals) else "") for ci in range(col_n)
                    ).strip()
                    parsed = _parse_tc_id_first_cell(vals[0] or "", tc_rules)
                    if not parsed:
                        try:
                            av = ws.cell(row=r_idx, column=1).value
                            parsed = _parse_tc_id_first_cell("" if av is None else str(av), tc_rules)
                        except Exception:
                            parsed = None

                    if not parsed:
                        rid = _xlsx_first_cell_risk_id_token(vals[0] or "")
                        if rid:
                            ex_r2 = _xlsx_find_row_by_id_any_col(ws, rid)
                            if ex_r2 >= 0:
                                before_joined2 = _xlsx_row_joined(ws, ex_r2, col_n)
                                for ci in range(1, col_n + 1):
                                    v = vals[ci - 1] if ci - 1 < len(vals) else ""
                                    ws.cell(row=ex_r2, column=ci, value=v)
                                joined2 = "\t".join(
                                    (vals[ci] if ci < len(vals) else "") for ci in range(len(vals))
                                )
                                report["changes"].append(
                                    {
                                        "type": "update_table_row_in_place",
                                        "anchor": anchor,
                                        "sheet": ws.title,
                                        "row_index": ex_r2,
                                        "tc_id": rid,
                                        "before": before_joined2,
                                        "after": joined2,
                                        "cells_preview": [str(x).strip() for x in vals],
                                        "note": "首列编号/风险号已存在，已原地更新整行（避免重复插入）",
                                        **_change_meta(op),
                                    }
                                )
                                report["applied"].append({"op": op, "hits": 1})
                                continue

                    if parsed:
                        prefix, _n, rule = parsed
                        max_n = _table_max_tc_suffix_for_prefix_xlsx(ws, prefix, tc_rules)
                        next_id = _next_unique_tc_id_xlsx(
                            ws, prefix, rule, tc_rules, max_n + 1
                        )
                        id_cell = (vals[0] or "").strip()
                        ex_r = _xlsx_find_row_by_id_any_col(ws, id_cell)
                        if ex_r >= 0:
                            before_joined = _xlsx_row_joined(ws, ex_r, col_n)
                            sim = SequenceMatcher(
                                None, new_joined, (before_joined or "").strip()
                            ).ratio()
                            if sim >= _TC_ROW_SIMILARITY_UPDATE_THRESHOLD:
                                for ci in range(1, col_n + 1):
                                    v = vals[ci - 1] if ci - 1 < len(vals) else ""
                                    ws.cell(row=ex_r, column=ci, value=v)
                                joined = "\t".join(
                                    (vals[ci] if ci < len(vals) else "") for ci in range(len(vals))
                                )
                                report["changes"].append(
                                    {
                                        "type": "update_table_row_in_place",
                                        "anchor": anchor,
                                        "sheet": ws.title,
                                        "row_index": ex_r,
                                        "tc_id": id_cell,
                                        "similarity": round(sim, 4),
                                        "before": before_joined,
                                        "after": joined,
                                        "cells_preview": [str(x).strip() for x in vals],
                                        "note": (
                                            f"首列编号已存在且整行相似度≥{_TC_ROW_SIMILARITY_UPDATE_THRESHOLD}，已原地更新该行"
                                        ),
                                        **_change_meta(op),
                                    }
                                )
                                report["applied"].append({"op": op, "hits": 1})
                                continue
                            vals[0] = next_id
                        else:
                            vals[0] = next_id

                    ws.insert_rows(r_idx + 1)
                    for ci, v in enumerate(vals, start=1):
                        ws.cell(row=r_idx + 1, column=ci, value=v)
                    joined = "\t".join(vals)
                    chx: Dict[str, Any] = {
                        "type": t,
                        "anchor": anchor,
                        "sheet": ws.title,
                        "row_index": r_idx,
                        "before": "",
                        "after": joined,
                        "cells_preview": [str(x).strip() for x in vals],
                    }
                    if parsed:
                        chx["tc_id_assigned"] = (vals[0] or "").strip()
                        chx["note"] = "首列已按表内同模块（前缀）最大编号 +1 分配"
                    chx.update(_change_meta(op))
                    report["changes"].append(chx)
                    report["applied"].append({"op": op, "hits": 1})
                except Exception as _ie:
                    report["errors"].append({"op": op, "error": f"插入行失败：{_ie}"})
                continue

            report["skipped"].append({"op": op, "reason": f"不支持的 type：{t}"})
        except Exception as e:
            report["errors"].append({"op": op, "error": str(e)})

    # 修订记录：仅当本次确实写入了修改时，向已有“修订记录”工作表追加一行；不新建工作表。
    try:
        if (len(report.get("applied") or []) > 0) or (len(report.get("changes") or []) > 0):
            if not _xlsx_append_revision_row(wb, meta):
                report["skipped"].append({"reason": "未找到基础文件中的“修订记录”工作表/表格区域，已跳过写入修订记录。"})
    except Exception as e:
        report["errors"].append({"error": f"写入修订记录失败：{e}"})

    wb.save(str(out))
    try:
        wb.close()
    except Exception:
        pass
    return str(out), report

