"""审核报告导出：HTML、PDF、Word、Excel；支持在原 Word 文档对应位置插入批注；待办任务 CSV"""

import csv
import io
import html
from pathlib import Path
from typing import List, Dict, Any, Optional, Callable

_SEVERITY_MAP = {"high": "高", "medium": "中", "low": "低", "info": "提示"}


def _escape(s: str) -> str:
    if s is None:
        return ""
    return html.escape(str(s), quote=True)


def _sev_label(s: str) -> str:
    return _SEVERITY_MAP.get((s or "info").lower(), s or "")


def _meta_info_lines(meta: dict) -> List[str]:
    """从 _review_meta 构建人类可读的审核信息行。"""
    if not meta:
        return []
    lines = []
    lines.append(f"审核类型：{meta.get('audit_type', '通用审核')}")
    if meta.get("project_name"):
        lines.append(f"项目名称：{meta['project_name']}")
    if meta.get("product_name"):
        lines.append(f"产品名称：{meta['product_name']}")
    if meta.get("registration_country"):
        lines.append(f"注册国家：{meta['registration_country']}")
    if meta.get("registration_type"):
        lines.append(f"注册类别：{meta['registration_type']}")
    if meta.get("registration_component"):
        lines.append(f"注册组成：{meta['registration_component']}")
    if meta.get("project_form"):
        lines.append(f"项目形态：{meta['project_form']}")
    if meta.get("document_language"):
        lines.append(f"文档语言：{meta['document_language']}")
    return lines


def _extract_meta(reports: List[Dict[str, Any]]) -> dict:
    for r in reports:
        m = r.get("_review_meta")
        if m:
            return m
    return {}


def _modify_docs_str(point: dict) -> str:
    docs = point.get("modify_docs") or []
    if not isinstance(docs, list):
        return str(docs)
    return "、".join(d for d in docs if d)


# ──────────────────────────── HTML ────────────────────────────

def report_to_html(reports: List[Dict[str, Any]]) -> str:
    meta = _extract_meta(reports)
    meta_lines = _meta_info_lines(meta)
    parts = [
        """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>审核报告</title>
<style>
body { font-family: "Microsoft YaHei", sans-serif; margin: 24px; }
h1 { color: #333; }
h2 { color: #555; margin-top: 24px; border-bottom: 1px solid #ddd; }
table { border-collapse: collapse; margin: 12px 0; width: 100%; }
th, td { border: 1px solid #ddd; padding: 8px 12px; text-align: left; vertical-align: top; }
th { background: #f5f5f5; }
.meta { background: #eef6ff; padding: 12px; border-radius: 4px; margin-bottom: 16px; }
.meta p { margin: 4px 0; }
.point { margin: 16px 0; padding: 12px; background: #fafafa; border-radius: 4px; border-left: 4px solid #ccc; }
.point.high { border-left-color: #e53935; }
.point.medium { border-left-color: #fb8c00; }
.point.low { border-left-color: #1e88e5; }
.summary { margin: 12px 0; line-height: 1.6; }
.tag { display: inline-block; padding: 2px 8px; border-radius: 3px; font-size: 0.85em; color: #fff; margin-right: 6px; }
.tag-high { background: #e53935; } .tag-medium { background: #fb8c00; }
.tag-low { background: #1e88e5; } .tag-info { background: #78909c; }
</style>
</head>
<body>
<h1>📋 审核报告</h1>
"""
    ]
    if meta_lines:
        parts.append('<div class="meta">')
        for line in meta_lines:
            parts.append(f"<p>{_escape(line)}</p>")
        parts.append("</div>")

    for report in reports:
        fn = _escape(report.get("original_filename", report.get("file_name", "未知")))
        parts.append(f"<h2>📄 {fn}</h2>")
        parts.append(
            "<table><tr><th>高风险</th><th>中风险</th><th>低风险</th><th>提示</th><th>总计</th></tr>"
            f"<tr><td>{report.get('high_count', 0)}</td><td>{report.get('medium_count', 0)}</td>"
            f"<td>{report.get('low_count', 0)}</td><td>{report.get('info_count', 0)}</td>"
            f"<td>{report.get('total_points', 0)}</td></tr></table>"
        )
        if report.get("summary"):
            parts.append(f"<p class='summary'><strong>总结：</strong> {_escape(report['summary'])}</p>")
        for i, point in enumerate(report.get("audit_points", []), 1):
            sev = (point.get("severity") or "info").lower()
            sev_cls = sev if sev in ("high", "medium", "low") else ""
            action = point.get("action") or ""
            docs_str = _modify_docs_str(point)
            parts.append(
                f'<div class="point {sev_cls}">'
                f'<h3><span class="tag tag-{sev}">{_sev_label(sev)}</span>'
                f"审核点 {i}：{_escape(point.get('category', '未分类'))}</h3>"
                f"<p><strong>位置：</strong> {_escape(point.get('location'))}</p>"
                f"<p><strong>问题描述：</strong> {_escape(point.get('description'))}</p>"
                f"<p><strong>法规依据：</strong> {_escape(point.get('regulation_ref'))}</p>"
                f"<p><strong>修改建议：</strong> {_escape(point.get('suggestion'))}</p>"
            )
            if docs_str:
                parts.append(f"<p><strong>需修改文档：</strong> {_escape(docs_str)}</p>")
            if action:
                parts.append(f"<p><strong>处理状态：</strong> {_escape(action)}</p>")
            parts.append("</div>")
    parts.append("</body>\n</html>")
    return "\n".join(parts)


# ──────────────────────────── Word ────────────────────────────

def report_to_docx(reports: List[Dict[str, Any]]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    doc.add_heading("审核报告", 0)

    meta = _extract_meta(reports)
    meta_lines = _meta_info_lines(meta)
    if meta_lines:
        for line in meta_lines:
            doc.add_paragraph(line)
        doc.add_paragraph("")

    for report in reports:
        fn = report.get("original_filename", report.get("file_name", "未知"))
        doc.add_heading(fn, level=1)
        tbl = doc.add_table(rows=2, cols=5)
        tbl.style = "Light Grid Accent 1"
        for ci, hdr in enumerate(["高风险", "中风险", "低风险", "提示", "总计"]):
            tbl.rows[0].cells[ci].text = hdr
        for ci, key in enumerate(["high_count", "medium_count", "low_count", "info_count", "total_points"]):
            tbl.rows[1].cells[ci].text = str(report.get(key, 0))
        doc.add_paragraph("")
        if report.get("summary"):
            p = doc.add_paragraph()
            p.add_run("总结：").bold = True
            p.add_run(report["summary"])
        for i, point in enumerate(report.get("audit_points", []), 1):
            sev = (point.get("severity") or "info").lower()
            doc.add_heading(f"[{_sev_label(sev)}] 审核点 {i}：{point.get('category', '未分类')}", level=2)
            doc.add_paragraph(f"位置：{point.get('location', '')}")
            p_desc = doc.add_paragraph()
            p_desc.add_run("问题描述：").bold = True
            p_desc.add_run(point.get("description", ""))
            p_reg = doc.add_paragraph()
            p_reg.add_run("法规依据：").bold = True
            p_reg.add_run(point.get("regulation_ref", ""))
            p_sug = doc.add_paragraph()
            p_sug.add_run("修改建议：").bold = True
            p_sug.add_run(point.get("suggestion", ""))
            docs_str = _modify_docs_str(point)
            if docs_str:
                doc.add_paragraph(f"需修改文档：{docs_str}")
            action = point.get("action") or ""
            if action:
                doc.add_paragraph(f"处理状态：{action}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────── PDF ────────────────────────────

def report_to_pdf(reports: List[Dict[str, Any]]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    cjk_font = "Helvetica"
    try:
        for font_path in [
            "C:/Windows/Fonts/msyh.ttf",
            "C:/Windows/Fonts/simhei.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        ]:
            try:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
                cjk_font = "CJK"
                styles["Normal"].fontName = "CJK"
                styles["Heading1"].fontName = "CJK"
                styles["Heading2"].fontName = "CJK"
                break
            except Exception:
                continue
    except Exception:
        pass

    body_style = ParagraphStyle(name="Body", fontName=cjk_font, fontSize=10, leading=14, spaceAfter=4)
    label_style = ParagraphStyle(name="Label", fontName=cjk_font, fontSize=10, leading=14, spaceAfter=4, textColor=HexColor("#333333"))

    story = []
    title_style = ParagraphStyle(name="Title", fontName=cjk_font, fontSize=16, spaceAfter=12)
    story.append(Paragraph("审核报告", title_style))
    story.append(Spacer(1, 8))

    meta = _extract_meta(reports)
    meta_lines = _meta_info_lines(meta)
    if meta_lines:
        for line in meta_lines:
            story.append(Paragraph(html.escape(line), body_style))
        story.append(Spacer(1, 10))

    for report in reports:
        fn = report.get("original_filename", report.get("file_name", "未知"))
        story.append(Paragraph(f"<b>文件：</b> {html.escape(fn)}", styles["Heading2"]))
        story.append(Spacer(1, 6))
        data = [
            ["高风险", "中风险", "低风险", "提示", "总计"],
            [str(report.get("high_count", 0)), str(report.get("medium_count", 0)),
             str(report.get("low_count", 0)), str(report.get("info_count", 0)),
             str(report.get("total_points", 0))],
        ]
        t = Table(data, colWidths=[3*cm]*5)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#f0f0f0")),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#999999")),
            ("FONTNAME", (0, 0), (-1, -1), cjk_font),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(t)
        story.append(Spacer(1, 6))
        if report.get("summary"):
            story.append(Paragraph(f"<b>总结：</b> {html.escape(report['summary'])}", body_style))
            story.append(Spacer(1, 6))
        for i, point in enumerate(report.get("audit_points", []), 1):
            sev = (point.get("severity") or "info").lower()
            story.append(Paragraph(
                f"<b>[{_sev_label(sev)}] 审核点 {i}：</b> {html.escape(point.get('category', '未分类'))}",
                styles["Heading2"] if cjk_font == "Helvetica" else ParagraphStyle("PH2", fontName=cjk_font, fontSize=12, leading=16, spaceBefore=8, spaceAfter=4),
            ))
            story.append(Paragraph(f"<b>位置：</b>{html.escape(point.get('location', ''))}", body_style))
            story.append(Paragraph(f"<b>问题描述：</b>{html.escape(point.get('description', ''))}", body_style))
            story.append(Paragraph(f"<b>法规依据：</b>{html.escape(point.get('regulation_ref', ''))}", body_style))
            story.append(Paragraph(f"<b>修改建议：</b>{html.escape(point.get('suggestion', ''))}", body_style))
            docs_str = _modify_docs_str(point)
            if docs_str:
                story.append(Paragraph(f"<b>需修改文档：</b>{html.escape(docs_str)}", body_style))
            action = point.get("action") or ""
            if action:
                story.append(Paragraph(f"<b>处理状态：</b>{html.escape(action)}", body_style))
            story.append(Spacer(1, 8))
        story.append(Spacer(1, 12))
    doc.build(story)
    buf.seek(0)
    return buf.read()


# ──────────────────────────── Excel ────────────────────────────

def report_to_excel(reports: List[Dict[str, Any]]) -> bytes:
    """导出审核报告为 Excel (.xlsx)，含汇总 sheet + 每文档一个 sheet。"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    meta = _extract_meta(reports)
    header_font = Font(name="Microsoft YaHei", bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(name="Microsoft YaHei", bold=True, size=11, color="FFFFFF")
    cell_font = Font(name="Microsoft YaHei", size=10)
    wrap_align = Alignment(wrap_text=True, vertical="top")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    sev_fills = {
        "high": PatternFill(start_color="FCE4EC", end_color="FCE4EC", fill_type="solid"),
        "medium": PatternFill(start_color="FFF3E0", end_color="FFF3E0", fill_type="solid"),
        "low": PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid"),
    }

    # ── 汇总 sheet ──
    ws_summary = wb.active
    ws_summary.title = "汇总"
    meta_lines = _meta_info_lines(meta)
    row_idx = 1
    if meta_lines:
        for line in meta_lines:
            ws_summary.cell(row=row_idx, column=1, value=line).font = Font(name="Microsoft YaHei", size=10, bold=True)
            row_idx += 1
        row_idx += 1

    summary_headers = ["序号", "文件名", "高风险", "中风险", "低风险", "提示", "总计"]
    for ci, h in enumerate(summary_headers, 1):
        cell = ws_summary.cell(row=row_idx, column=ci, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border
    row_idx += 1
    for ri, report in enumerate(reports, 1):
        fn = report.get("original_filename", report.get("file_name", "未知"))
        vals = [ri, fn, report.get("high_count", 0), report.get("medium_count", 0),
                report.get("low_count", 0), report.get("info_count", 0), report.get("total_points", 0)]
        for ci, v in enumerate(vals, 1):
            cell = ws_summary.cell(row=row_idx, column=ci, value=v)
            cell.font = cell_font
            cell.border = thin_border
            if ci >= 3:
                cell.alignment = Alignment(horizontal="center")
        row_idx += 1

    ws_summary.column_dimensions["A"].width = 6
    ws_summary.column_dimensions["B"].width = 40
    for col_letter in ["C", "D", "E", "F", "G"]:
        ws_summary.column_dimensions[col_letter].width = 10

    # ── 每文档 sheet ──
    detail_headers = ["序号", "严重程度", "类别", "位置", "问题描述", "法规依据", "修改建议", "需修改文档", "处理状态"]
    detail_widths = [6, 10, 12, 18, 40, 30, 40, 25, 12]

    for r_idx, report in enumerate(reports):
        fn = report.get("original_filename", report.get("file_name", f"文档{r_idx+1}"))
        safe_sheet = "".join(c for c in fn if c.isalnum() or c in " _-().").strip()[:28] or f"Doc{r_idx+1}"
        existing = [ws.title for ws in wb.worksheets]
        if safe_sheet in existing:
            safe_sheet = f"{safe_sheet}_{r_idx+1}"
        ws = wb.create_sheet(title=safe_sheet)

        ws.cell(row=1, column=1, value=f"文件：{fn}").font = Font(name="Microsoft YaHei", size=12, bold=True)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(detail_headers))

        stats_row = 2
        stat_labels = ["高风险", "中风险", "低风险", "提示", "总计"]
        stat_keys = ["high_count", "medium_count", "low_count", "info_count", "total_points"]
        for ci, (lbl, key) in enumerate(zip(stat_labels, stat_keys)):
            ws.cell(row=stats_row, column=ci*2+1, value=lbl).font = Font(name="Microsoft YaHei", bold=True, size=10)
            ws.cell(row=stats_row, column=ci*2+2, value=report.get(key, 0)).font = cell_font

        if report.get("summary"):
            ws.cell(row=3, column=1, value=f"总结：{report['summary']}").font = cell_font
            ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=len(detail_headers))

        hdr_row = 5
        for ci, h in enumerate(detail_headers, 1):
            cell = ws.cell(row=hdr_row, column=ci, value=h)
            cell.font = header_font_white
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
            cell.border = thin_border

        for ci, w in enumerate(detail_widths, 1):
            ws.column_dimensions[chr(64 + ci) if ci <= 9 else ""].width = w
        col_letters = ["A", "B", "C", "D", "E", "F", "G", "H", "I"]
        for ci, w in enumerate(detail_widths):
            ws.column_dimensions[col_letters[ci]].width = w

        data_row = hdr_row + 1
        for pi, point in enumerate(report.get("audit_points", []), 1):
            sev = (point.get("severity") or "info").lower()
            vals = [
                pi,
                _sev_label(sev),
                point.get("category", ""),
                point.get("location", ""),
                point.get("description", ""),
                point.get("regulation_ref", ""),
                point.get("suggestion", ""),
                _modify_docs_str(point),
                point.get("action", ""),
            ]
            fill = sev_fills.get(sev)
            for ci, v in enumerate(vals, 1):
                cell = ws.cell(row=data_row, column=ci, value=v)
                cell.font = cell_font
                cell.alignment = wrap_align
                cell.border = thin_border
                if fill:
                    cell.fill = fill
            data_row += 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def _find_paragraph_for_comment(doc, location: str, description: str) -> Optional[Any]:
    """在 doc 的 body 段落中查找与 location 或 description 片段匹配的段落，用于挂批注。"""
    location = (location or "").strip()
    description = (description or "").strip()
    search_parts = []
    if location:
        search_parts.append(location[:50])
    if description:
        search_parts.append(description[:40])
    if not search_parts:
        return None
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        for part in search_parts:
            if part and part in text:
                if para.runs:
                    return para
                break
    for para in doc.paragraphs:
        if para.runs and (para.text or "").strip():
            return para
    return None


def report_to_docx_with_comments(original_docx_path_or_bytes, report: Dict[str, Any], author: str = "审核") -> bytes:
    """
    基于原 Word 文档生成带批注的副本：在每个审核点对应位置插入批注（位置/描述匹配段落）。
    original_docx_path_or_bytes: 待审 docx 文件路径(str) 或 docx 二进制内容(bytes)
    report: 单份审核报告 dict（含 audit_points）
    author: 批注作者显示名
    返回带批注的 docx 的 bytes。
    """
    from docx import Document

    if isinstance(original_docx_path_or_bytes, bytes):
        doc = Document(io.BytesIO(original_docx_path_or_bytes))
    else:
        path = Path(str(original_docx_path_or_bytes))
        if not path.exists() or path.suffix.lower() not in (".docx", ".doc"):
            raise ValueError("仅支持 .docx 原文件")
        doc = Document(str(path))
    used_paragraphs = set()
    for i, point in enumerate(report.get("audit_points", []), 1):
        location = point.get("location") or ""
        description = point.get("description") or ""
        severity = point.get("severity") or ""
        regulation_ref = point.get("regulation_ref") or ""
        suggestion = point.get("suggestion") or ""
        comment_text = f"[{severity}] {description}"
        if regulation_ref:
            comment_text += f"\n法规依据：{regulation_ref}"
        if suggestion:
            comment_text += f"\n建议：{suggestion}"
        para = _find_paragraph_for_comment(doc, location, description)
        if para is None or id(para) in used_paragraphs:
            for p in doc.paragraphs:
                if p.runs and id(p) not in used_paragraphs:
                    para = p
                    break
        if para is not None and para.runs:
            try:
                doc.add_comment(runs=para.runs, text=comment_text, author=author, initials="审")
                used_paragraphs.add(id(para))
            except Exception:
                pass
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# 待办导入模板表头（与制氧机 web 管理系统等导入格式一致）
TODO_CSV_HEADERS = [
    "项目名称", "项目编号", "影响业务方", "产品", "国家", "项目备注",
    "文件名称", "任务类型", "文档链接", "文件版本号", "编写人员", "负责人",
    "截止日期", "下发任务备注", "文档体现日期", "审核人员", "批准人员", "所属模块",
]


def _default_action_for_severity(severity: str) -> str:
    """与 app 中多文档默认状态一致，用于导出时筛选「立即修改」项。"""
    severity = (severity or "info").lower()
    defaults = {"high": "立即修改", "medium": "立即修改", "low": "延期修改", "info": "无需修改"}
    return defaults.get(severity, "无需修改")


def _todo_rows_from_reports(
    reports: List[Dict[str, Any]],
    only_immediate: bool = True,
    get_default_action: Optional[Callable[[str], str]] = None,
    project_name: str = "",
    product: str = "",
    country: str = "",
    impact_business: str = "",
    module: str = "",
) -> List[Dict[str, Any]]:
    """
    从报告中提取待办行（仅「立即修改」或全部），返回 dict 列表，每项含：
    序号、项目名称、产品、国家、文件名称、任务类型、文档链接、类别、严重程度、位置、
    问题描述、法规依据、修改建议、需修改文档、下发任务备注。
    """
    default_action_fn = get_default_action or _default_action_for_severity

    def _meta_from_reports() -> dict:
        for r in reports:
            m = r.get("_review_meta")
            if m:
                return m
        return {}

    _meta = _meta_from_reports()
    _project_name = project_name or _meta.get("project_name", "")
    _product = product or _meta.get("product_name", "")
    _country = country or _meta.get("registration_country", "")

    rows = []
    n = 0
    for report in reports:
        file_name = report.get("original_filename", report.get("file_name", "未知"))
        doc_link = report.get("_kdocs_view_url", "") or ""
        r_meta = report.get("_review_meta") or {}
        r_project = _project_name or r_meta.get("project_name", "")
        r_product = _product or r_meta.get("product_name", "")
        r_country = _country or r_meta.get("registration_country", "")
        points = report.get("audit_points") or []
        for p in points:
            action = p.get("action") or default_action_fn(p.get("severity", "info"))
            if only_immediate and action != "立即修改":
                continue
            n += 1
            location = (p.get("location") or "").strip()
            desc = (p.get("description") or "").strip()
            sug = (p.get("suggestion") or "").strip()
            reg_ref = (p.get("regulation_ref") or "").strip()
            remark_parts = []
            if location:
                remark_parts.append(f"位置：{location}")
            if desc:
                remark_parts.append(f"问题描述：{desc}")
            if reg_ref:
                remark_parts.append(f"法规依据：{reg_ref}")
            if sug:
                remark_parts.append(f"修改建议：{sug}")
            remark = "；".join(remark_parts) if remark_parts else "审核待修改"
            docs = p.get("modify_docs") or []
            modify_docs_str = "、".join(d for d in docs if d) if isinstance(docs, list) else str(docs)

            rows.append({
                "序号": n,
                "项目名称": r_project,
                "产品": r_product,
                "国家": r_country,
                "文件名称": file_name,
                "任务类型": "审核待修改",
                "文档链接": doc_link,
                "类别": p.get("category", ""),
                "严重程度": _sev_label(p.get("severity", "info")),
                "位置": location,
                "问题描述": desc,
                "法规依据": reg_ref,
                "修改建议": sug,
                "需修改文档": modify_docs_str,
                "下发任务备注": remark,
                "影响业务方": impact_business,
                "所属模块": module,
            })
    return rows


def report_todo_to_pdf(
    reports: List[Dict[str, Any]],
    only_immediate: bool = True,
    get_default_action: Optional[Callable[[str], str]] = None,
    project_name: str = "",
    product: str = "",
    country: str = "",
    impact_business: str = "",
    module: str = "",
) -> bytes:
    """导出待办任务为 PDF，采用与审核报告相同的逐条段落布局，完整展示全部字段。"""
    rows = _todo_rows_from_reports(
        reports, only_immediate=only_immediate, get_default_action=get_default_action,
        project_name=project_name, product=product, country=country,
        impact_business=impact_business, module=module,
    )

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    cjk_font = "Helvetica"
    try:
        for font_path in [
            "C:/Windows/Fonts/msyh.ttf",
            "C:/Windows/Fonts/simhei.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        ]:
            try:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
                cjk_font = "CJK"
                styles["Normal"].fontName = "CJK"
                styles["Heading1"].fontName = "CJK"
                styles["Heading2"].fontName = "CJK"
                break
            except Exception:
                continue
    except Exception:
        pass

    body_style = ParagraphStyle(name="Body", fontName=cjk_font, fontSize=10, leading=14, spaceAfter=4)
    heading_style = ParagraphStyle(name="ItemHeading", fontName=cjk_font, fontSize=12, leading=16, spaceBefore=10, spaceAfter=4)

    story = []
    title_style = ParagraphStyle(name="Title", fontName=cjk_font, fontSize=16, spaceAfter=12)
    story.append(Paragraph("待办任务（审核待修改）", title_style))
    story.append(Spacer(1, 6))

    meta = _extract_meta(reports)
    meta_lines = _meta_info_lines(meta)
    if meta_lines:
        for line in meta_lines:
            story.append(Paragraph(html.escape(line), body_style))
        story.append(Spacer(1, 8))

    story.append(Paragraph(f"共 <b>{len(rows)}</b> 项待办", body_style))
    story.append(Spacer(1, 8))

    if not rows:
        story.append(Paragraph("暂无待办任务。", body_style))
    else:
        sev_colors = {"高": "#e53935", "中": "#fb8c00", "低": "#1e88e5", "提示": "#78909c"}
        for r in rows:
            sev = str(r.get("严重程度", ""))
            color = sev_colors.get(sev, "#333333")
            story.append(Paragraph(
                f'<b><font color="{color}">[{html.escape(sev)}]</font> '
                f'待办 {r.get("序号", "")}：{html.escape(str(r.get("类别", "")))}</b>',
                heading_style,
            ))
            if r.get("文件名称"):
                story.append(Paragraph(f"<b>文件名称：</b>{html.escape(str(r['文件名称']))}", body_style))
            if r.get("位置"):
                story.append(Paragraph(f"<b>位置：</b>{html.escape(str(r['位置']))}", body_style))
            if r.get("问题描述"):
                story.append(Paragraph(f"<b>问题描述：</b>{html.escape(str(r['问题描述']))}", body_style))
            if r.get("法规依据"):
                story.append(Paragraph(f"<b>法规依据：</b>{html.escape(str(r['法规依据']))}", body_style))
            if r.get("修改建议"):
                story.append(Paragraph(f"<b>修改建议：</b>{html.escape(str(r['修改建议']))}", body_style))
            if r.get("需修改文档"):
                story.append(Paragraph(f"<b>需修改文档：</b>{html.escape(str(r['需修改文档']))}", body_style))
            info_parts = []
            if r.get("项目名称"):
                info_parts.append(f"项目：{html.escape(str(r['项目名称']))}")
            if r.get("产品"):
                info_parts.append(f"产品：{html.escape(str(r['产品']))}")
            if r.get("国家"):
                info_parts.append(f"国家：{html.escape(str(r['国家']))}")
            if info_parts:
                story.append(Paragraph("　|　".join(info_parts), body_style))
            story.append(Spacer(1, 8))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def report_todo_to_docx(
    reports: List[Dict[str, Any]],
    only_immediate: bool = True,
    get_default_action: Optional[Callable[[str], str]] = None,
    project_name: str = "",
    product: str = "",
    country: str = "",
    impact_business: str = "",
    module: str = "",
) -> bytes:
    """导出待办任务为 Word (.docx)，采用与审核报告相同的逐条段落布局，完整展示全部字段。"""
    rows = _todo_rows_from_reports(
        reports, only_immediate=only_immediate, get_default_action=get_default_action,
        project_name=project_name, product=product, country=country,
        impact_business=impact_business, module=module,
    )
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    doc.add_heading("待办任务（审核待修改）", 0)

    meta = _extract_meta(reports)
    meta_lines = _meta_info_lines(meta)
    if meta_lines:
        for line in meta_lines:
            doc.add_paragraph(line)
        doc.add_paragraph("")

    doc.add_paragraph(f"共 {len(rows)} 项待办")

    if not rows:
        doc.add_paragraph("暂无待办任务。")
    else:
        for r in rows:
            sev = str(r.get("严重程度", ""))
            doc.add_heading(f"[{sev}] 待办 {r.get('序号', '')}：{r.get('类别', '')}", level=2)
            if r.get("文件名称"):
                p = doc.add_paragraph()
                p.add_run("文件名称：").bold = True
                p.add_run(str(r["文件名称"]))
            if r.get("位置"):
                p = doc.add_paragraph()
                p.add_run("位置：").bold = True
                p.add_run(str(r["位置"]))
            if r.get("问题描述"):
                p = doc.add_paragraph()
                p.add_run("问题描述：").bold = True
                p.add_run(str(r["问题描述"]))
            if r.get("法规依据"):
                p = doc.add_paragraph()
                p.add_run("法规依据：").bold = True
                p.add_run(str(r["法规依据"]))
            if r.get("修改建议"):
                p = doc.add_paragraph()
                p.add_run("修改建议：").bold = True
                p.add_run(str(r["修改建议"]))
            if r.get("需修改文档"):
                p = doc.add_paragraph()
                p.add_run("需修改文档：").bold = True
                p.add_run(str(r["需修改文档"]))
            info_parts = []
            if r.get("项目名称"):
                info_parts.append(f"项目：{r['项目名称']}")
            if r.get("产品"):
                info_parts.append(f"产品：{r['产品']}")
            if r.get("国家"):
                info_parts.append(f"国家：{r['国家']}")
            if info_parts:
                doc.add_paragraph("　|　".join(info_parts))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def report_todo_to_excel(
    reports: List[Dict[str, Any]],
    only_immediate: bool = True,
    get_default_action: Optional[Callable[[str], str]] = None,
    project_name: str = "",
    product: str = "",
    country: str = "",
    impact_business: str = "",
    module: str = "",
) -> bytes:
    """导出待办任务为 Excel (.xlsx)。"""
    rows = _todo_rows_from_reports(
        reports, only_immediate=only_immediate, get_default_action=get_default_action,
        project_name=project_name, product=product, country=country,
        impact_business=impact_business, module=module,
    )
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "待办任务"
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(name="Microsoft YaHei", bold=True, size=10, color="FFFFFF")
    cell_font = Font(name="Microsoft YaHei", size=10)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    wrap_align = Alignment(wrap_text=True, vertical="top")

    headers = ["序号", "项目名称", "产品", "国家", "文件名称", "任务类型", "文档链接", "类别", "严重程度", "位置", "问题描述", "法规依据", "修改建议", "需修改文档", "下发任务备注"]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = thin_border
    widths = [6, 14, 12, 10, 28, 12, 20, 14, 8, 18, 35, 25, 35, 20, 40]
    for ci, w in enumerate(widths, 1):
        col_letter = chr(64 + ci) if ci <= 26 else ("A" + chr(64 + ci - 26))
        try:
            ws.column_dimensions[col_letter].width = min(w, 50)
        except Exception:
            pass
    for ri, r in enumerate(rows, 2):
        vals = [
            r.get("序号", ""),
            r.get("项目名称", ""),
            r.get("产品", ""),
            r.get("国家", ""),
            r.get("文件名称", ""),
            r.get("任务类型", ""),
            r.get("文档链接", ""),
            r.get("类别", ""),
            r.get("严重程度", ""),
            r.get("位置", ""),
            r.get("问题描述", ""),
            r.get("法规依据", ""),
            r.get("修改建议", ""),
            r.get("需修改文档", ""),
            r.get("下发任务备注", ""),
        ]
        for ci, v in enumerate(vals, 1):
            cell = ws.cell(row=ri, column=ci, value=v)
            cell.font = cell_font
            cell.alignment = wrap_align
            cell.border = thin_border
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def report_todo_to_csv(
    reports: List[Dict[str, Any]],
    only_immediate: bool = True,
    get_default_action: Optional[Callable[[str], str]] = None,
    project_name: str = "",
    product: str = "",
    country: str = "",
    impact_business: str = "",
    module: str = "",
) -> bytes:
    """
    将审核报告中的待办项（仅「立即修改」或全部审核点）导出为与待办导入模板同格式的 CSV。
    每行一条待办：文件名称、任务类型=审核待修改、文档链接、下发任务备注=位置+描述+建议。
    only_immediate: True 时只导出处理状态为「立即修改」的审核点；False 时导出全部审核点。
    get_default_action: 用于判断默认处理状态的函数 severity -> "立即修改"|"延期修改"|...；不传则用内置默认。
    project_name, product, country 等：可选默认列值，可留空。
    返回 UTF-8-BOM 编码的 CSV 字节，便于 Excel 正确打开中文。
    """
    rows = _todo_rows_from_reports(
        reports, only_immediate=only_immediate, get_default_action=get_default_action,
        project_name=project_name, product=product, country=country,
        impact_business=impact_business, module=module,
    )
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(TODO_CSV_HEADERS)
    for r in rows:
        row = [
            r.get("项目名称", ""),
            "",
            r.get("影响业务方", ""),
            r.get("产品", ""),
            r.get("国家", ""),
            "",
            r.get("文件名称", ""),
            r.get("任务类型", "审核待修改"),
            r.get("文档链接", ""),
            "",
            "",
            "",
            "",
            r.get("下发任务备注", ""),
            "",
            "",
            "",
            r.get("所属模块", ""),
        ]
        writer.writerow(row)
    return buf.getvalue().encode("utf-8-sig")