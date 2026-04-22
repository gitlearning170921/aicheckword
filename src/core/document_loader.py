"""文档加载器：支持 PDF、Word、Excel、TXT、Markdown 等格式，以及压缩包自动解压"""

import os
import re
import shutil
import subprocess
import zipfile
import tempfile
import tarfile
import base64
from pathlib import Path
from typing import List, Optional, Tuple

import httpx
from .langchain_compat import Document, RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredExcelLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
)

from config import settings
from config.settings import get_pdf_ocr_llm_model
from config.cursor_overrides import get_llm_verify_ssl, get_llm_trust_env

LOADER_MAP = {
    ".pdf": PyPDFLoader,
    ".docx": Docx2txtLoader,
    ".doc": UnstructuredWordDocumentLoader,  # 旧版 .doc 用 Unstructured（需系统安装 LibreOffice 时效果最佳）
    ".xlsx": UnstructuredExcelLoader,
    ".xls": UnstructuredExcelLoader,
    ".txt": TextLoader,
    ".md": UnstructuredMarkdownLoader,
}

# 支持的压缩格式（.rar 需安装 rarfile 及系统 UnRAR 可执行文件）
ARCHIVE_EXTENSIONS = (".zip", ".tar", ".tgz", ".gz", ".rar")
SUPPORTED_DOC_EXTENSIONS = tuple(LOADER_MAP.keys())

# 标了「废弃」的文件或文件夹在训练、审核时跳过（路径或文件名中含此标记即视为废弃）
DEPRECATED_MARKER = "废弃"


def is_deprecated_path(path_or_name) -> bool:
    """路径或文件名中若包含「废弃」则视为废弃，跳过训练/审核。"""
    s = str(path_or_name) if path_or_name is not None else ""
    return DEPRECATED_MARKER in s


def extract_section_outline_from_texts(texts: List[str], max_sections: int = 80) -> str:
    """从多段文档内容中提取章节标题，形成「应有章节」参考列表，用于文档内容完整性审核。
    识别常见格式：第X章、1. 1.1、一、二、# 标题 等。去重并按出现顺序保留。"""
    import re
    seen = set()
    sections = []
    # 常见章节模式（一行视为一个候选标题）
    patterns = [
        re.compile(r"^第[一二三四五六七八九十百零\d]+章\s*.+$", re.MULTILINE),
        re.compile(r"^\d+(\.\d+)*[\.\s、]\s*.+$", re.MULTILINE),  # 1. 1.1 1.1.1
        re.compile(r"^[一二三四五六七八九十]+[、．.]\s*.+$", re.MULTILINE),
        re.compile(r"^#+\s*.+$", re.MULTILINE),  # Markdown
        re.compile(r"^[（(]\s*[一二三四五六七八九十\d]+\s*[)）]\s*.+$", re.MULTILINE),
    ]
    for text in (texts or []):
        if not (text and isinstance(text, str)):
            continue
        for line in text.splitlines():
            line = line.strip()
            if len(line) < 2 or len(line) > 200:
                continue
            for pat in patterns:
                if pat.match(line):
                    key = re.sub(r"\s+", " ", line).strip()
                    if key and key not in seen:
                        seen.add(key)
                        sections.append(line if len(line) <= 100 else line[:97] + "...")
                    break
            if len(sections) >= max_sections:
                break
        if len(sections) >= max_sections:
            break
    if not sections:
        return ""
    return "\n".join(sections)


# PDF：仅支持文本型；含图片/扫描版无法提取文字。限制最大页数避免大文件卡死
MAX_PDF_PAGES = 500
MIN_PDF_TEXT_LEN = 20  # 提取文字少于此长度视为“图片版/扫描版”
MAX_PDF_OCR_PAGES = 30  # AI OCR 兜底最多识别前 N 页，避免超时与费用失控

# DashScope 多模态 OCR 默认模型（与纯文本审核模型 qwen-plus 等不同）
_DASHSCOPE_OCR_VL_DEFAULT = "qwen-vl-plus"


def _dashscope_ocr_vl_model(ocr_field: str, chat_model: str) -> str:
    """从侧栏「OCR 模型」或审核模型名中选出可用的通义 VL 模型名；非 VL 时回退默认。"""
    for candidate in ((ocr_field or "").strip(), (chat_model or "").strip()):
        if not candidate:
            continue
        cl = candidate.lower()
        if "vl" in cl or "qwen-vl" in cl or "qwen2.5-vl" in cl or "qwen2-vl" in cl:
            return candidate
    return _DASHSCOPE_OCR_VL_DEFAULT


def _parse_dashscope_multimodal_text(resp) -> Tuple[str, str]:
    """解析 MultiModalConversation.call 返回值，成功 (text, '')，失败 ('', err)。"""
    try:
        sc = int(getattr(resp, "status_code", 0) or 0)
        if sc != 200:
            return "", f"HTTP {sc}: {getattr(resp, 'message', '') or resp}"
        code = (getattr(resp, "code", None) or "").strip()
        if code:
            return "", f"{code}: {getattr(resp, 'message', '') or '请求失败'}"
        out = getattr(resp, "output", None)
        if not out:
            return "", "DashScope 返回无 output"
        tx = getattr(out, "text", None)
        if isinstance(tx, str) and tx.strip():
            return tx.strip(), ""
        choices = getattr(out, "choices", None) or []
        if choices:
            msg = getattr(choices[0], "message", None)
            if msg is not None:
                c = getattr(msg, "content", None)
                if isinstance(c, str) and c.strip():
                    return c.strip(), ""
                if isinstance(c, list):
                    parts = []
                    for it in c:
                        if isinstance(it, dict):
                            t = (it.get("text") or "").strip()
                            if t:
                                parts.append(t)
                    if parts:
                        return "\n".join(parts).strip(), ""
        return "", "模型未返回可解析文本"
    except Exception as e:
        return "", str(e)


def _dashscope_ocr_one_page(img_b64: str, api_key: str, vl_model: str) -> Tuple[str, str]:
    """通义 DashScope 多模态单页 OCR（content 为 image + text，非 OpenAI image_url）。"""
    try:
        from dashscope import MultiModalConversation
    except Exception as e:
        return "", f"未安装 dashscope，无法调用通义多模态 OCR：{e!s}"
    try:
        resp = MultiModalConversation.call(
            model=vl_model,
            api_key=api_key,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"image": f"data:image/png;base64,{img_b64}"},
                        {
                            "text": (
                                "请做 OCR，仅输出图片中的正文文字。要求：1) 保留原有换行；"
                                "2) 不要解释；3) 看不清处写[不清晰]。"
                            ),
                        },
                    ],
                }
            ],
        )
    except Exception as e:
        return "", str(e)
    return _parse_dashscope_multimodal_text(resp)


def _extract_openai_choice_text(data: dict) -> str:
    """兼容 OpenAI 风格返回：message.content 可能是字符串或分段数组。"""
    try:
        choice = (data.get("choices") or [None])[0] or {}
        msg = (choice.get("message") or {})
        c = msg.get("content")
        if isinstance(c, str):
            return c.strip()
        if isinstance(c, list):
            parts = []
            for it in c:
                if isinstance(it, dict):
                    t = (it.get("text") or "").strip()
                    if t:
                        parts.append(t)
            return "\n".join(parts).strip()
    except Exception:
        return ""
    return ""


def _pdf_ocr_with_multimodal_llm(path: Path, max_pages: int = MAX_PDF_OCR_PAGES) -> Tuple[List[Document], str]:
    """
    当 PDF 文本层几乎为空时，尝试用多模态模型 OCR。
    返回 (docs, err_msg)：成功时 err_msg 为空；失败时 docs 为空并给出可读错误。
    """
    p = (settings.provider or "").strip().lower()
    # Cursor 主对话不走本函数；扫描 PDF 的 OCR 按「向量化使用」走 Ollama 多模态或 OpenAI Vision
    if p == "cursor":
        emb = (getattr(settings, "cursor_embedding", None) or "ollama").strip().lower()
        p = "openai" if emb == "openai" else "ollama"
    _vision = get_pdf_ocr_llm_model()
    # OpenAI 兼容多模态（含零一）；DeepSeek 公网 Chat 不支持 image_url，见下方分支
    model = _vision or (settings.llm_model or "").strip() or ("deepseek-chat" if p == "deepseek" else "gpt-4o-mini")

    try:
        import fitz  # PyMuPDF
    except Exception:
        return [], "未安装 PyMuPDF（fitz），无法将 PDF 渲染为图片做 AI OCR。请先安装：pip install pymupdf"

    try:
        pdf = fitz.open(str(path))
    except Exception as e:
        return [], f"PDF 打开失败，无法执行 AI OCR：{e!s}"

    docs: List[Document] = []
    timeout = httpx.Timeout(180.0, connect=45.0)
    err_detail = ""
    with httpx.Client(verify=get_llm_verify_ssl(), trust_env=get_llm_trust_env(), timeout=timeout) as client:
        try:
            total = min(len(pdf), max_pages)
            for i in range(total):
                page = pdf.load_page(i)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                img_b64 = base64.b64encode(pix.tobytes("png")).decode("ascii")
                text = ""
                if p == "tongyi":
                    ds_key = (settings.dashscope_api_key or "").strip()
                    if not ds_key:
                        err_detail = "provider=tongyi 未配置 DASHSCOPE_API_KEY，无法执行 AI OCR。"
                        break
                    vl_model = _dashscope_ocr_vl_model(_vision, settings.llm_model or "")
                    text, page_err = _dashscope_ocr_one_page(img_b64, ds_key, vl_model)
                    if page_err:
                        err_detail = page_err
                        break
                elif p == "deepseek":
                    # DeepSeek Chat Completions 不接受 OpenAI Vision 的 image_url 块，否则会 HTTP 400。
                    ds_key = (settings.dashscope_api_key or "").strip()
                    if not ds_key:
                        err_detail = (
                            "DeepSeek 接口不支持多模态 OCR（image_url）。"
                            "请在 .env 配置 DASHSCOPE_API_KEY，系统将用通义千问 VL 识别扫描页；"
                            "或改用 Ollama 多模态（如 qwen2.5vl），或侧栏切换为「通义」提供方。"
                        )
                        break
                    vl_model = _dashscope_ocr_vl_model(_vision, settings.llm_model or "")
                    text, page_err = _dashscope_ocr_one_page(img_b64, ds_key, vl_model)
                    if page_err:
                        err_detail = page_err
                        break
                elif p in ("openai", "lingyi"):
                    api_key = (
                        (settings.lingyi_api_key if p == "lingyi" else settings.openai_api_key)
                        or settings.openai_api_key
                        or ""
                    ).strip()
                    if not api_key:
                        err_detail = f"provider={p} 未配置 API Key，无法执行 AI OCR。"
                        break
                    base_url = (
                        settings.lingyi_base_url if p == "lingyi" else settings.openai_base_url
                    ) or "https://api.openai.com/v1"
                    base_url = base_url.rstrip("/")
                    payload = {
                        "model": model,
                        "temperature": 0,
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "请做 OCR，仅输出图片中的正文文字。要求：1) 保留原有换行；2) 不要解释；3) 看不清处写[不清晰]。",
                                    },
                                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}},
                                ],
                            }
                        ],
                    }
                    r = client.post(
                        f"{base_url}/chat/completions",
                        json=payload,
                        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    )
                    if r.status_code >= 400:
                        err_detail = f"HTTP {r.status_code}: {(r.text or '')[:300]}"
                        break
                    text = _extract_openai_choice_text(r.json())
                elif p == "ollama":
                    # Ollama 多模态：需使用支持图像输入的模型（如 qwen2.5vl/llava 等）
                    base_url = (settings.ollama_base_url or "http://localhost:11434").rstrip("/")
                    payload = {
                        "model": model,
                        "stream": False,
                        "messages": [
                            {
                                "role": "user",
                                "content": "请做 OCR，仅输出图片中的正文文字。要求：保留原有换行，不要解释，看不清处写[不清晰]。",
                                "images": [img_b64],
                            }
                        ],
                    }
                    r = client.post(f"{base_url}/api/chat", json=payload)
                    if r.status_code >= 400:
                        err_detail = f"HTTP {r.status_code}: {(r.text or '')[:300]}"
                        break
                    data = r.json() or {}
                    text = ((data.get("message") or {}).get("content") or "").strip()
                else:
                    err_detail = f"当前 provider={p or 'unknown'} 未实现图像 OCR 兜底。"
                    break
                if not text:
                    text = "[该页 OCR 未识别到可用文本]"
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source_file": path.name, "file_type": ".pdf", "page": i + 1, "ocr_fallback": True},
                    )
                )
        except Exception as e:
            err_detail = str(e)
        finally:
            pdf.close()

    if not docs:
        return [], f"AI OCR 失败：{err_detail or '模型未返回可用内容'}"
    return docs, ""


def _xlsx_cell_to_text(c) -> str:
    """单元格转字符串：尽量保留文本型编号（避免无意义的小数尾）；数值型无法恢复 Excel 自定义显示格式中的前导零。"""
    if c is None:
        return ""
    if isinstance(c, bool):
        return str(c)
    if isinstance(c, str):
        return c.strip()
    if isinstance(c, float):
        if c == int(c) and abs(c) < 1e15:
            return str(int(c))
        return repr(c) if c != c else ""  # NaN
    if isinstance(c, int):
        return str(c)
    return str(c).strip()


def _excel_risk_sheet_content_hint(sheet_title: str) -> str:
    """在风险类工作表摘录前加一句，引导模型按「风险需求」或英文 Risk demand 列取 CS 原文。"""
    st = (sheet_title or "").strip()
    snc = st.replace(" ", "").replace("\u3000", "")
    sl = snc.lower()
    zh = "风险" in snc and ("分析" in snc or "管理" in snc or "评估" in snc)
    en = "risk" in sl and any(
        k in sl for k in ("analysis", "analys", "management", "assessment", "mitigation", "control")
    )
    if zh or en:
        return (
            "【提示：CS 等措施编号常见列名为中文「风险需求」或英文「Risk demand (Measure) ID」/「Risk demand」等；"
            "请按表头定位对应列并摘取单元格原文，勿改写或与其他列混淆。】\n"
        )
    return ""


def _prioritize_excel_sheet_documents(docs: List[Document]) -> List[Document]:
    """
    将名称含「风险分析/风险管理」等（**含英文 Risk Analysis 等**）的工作表排在前面，
    避免跨文档追溯时总字数截断把后面的风险分析表切没。
    """
    if not docs or len(docs) < 2:
        return docs
    scored = []
    for i, d in enumerate(docs):
        name = ((d.metadata or {}).get("sheet") or "").strip()
        n = name.replace(" ", "").replace("\u3000", "")
        nl = n.lower()
        tier0_zh = "风险" in n and ("分析" in n or "评估" in n or "管理" in n)
        tier0_en = "risk" in nl and any(
            k in nl for k in ("analysis", "analys", "management", "assessment", "mitigation")
        )
        if tier0_zh or tier0_en:
            tier = 0
        elif "风险" in n or "risk" in nl:
            tier = 1
        else:
            tier = 2
        scored.append((tier, i, d))
    scored.sort(key=lambda x: (x[0], x[1]))
    return [t[2] for t in scored]


def _load_xlsx_with_openpyxl(path: Path) -> List[Document]:
    """使用 openpyxl 读取 xlsx **全部工作表**（每表一段文本，带表名标题）。审核/追溯须覆盖所有 sheet，勿依赖 Unstructured 默认仅前几表的行为。"""
    import openpyxl
    path = Path(path)
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    docs = []
    try:
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append("\t".join(_xlsx_cell_to_text(c) for c in row))
            text = "\n".join(rows).strip()
            if text:
                hint = _excel_risk_sheet_content_hint(sheet.title)
                block = f"【Excel 工作表：{sheet.title}】\n{hint}{text}"
                docs.append(Document(
                    page_content=block,
                    metadata={"source_file": path.name, "file_type": ".xlsx", "sheet": sheet.title},
                ))
    finally:
        wb.close()
    docs = _prioritize_excel_sheet_documents(docs)
    return docs if docs else [Document(page_content="(空表)", metadata={"source_file": path.name, "file_type": ".xlsx"})]


def _load_xls_all_sheets_with_xlrd(path: Path) -> List[Document]:
    """使用 xlrd 读取旧版 .xls 的**全部工作表**（每表带标题）。未安装 xlrd 时返回空列表。"""
    path = Path(path)
    try:
        import xlrd
    except ImportError:
        return []
    try:
        book = xlrd.open_workbook(str(path))
    except Exception:
        return []
    docs: List[Document] = []
    for si in range(book.nsheets):
        sh = book.sheet_by_index(si)
        rows = []
        for ri in range(sh.nrows):
            rows.append(
                "\t".join(_xlsx_cell_to_text(sh.cell_value(ri, ci)) for ci in range(sh.ncols))
            )
        text = "\n".join(rows).strip()
        if text:
            hint = _excel_risk_sheet_content_hint(sh.name)
            block = f"【Excel 工作表：{sh.name}】\n{hint}{text}"
            docs.append(
                Document(
                    page_content=block,
                    metadata={"source_file": path.name, "file_type": ".xls", "sheet": sh.name},
                )
            )
    return _prioritize_excel_sheet_documents(docs)


def _find_soffice() -> Optional[str]:
    """查找系统已安装的 LibreOffice soffice 可执行路径（Windows 常见路径或 PATH）。"""
    # 环境变量优先，便于用户自定义安装路径
    env_path = os.environ.get("LIBREOFFICE_PATH", "").strip()
    if env_path:
        p = Path(env_path)
        if p.is_file() and p.exists():
            return str(p.resolve())
        if p.is_dir():
            # Windows 下 soffice 在 program 子目录
            for subdir, exe in (("", "soffice.exe"), ("", "soffice"), ("program", "soffice.exe"), ("program", "soffice")):
                candidate = p / subdir / exe if subdir else p / exe
                if candidate.exists():
                    return str(candidate.resolve())
    if os.name == "nt":
        for base in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(base).exists():
                return base
        for name in ("soffice.exe", "soffice", "libreoffice.exe", "libreoffice"):
            found = shutil.which(name)
            if found:
                return found
    else:
        for name in ("soffice", "libreoffice"):
            found = shutil.which(name)
            if found:
                return found
    return None


def _find_wps() -> Optional[str]:
    """查找 Windows 下已安装的 WPS 文字（wps.exe）。可通过环境变量 WPS_PATH 指定。"""
    if os.name != "nt":
        return None
    env_path = os.environ.get("WPS_PATH", "").strip()
    if env_path:
        p = Path(env_path)
        if p.is_file() and p.suffix.lower() in (".exe", ""):
            return str(p.resolve())
        if p.is_dir():
            for name in ("wps.exe", "et.exe"):
                candidate = p / name
                if candidate.exists():
                    return str(candidate.resolve())
    # 常见安装路径：Kingsoft WPS Office
    for base in (
        Path(os.environ.get("ProgramFiles(X86)", r"C:\Program Files (x86)")) / "Kingsoft" / "WPS Office",
        Path(os.environ.get("ProgramFiles", r"C:\Program Files")) / "Kingsoft" / "WPS Office",
    ):
        if not base.exists():
            continue
        for sub in base.iterdir():
            if sub.is_dir():
                wps_exe = sub / "wps.exe"
                if wps_exe.exists():
                    return str(wps_exe.resolve())
        return None


def _convert_doc_to_docx_with_pandoc(doc_path: Path) -> Tuple[Optional[Path], str]:
    """使用 pandoc 将 .doc 转为 .docx。返回 (生成的 .docx 路径或 None, 失败原因)。"""
    doc_path = doc_path.resolve()
    if not doc_path.exists():
        return None, "文件不存在"
    pandoc_exe = shutil.which("pandoc")
    if not pandoc_exe:
        return None, "未找到 pandoc（请安装并加入 PATH）"
    fd, out_file = tempfile.mkstemp(suffix=".docx", prefix="aicheckword_pandoc_")
    os.close(fd)
    try:
        result = subprocess.run(
            [pandoc_exe, "-f", "doc", "-t", "docx", "-o", out_file, str(doc_path)],
            capture_output=True,
            text=True,
            timeout=30,
            encoding="utf-8",
            errors="replace",
        )
        if result.returncode == 0 and Path(out_file).exists() and Path(out_file).stat().st_size > 0:
            return Path(out_file), ""
        try:
            Path(out_file).unlink(missing_ok=True)
        except Exception:
            pass
        err = (result.stderr or result.stdout or "").strip()[:400]
        return None, f"pandoc 失败: {err}" if err else f"pandoc 退出码 {result.returncode}"
    except subprocess.TimeoutExpired:
        try:
            Path(out_file).unlink(missing_ok=True)
        except Exception:
            pass
        return None, "pandoc 超时"
    except Exception as e:
        try:
            Path(out_file).unlink(missing_ok=True)
        except Exception:
            pass
        return None, str(e).strip() or type(e).__name__


def _convert_doc_to_docx_with_wps(doc_path: Path) -> Tuple[Optional[Path], str]:
    """使用 WPS Office COM 接口将 .doc 转为 .docx（仅 Windows，需已安装 WPS 与 pywin32）。
    多组 ProgID 依次尝试（部分 WPS 版本或安装方式注册名不同）。返回 (生成的 docx 路径或 None, 失败原因)。"""
    if os.name != "nt":
        return None, "非 Windows 平台"
    doc_path = doc_path.resolve()
    if not doc_path.exists():
        return None, "文件不存在"
    try:
        import win32com.client
    except ImportError:
        return None, "未安装 pywin32（pip install pywin32）"
    # 多种 ProgID：Kwps.Application 常见；部分环境为小写或 wps.Application；新版本可能禁用 COM
    for prog_id in ("Kwps.Application", "kwps.application", "WPS.Application", "wps.Application"):
        wps = None
        doc = None
        out_file = None
        try:
            wps = win32com.client.dynamic.Dispatch(prog_id)
            if not wps:
                continue
            # WPS 为兼容性会将 Name 报告为 "Microsoft Word"，不能以此判断是否真正的 MS Word。
            try:
                exe_path = (getattr(wps, "Path", None) or "").strip()
                is_genuine_msword = exe_path and "Microsoft" in exe_path and "Kingsoft" not in exe_path and "WPS" not in exe_path
                if is_genuine_msword:
                    wps.Quit()
                    continue
            except Exception:
                pass
            wps.Visible = False
            doc_full = str(doc_path.resolve())
            doc = wps.Documents.Open(doc_full)
            fd, out_file = tempfile.mkstemp(suffix=".docx", prefix="aicheckword_wps_")
            os.close(fd)
            doc.SaveAs2(out_file, FileFormat=12)
            doc.Close(False)
            doc = None
            wps.Quit()
            wps = None
            return Path(out_file), ""
        except Exception as e:
            err = str(e).strip() or type(e).__name__
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if wps is not None:
                try:
                    wps.Quit()
                except Exception:
                    pass
            if out_file and Path(out_file).exists():
                try:
                    Path(out_file).unlink(missing_ok=True)
                except Exception:
                    pass
            if prog_id == "wps.Application":
                return None, (
                    f"WPS COM 失败: {err}。"
                    "若为「无效的类字符串」：部分 WPS 版本默认关闭 COM，请在 WPS 设置中开启「兼容 Microsoft Office 的 COM 加载项」或改用 LibreOffice。"
                )
    return None, "未找到 WPS COM（已尝试 Kwps/kwps/WPS/wps.Application）；可改用 LibreOffice 或手动将 .doc 另存为 .docx。"


def _convert_doc_to_docx_with_libreoffice(doc_path: Path) -> Tuple[Optional[Path], str]:
    """使用 LibreOffice 将 .doc 转为 .docx。返回 (生成的 .docx 路径或 None, 失败原因)。"""
    soffice = _find_soffice()
    if not soffice:
        return None, "未找到 soffice（请设置环境变量 LIBREOFFICE_PATH 指向安装目录或 program/soffice.exe）"
    doc_path = doc_path.resolve()
    if not doc_path.exists():
        return None, "文件不存在"
    out_dir = tempfile.mkdtemp(prefix="aicheckword_doc_")
    result = None
    try:
        env = os.environ.copy()
        env["SAL_USE_VCLPLUGIN"] = "headless"
        user_inst = tempfile.mkdtemp(prefix="LibreOffice_Conversion_")
        try:
            user_inst_uri = Path(user_inst).as_uri()
            cmd = [
                soffice,
                "--headless",
                f"-env:UserInstallation={user_inst_uri}",
                "--convert-to",
                "docx",
                "--outdir",
                out_dir,
                str(doc_path),
            ]
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
                env=env,
                cwd=out_dir,
                encoding="utf-8",
                errors="replace",
            )
        finally:
            shutil.rmtree(user_inst, ignore_errors=True)
        if result is None or result.returncode != 0:
            err = (result.stderr or result.stdout or "").strip() if result else ""
            if err:
                err = err[:500] + "..." if len(err) > 500 else err
            return None, f"soffice 转换退出码 {getattr(result, 'returncode', None)}" + (f": {err}" if err else "")
        docx_name = doc_path.stem + ".docx"
        docx_path = Path(out_dir) / docx_name
        if docx_path.exists():
            fd, tmp = tempfile.mkstemp(suffix=".docx", prefix="aicheckword_")
            os.close(fd)
            shutil.copy2(docx_path, tmp)
            shutil.rmtree(out_dir, ignore_errors=True)
            return Path(tmp), ""
        return None, "soffice 未生成 .docx 文件"
    except subprocess.TimeoutExpired:
        try:
            shutil.rmtree(out_dir, ignore_errors=True)
        except Exception:
            pass
        return None, "soffice 转换超时（60 秒）"
    except Exception as e:
        try:
            shutil.rmtree(out_dir, ignore_errors=True)
        except Exception:
            pass
        return None, str(e).strip() or type(e).__name__


def _load_word_with_unstructured(path: Path, suffix: str) -> List[Document]:
    """使用 UnstructuredWordDocumentLoader 加载 .doc 或 .docx（兼容旧版 .doc，需 unstructured 且推荐安装 LibreOffice）。"""
    loader = UnstructuredWordDocumentLoader(str(path))
    docs = loader.load()
    for d in docs:
        d.metadata.setdefault("source_file", path.name)
        d.metadata.setdefault("file_type", suffix)
    return docs if docs else [Document(page_content="(空)", metadata={"source_file": path.name, "file_type": suffix})]


# 签批、日期等关键词：用于判断表格行/段落是否与「签署」相关（中英）
_SIGNOFF_KW_RE = re.compile(
    r"(编制|拟制|编写|作成|审核|校核|核对|批准|审定|核准|会签|签发|授权|"
    r"签名|签章|签署|签批|编制人|审核人|批准人|"
    r"日期|年\s*月\s*日|签署日期|生效日期|"
    r"author|reviewer|approver|prepared\s*by|checked\s*by|approved\s*by|date\b)",
    re.I,
)


def _ooxml_element_has_visual(element) -> bool:
    """段落/单元格 OOXML 中是否含图、嵌入对象等（Docx2txt 通常无法转为文字）。"""
    try:
        for node in element.iter():
            tag = node.tag
            if not isinstance(tag, str):
                continue
            tl = tag.lower()
            tail = tag.split("}")[-1].lower()
            if tail in ("drawing", "pict", "object", "binobject", "shape", "imagedata", "oleobject", "inline"):
                return True
            if "drawing" in tl or "picture" in tl or "imagedata" in tl or "wps:wsp" in tl:
                return True
    except Exception:
        return False
    return False


def _build_docx_signoff_image_supplement(docx_path: Path) -> str:
    """
    扫描 .docx 中签批相关区域附近的嵌入图/对象，生成简短说明附加到提取文本后，
    避免模型将「图片签名/日期」误判为未签署（docx2txt 不输出图片）。
    """
    docx_path = Path(docx_path)
    if not docx_path.is_file() or docx_path.suffix.lower() != ".docx":
        return ""
    try:
        from docx import Document as DocxDocument
    except Exception:
        return ""
    try:
        doc = DocxDocument(str(docx_path))
    except Exception:
        return ""

    hints: List[str] = []
    seen = set()

    def add_hint(msg: str) -> None:
        if msg not in seen:
            seen.add(msg)
            hints.append(msg)

    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        text = (p.text or "").strip()
        vis = _ooxml_element_has_visual(p._element)
        if vis and _SIGNOFF_KW_RE.search(text):
            add_hint(
                f"正文段落含签批相关用语且该段含嵌入图/对象（可能为签名或日期图），勿判「签署空白」。摘录：{text[:100]}"
            )
        elif vis and len(text) < 12 and i > 0:
            prev_t = (paras[i - 1].text or "").strip()
            if _SIGNOFF_KW_RE.search(prev_t):
                add_hint(
                    "签批标签行紧邻段落几乎无文字但含嵌入图/对象，常见于签名/日期占位列，勿判「签署空白」。"
                )

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            row_text = " ".join(
                ((c.text or "").replace("\n", " ").strip()) for c in row.cells
            )
            if not _SIGNOFF_KW_RE.search(row_text):
                continue
            for ci, cell in enumerate(row.cells):
                cell_has_vis = any(
                    _ooxml_element_has_visual(p._element) for p in cell.paragraphs
                )
                if cell_has_vis:
                    add_hint(
                        f"第 {ti + 1} 个表格第 {ri + 1} 行含编审批/日期等关键词，且第 {ci + 1} 列存在嵌入图/对象（视为可能已图片签署）。"
                    )

    if not hints:
        return ""

    cap = 28
    body = "\n".join(f"- {h}" for h in hints[:cap])
    more = f"\n- …共检出 {len(hints)} 处，其余已省略" if len(hints) > cap else ""
    return (
        "【Word 版式检测结果（嵌入图像，非 OCR）】\n"
        "说明：正文由文本提取得到，手写签名/电子签章/日期常以嵌入图片存在；以下由程序扫描 Word 结构（OOXML）得到。\n"
        f"{body}{more}"
    )


def _append_docx_signoff_supplement(docs: List[Document], docx_path: Path) -> None:
    """将签批区嵌入图检测结果拼入首个 Document 的 page_content。"""
    if not docs:
        return
    sup = _build_docx_signoff_image_supplement(docx_path)
    if not sup:
        return
    first = docs[0]
    base = (first.page_content or "").rstrip()
    first.page_content = f"{base}\n\n{sup}"


# docx2txt 对复杂 Word 表格列（如 URS 编号）常漏字或错位；用 python-docx 按单元格再导出一遍
_DOCX_TABLE_MAX_ROWS = 4000
_DOCX_TABLE_MAX_CELL_CHARS = 4000


def _docx_paragraph_ooxml_text(paragraph) -> str:
    """
    段落纯文本，近似 Word「最终状态」：**不采集** w:del / w:moveFrom 子树内的文字。
    若把 w:delText（已删除的旧编号如 C201）与 w:ins 内新编号（CS02）一并拼接，会得到 C201CS02，
    模型易误读为 CS012；审核应以**修订后**为准。
    w:ins 与未删除 run 中的 w:t、w:instrText（域代码可见文本）仍保留。
    """
    try:
        from docx.oxml.ns import qn

        root = paragraph._p
    except Exception:
        return (getattr(paragraph, "text", None) or "").strip()
    skip_roots = frozenset({qn("w:del"), qn("w:moveFrom")})

    def walk(node) -> str:
        if node.tag in skip_roots:
            return ""
        if node.tag == qn("w:t"):
            return node.text or ""
        if node.tag == qn("w:delText"):
            return ""
        if node.tag == qn("w:instrText"):
            return node.text or ""
        if node.tag == qn("w:tab"):
            return "\t"
        if node.tag in (qn("w:br"), qn("w:cr")):
            return "\n"
        return "".join(walk(ch) for ch in node)

    raw = walk(root)
    raw = raw.replace("\t", " ")
    raw = " ".join(raw.split())
    return raw.strip()


# 无分隔符时「旧号+新号」粘连（含未开修订、纯粘贴导致）
_REVISION_ID_COLLISION_PATTERNS = (
    (re.compile(r"C(\d{3})CS(\d{1,4})\b", re.I), r"【以修订后为准】CS\2（已废止 C\1）"),
    (re.compile(r"CE(\d{1,4})CS(\d{1,4})\b", re.I), r"【以修订后为准】CS\2（已废止 CE\1）"),
    (re.compile(r"CS(\d{1,4})CS(\d{1,4})\b", re.I), r"【以修订后为准】CS\2（已取代 CS\1）"),
)


def _normalize_revision_id_collisions(text: str) -> str:
    """将明显为「修订串联」的粘连编号标成以新号为准，供模型只审核最终编号。"""
    if not text:
        return text
    s = text
    for _ in range(8):
        prev = s
        for rx, repl in _REVISION_ID_COLLISION_PATTERNS:
            s = rx.sub(repl, s)
        if s == prev:
            break
    return s


def _build_docx_final_revision_paragraphs_text(docx_path: Path) -> str:
    """正文段落（不含表格），与表格结构化摘录分工；均使用跳过删除修订的 OOXML 规则。"""
    try:
        from docx import Document as DocxDocument
    except Exception:
        return ""
    try:
        doc = DocxDocument(str(docx_path))
    except Exception:
        return ""
    parts: List[str] = []
    for p in doc.paragraphs:
        t = _docx_paragraph_ooxml_text(p)
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _try_replace_docx_body_with_final_revision_view(docs: List[Document], docx_path: Path) -> None:
    """用「忽略删除修订」的段落文本替换 Docx2txt 首块，减少旧编号参与拼接。"""
    if not docs:
        return
    raw = (docs[0].page_content or "").strip()
    alt = _build_docx_final_revision_paragraphs_text(docx_path).strip()
    if len(alt) < 40:
        return
    suspicious = bool(re.search(r"C\d{3}CS\d", raw, re.I)) or bool(
        re.search(r"CS\d{1,4}CS\d{1,4}", raw, re.I)
    )
    if suspicious:
        docs[0].page_content = alt
        return
    if len(alt) >= max(120, int(len(raw) * 0.18)):
        docs[0].page_content = alt


def _apply_revision_collision_normalize_to_docs(docs: List[Document]) -> None:
    for d in docs or []:
        if d.page_content:
            d.page_content = _normalize_revision_id_collisions(d.page_content)


def _docx_cell_plaintext(cell) -> str:
    parts = []
    for p in cell.paragraphs:
        t = _docx_paragraph_ooxml_text(p)
        if t:
            parts.append(t)
    s = " ".join(parts).strip()
    if len(s) > _DOCX_TABLE_MAX_CELL_CHARS:
        s = s[:_DOCX_TABLE_MAX_CELL_CHARS] + "…"
    return s


def _docx_row_cells_merged_aware(row) -> List[str]:
    """
    水平合并单元格在 python-docx 中会对每个网格列返回同一底层单元格（_tc 相同），
    须按 _tc 去重；**禁止**按「文本相同」去重，否则两列恰好同字（如占位「/」、重复列名）
    会被错误合并，导致后续列左移，出现 CE01→CS0、URS001→URS1 等错位误读。
    """
    out: List[str] = []
    last_tc_id = None
    for c in row.cells:
        try:
            tc_id = id(c._tc)
        except Exception:
            tc_id = id(c)
        if last_tc_id is not None and tc_id == last_tc_id:
            continue
        last_tc_id = tc_id
        out.append(_docx_cell_plaintext(c))
    return out


def _docx_table_to_tsv(table) -> str:
    lines: List[str] = []
    for ri, row in enumerate(table.rows):
        if ri >= _DOCX_TABLE_MAX_ROWS:
            lines.append(f"…（该表仅摘录前 {_DOCX_TABLE_MAX_ROWS} 行）…")
            break
        cells = _docx_row_cells_merged_aware(row)
        lines.append("\t".join(cells))
    return "\n".join(lines).strip()


def _build_docx_tables_plaintext(docx_path: Path, max_total_chars: int = 450_000) -> str:
    """
    遍历 .docx 中全部表格（正文 + 各节页眉/页脚/首页页眉脚），按行导出为 TSV，
    供审核/追溯使用；与 Docx2txtLoader 正文互补，减少「表格内 URS 未检索到」误报。
    """
    docx_path = Path(docx_path)
    if not docx_path.is_file() or docx_path.suffix.lower() != ".docx":
        return ""
    try:
        from docx import Document as DocxDocument
    except Exception:
        return ""
    try:
        doc = DocxDocument(str(docx_path))
    except Exception:
        return ""

    seen_tbl_ids = set()
    blocks: List[str] = []
    n = 0

    def emit_block(title: str, table) -> None:
        nonlocal n
        try:
            tid = id(table._tbl)
        except Exception:
            tid = id(table)
        if tid in seen_tbl_ids:
            return
        text = _docx_table_to_tsv(table)
        if not text:
            return
        seen_tbl_ids.add(tid)
        n += 1
        blocks.append(f"{title} {n}】\n{text}")

    for table in doc.tables:
        emit_block("【Word 正文表格", table)

    for si, sec in enumerate(doc.sections):
        for hf_label, part_attr in (
            ("页眉", "header"),
            ("页脚", "footer"),
        ):
            try:
                part = getattr(sec, part_attr)
            except Exception:
                continue
            if part is None:
                continue
            try:
                for table in part.tables:
                    emit_block(f"【Word 第{si + 1}节{hf_label}内表格", table)
            except Exception:
                continue
        if getattr(sec, "different_first_page_header_footer", False):
            for hf_label, part_attr in (
                ("首页页眉", "first_page_header"),
                ("首页页脚", "first_page_footer"),
            ):
                try:
                    part = getattr(sec, part_attr)
                except Exception:
                    continue
                if part is None:
                    continue
                try:
                    for table in part.tables:
                        emit_block(f"【Word 第{si + 1}节{hf_label}内表格", table)
                except Exception:
                    continue

    if not blocks:
        return ""
    intro = (
        "【Word 表格结构化摘录（python-docx 按单元格提取；合并单元格已按网格去重，两列同字也会分栏保留，避免编号错位）】\n"
    )
    body = "\n\n".join(blocks)
    full = intro + body
    if len(full) > max_total_chars:
        full = (
            full[: max_total_chars - 100].rstrip()
            + "\n\n……（表格摘录总长已达上限，后续表已省略）……"
        )
    return full


def _append_docx_tables_plaintext(docs: List[Document], docx_path: Path) -> None:
    """将结构化表格文本拼入首个 Document（与签批补充相同策略）。"""
    if not docs:
        return
    extra = _build_docx_tables_plaintext(docx_path)
    if not extra:
        return
    first = docs[0]
    base = (first.page_content or "").rstrip()
    first.page_content = f"{base}\n\n{extra}"


def _append_signoff_block_to_first_doc(docs: List[Document], block: str) -> None:
    if not docs or not (block or "").strip():
        return
    first = docs[0]
    first.page_content = (first.page_content or "").rstrip() + "\n\n" + block.strip()


def _build_xlsx_signoff_image_supplement(xlsx_path: Path) -> str:
    """
    检测 .xlsx 工作簿中是否存在嵌入图片（openpyxl），且工作簿文本中含编审批/日期等关键词。
    xls 老格式无法用本逻辑扫描，返回空字符串。
    """
    xlsx_path = Path(xlsx_path)
    if not xlsx_path.is_file() or xlsx_path.suffix.lower() != ".xlsx":
        return ""
    import openpyxl

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=False)
    except Exception:
        return ""

    hints: List[str] = []
    seen = set()
    sheet_info: List[Tuple[str, str, int]] = []

    def add_hint(msg: str) -> None:
        if msg not in seen:
            seen.add(msg)
            hints.append(msg)

    try:
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append("\t".join(str(c) if c is not None else "" for c in row))
            sheet_blob = "\n".join(rows)
            nimg = len(getattr(ws, "_images", None) or [])
            sheet_info.append((ws.title, sheet_blob, nimg))
    finally:
        try:
            wb.close()
        except Exception:
            pass

    wb_blob = "\n".join(b for _, b, _ in sheet_info)
    wb_kw = bool(_SIGNOFF_KW_RE.search(wb_blob))

    for title, blob, nimg in sheet_info:
        if nimg <= 0 or not _SIGNOFF_KW_RE.search(blob):
            continue
        add_hint(
            f"工作表「{title}」检出 {nimg} 处嵌入图片，且该表存在编审批/日期等相关文字；"
            f"签名或签署日期可能位于图片中，勿仅凭单元格无文字判定「签署空白」。"
        )

    if not hints and wb_kw:
        titled = [(t, n) for t, b, n in sheet_info if n > 0]
        total_imgs = sum(n for _, n in titled)
        if total_imgs > 0:
            parts = [f"「{t}」{n} 张" for t, n in titled[:12]]
            tail = f" 等共 {total_imgs} 张" if len(titled) > 12 else ""
            add_hint(
                "工作簿内存在编审批/日期等相关文字，且在 "
                + "、".join(parts)
                + tail
                + " 嵌入图片；签批可能以图片形式出现在任意工作表，勿判「签署空白」。"
            )

    if not hints:
        return ""

    body = "\n".join(f"- {h}" for h in hints[:20])
    more = f"\n- …共 {len(hints)} 条，其余已省略" if len(hints) > 20 else ""
    return (
        "【Excel 嵌入图像检测（非单元格文字）】\n"
        "说明：表格文本由单元格值拼接得到，手写签名/电子签/日期常以浮动或单元格内图片存在，以下由 openpyxl 扫描工作簿得到。\n"
        f"{body}{more}"
    )


def _pdf_resources_has_raster_image(resources, depth: int = 0, max_depth: int = 5) -> bool:
    """递归检查 Resources 下 /XObject 是否含 /Image 或 Form 内嵌图片。"""
    if depth > max_depth or resources is None:
        return False
    try:
        if hasattr(resources, "get_object"):
            resources = resources.get_object()
        xobjs = resources.get("/XObject")
        if xobjs is None:
            return False
        if hasattr(xobjs, "get_object"):
            xobjs = xobjs.get_object()
        for _k in xobjs:
            obj = xobjs[_k]
            if hasattr(obj, "get_object"):
                obj = obj.get_object()
            try:
                sub = obj.get("/Subtype")
            except Exception:
                continue
            if sub == "/Image":
                return True
            if sub == "/Form":
                subres = obj.get("/Resources")
                if subres and _pdf_resources_has_raster_image(subres, depth + 1, max_depth):
                    return True
    except Exception:
        return False
    return False


def _build_pdf_signoff_image_supplement(pdf_path: Path, max_pages: int = MAX_PDF_PAGES) -> str:
    """检测 PDF 各页是否含嵌入位图对象，且全文含签批关键词。"""
    pdf_path = Path(pdf_path)
    if not pdf_path.is_file() or pdf_path.suffix.lower() != ".pdf":
        return ""
    try:
        from pypdf import PdfReader
    except Exception:
        return ""

    try:
        reader = PdfReader(str(pdf_path))
    except Exception:
        return ""

    pages = list(reader.pages)[:max_pages]
    if not pages:
        return ""

    full_text = ""
    for p in pages:
        try:
            full_text += (p.extract_text() or "") + "\n"
        except Exception:
            full_text += "\n"

    if not _SIGNOFF_KW_RE.search(full_text):
        return ""

    img_pages: List[int] = []
    for i, page in enumerate(pages):
        try:
            res = page.get("/Resources")
            if res and _pdf_resources_has_raster_image(res):
                img_pages.append(i + 1)
        except Exception:
            continue

    if not img_pages:
        return ""

    if len(img_pages) <= 25:
        pages_str = "、".join(str(n) for n in img_pages)
    else:
        pages_str = f"第 1—{max(img_pages)} 页中至少 {len(img_pages)} 页（已抽样提示）"

    return (
        "【PDF 嵌入图像检测（页面 XObject）】\n"
        "说明：正文由 PDF 文本层提取；签名/日期常以页面内嵌图片存在。以下页检测到嵌入位图对象，且文档中出现编审批/日期等相关文字。\n"
        f"- 含嵌入图片的页码（约）：{pages_str}\n"
        "- 请勿仅凭文本层无手写姓名字符即报「签署空白」；除非规范明确要求必须为可检索文本。"
    )


def load_single_file(file_path) -> List[Document]:
    """加载单个文件，返回 Document 列表。失败时抛出带完整原因的异常。"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in LOADER_MAP:
        raise ValueError(f"不支持的文件格式: {suffix}，支持的格式: {list(LOADER_MAP.keys())}")

    loader_cls = LOADER_MAP[suffix]
    last_error = None

    # Excel：xlsx **优先 openpyxl 全表**（Unstructured 常只解析部分 sheet）；失败再回退 Unstructured。
    # .xls：优先 xlrd 全表；失败再 Unstructured。
    if suffix in (".xlsx", ".xls"):
        if suffix == ".xlsx":
            docs = []
            try:
                docs = _load_xlsx_with_openpyxl(path)
            except Exception as e:
                last_error = e
                docs = []
            total_chars = sum(len((d.page_content or "").replace("(空表)", "")) for d in docs)
            if not docs or all("(空表)" in (d.page_content or "") for d in docs) or total_chars < 20:
                try:
                    loader = loader_cls(str(path))
                    docs = loader.load()
                    for d in docs:
                        st = (d.page_content or "").strip()
                        if st and not st.startswith("【Excel 工作表："):
                            sheet = (d.metadata or {}).get("sheet") or (d.metadata or {}).get("page_name") or ""
                            prefix = f"【Excel 工作表：{sheet}】\n" if sheet else "【Excel 工作表】\n"
                            d.page_content = prefix + st
                except Exception as e2:
                    if not docs:
                        err_msg = f"Excel(xlsx) 加载失败: openpyxl 未得到有效内容；Unstructured 也失败: {e2!s}"
                        raise RuntimeError(err_msg) from (last_error if last_error is not None else e2)
            docs = _prioritize_excel_sheet_documents(docs)
            try:
                _append_signoff_block_to_first_doc(docs, _build_xlsx_signoff_image_supplement(path))
            except Exception:
                pass
        else:
            docs = _load_xls_all_sheets_with_xlrd(path)
            total_chars = sum(len(d.page_content or "") for d in docs) if docs else 0
            if not docs or total_chars < 20:
                try:
                    loader = loader_cls(str(path))
                    docs = loader.load()
                    for d in docs:
                        st = (d.page_content or "").strip()
                        if st and not st.startswith("【Excel 工作表："):
                            sheet = (d.metadata or {}).get("sheet") or (d.metadata or {}).get("page_name") or ""
                            prefix = f"【Excel 工作表：{sheet}】\n" if sheet else "【Excel 工作表】\n"
                            d.page_content = prefix + st
                except Exception as e:
                    raise RuntimeError(
                        f"xls 文件加载失败（已尝试 xlrd 全表，仍失败；可 pip install xlrd）：{e!s}"
                    ) from e
            docs = _prioritize_excel_sheet_documents(docs)
    elif suffix == ".pdf":
        # PDF：用 lazy_load + 页数上限，避免大文件或复杂版式卡死；检测图片版/扫描版
        try:
            loader = loader_cls(str(path))
            docs = []
            for i, doc in enumerate(loader.lazy_load()):
                if i >= MAX_PDF_PAGES:
                    break
                doc.metadata.setdefault("source_file", path.name)
                doc.metadata.setdefault("file_type", ".pdf")
                doc.metadata.setdefault("page", i + 1)
                docs.append(doc)
            if not docs:
                raise RuntimeError(
                    "该 PDF 无法解析出任何页面。若为扫描件或纯图片 PDF，当前不支持（需先做 OCR 或使用文本型 PDF）。"
                )
            total_text = "".join(d.page_content for d in docs)
            if len(total_text.strip()) < MIN_PDF_TEXT_LEN:
                # 先尝试 AI OCR 兜底，避免扫描版直接失败
                ocr_docs, ocr_err = _pdf_ocr_with_multimodal_llm(path, max_pages=MAX_PDF_OCR_PAGES)
                if ocr_docs:
                    docs = ocr_docs
                else:
                    raise RuntimeError(
                        f"该 PDF 几乎未提取到文字（仅 {len(total_text.strip())} 字），已尝试 AI OCR 兜底但失败：{ocr_err}"
                    )
        except Exception as e:
            if "无法解析" in str(e) or "几乎未提取" in str(e):
                raise
            raise RuntimeError(f"PDF 加载失败: {e!s}") from e
    elif suffix in (".docx", ".doc"):
        # Word：.docx 优先用 Docx2txtLoader；遇 BadZipFile（实为 .doc 或损坏）或 .doc 则用 Unstructured 兼容旧版 .doc
        if suffix == ".docx":
            try:
                loader = Docx2txtLoader(str(path))
                docs = loader.load()
            except zipfile.BadZipFile:
                try:
                    docs = _load_word_with_unstructured(path, suffix)
                except Exception as e2:
                    raise RuntimeError(
                        "该文件不是有效的 .docx 格式（可能为旧版 .doc 或损坏）。"
                        "已尝试用 Unstructured 解析仍失败，请确认文件可被 Word 打开，或安装 LibreOffice 后重试。"
                        f" 详情: {e2!s}"
                    ) from e2
            except Exception as e:
                if "zip" in str(e).lower() or "not a zip" in str(e).lower():
                    try:
                        docs = _load_word_with_unstructured(path, suffix)
                    except Exception as e2:
                        raise RuntimeError(
                            "该 Word 文件无法按 .docx 解析，已尝试 Unstructured 仍失败。"
                            "若为旧版 .doc 或损坏，请安装 LibreOffice 或改用 .docx。"
                            f" 详情: {e2!s}"
                        ) from e2
                else:
                    raise RuntimeError(f"Word 加载失败: {e!s}") from e
        else:
            # .doc 旧版格式：先统一尝试转为 .docx 再训练（WPS / LibreOffice / Pandoc），收集失败原因便于排查
            docx_path = None
            reasons = []
            if os.name == "nt":
                docx_path, wps_err = _convert_doc_to_docx_with_wps(path)
                if not (docx_path and docx_path.exists()) and wps_err:
                    reasons.append("WPS: " + wps_err)
            if not (docx_path and docx_path.exists()):
                docx_path, lo_err = _convert_doc_to_docx_with_libreoffice(path)
                if not (docx_path and docx_path.exists()) and lo_err:
                    reasons.append("LibreOffice: " + lo_err)
            if not (docx_path and docx_path.exists()):
                docx_path, pd_err = _convert_doc_to_docx_with_pandoc(path)
                if not (docx_path and docx_path.exists()) and pd_err:
                    reasons.append("Pandoc: " + pd_err)
            if docx_path and docx_path.exists():
                try:
                    loader = Docx2txtLoader(str(docx_path))
                    docs = loader.load()
                    for d in docs:
                        d.metadata.setdefault("source_file", path.name)
                        d.metadata.setdefault("file_type", ".doc")
                    if not docs:
                        docs = [Document(page_content="(空)", metadata={"source_file": path.name, "file_type": ".doc"})]
                    try:
                        _try_replace_docx_body_with_final_revision_view(docs, Path(docx_path))
                    except Exception:
                        pass
                    try:
                        _append_docx_signoff_supplement(docs, Path(docx_path))
                    except Exception:
                        pass
                    try:
                        _append_docx_tables_plaintext(docs, Path(docx_path))
                    except Exception:
                        pass
                    try:
                        _apply_revision_collision_normalize_to_docs(docs)
                    except Exception:
                        pass
                finally:
                    try:
                        Path(docx_path).unlink(missing_ok=True)
                    except Exception:
                        pass
            else:
                # 转换均失败时再试 Unstructured；若仍失败则报错并附上各方式失败原因
                err_unstruct = None
                try:
                    docs = _load_word_with_unstructured(path, suffix)
                except Exception as e:
                    err_unstruct = e
                    docs = None
                if docs is None or (docs and all((d.page_content or "").strip() in ("", "(空)") for d in docs)):
                    detail = "; ".join(reasons) if reasons else "未找到可用转换方式"
                    hint = (
                        "推荐做法：安装 LibreOffice，并设置环境变量 LIBREOFFICE_PATH 指向安装目录（如 C:\\Program Files\\LibreOffice）或 program\\soffice.exe，"
                        "然后重启本程序。若使用 WPS：需安装 pywin32，且部分 WPS 版本需在设置中开启 COM；也可在 WPS/Word 中把 .doc 另存为 .docx 后上传。"
                    )
                    raise RuntimeError(
                        "旧版 .doc 无法转为 .docx。失败原因: %s。%s" % (detail, hint)
                    ) from err_unstruct
        if suffix == ".docx":
            try:
                _try_replace_docx_body_with_final_revision_view(docs, path)
            except Exception:
                pass
            try:
                _append_docx_signoff_supplement(docs, path)
            except Exception:
                pass
            try:
                _append_docx_tables_plaintext(docs, path)
            except Exception:
                pass
            try:
                _apply_revision_collision_normalize_to_docs(docs)
            except Exception:
                pass
        for doc in docs:
            doc.metadata.setdefault("source_file", path.name)
            doc.metadata.setdefault("file_type", suffix)
    else:
        try:
            loader = loader_cls(str(path))
            docs = loader.load()
        except Exception as e:
            # 仅对纯文本类用 TextLoader 兜底
            if suffix == ".txt":
                try:
                    loader = TextLoader(str(path), encoding="utf-8")
                    docs = loader.load()
                except Exception as e2:
                    raise RuntimeError(f"文本加载失败: {e!s}; 备用编码失败: {e2!s}") from e
            else:
                raise RuntimeError(f"加载失败: {e!s}") from e

    for doc in docs:
        doc.metadata.setdefault("source_file", path.name)
        doc.metadata.setdefault("file_type", suffix)

    return docs


def is_archive(path) -> bool:
    """判断是否为支持的压缩文件"""
    path = Path(path)
    name = path.name.lower()
    if name.endswith(".tar.gz") or name.endswith(".tgz"):
        return True
    return path.suffix.lower() in (".zip", ".tar", ".gz", ".rar")


def _open_zip_with_encoding(archive_path) -> zipfile.ZipFile:
    """用正确编码打开 ZIP，避免中文等文件名乱码。先尝试 gbk（常见于含文件+文件夹的 Windows 压缩包），再 utf-8。"""
    path = Path(archive_path)
    try:
        _ = zipfile.ZipFile(path, "r", metadata_encoding="utf-8")
        _.close()
    except TypeError:
        return zipfile.ZipFile(path, "r")
    for enc in ("gbk", "utf-8", "cp437"):
        try:
            zf = zipfile.ZipFile(path, "r", metadata_encoding=enc)
            names = zf.namelist()
            if names and any("\ufffd" in n for n in names):
                zf.close()
                continue
            return zf
        except (ValueError, UnicodeDecodeError):
            continue
    return zipfile.ZipFile(path, "r", metadata_encoding="utf-8")


def extract_archive(archive_path) -> Tuple[str, List[Path]]:
    """解压到临时目录，返回 (临时目录路径, 文档路径列表)。调用方需在完成后删除临时目录。文件名编码兼容 utf-8/gbk，避免乱码。"""
    archive_path = Path(archive_path)
    temp_dir = tempfile.mkdtemp(prefix="aicheckword_")
    temp_path = Path(temp_dir)
    doc_files = []

    try:
        if archive_path.suffix.lower() == ".zip":
            zf = _open_zip_with_encoding(archive_path)
            try:
                for name in zf.namelist():
                    if name.endswith("/") or "/__MACOSX" in name:
                        continue
                    base = Path(name).name
                    if not base or base.startswith("."):
                        continue
                    zf.extract(name, temp_dir)
                    extracted = temp_path / name
                    if extracted.is_file():
                        suf = extracted.suffix.lower()
                        if suf in SUPPORTED_DOC_EXTENSIONS and not is_deprecated_path(name):
                            doc_files.append(extracted)
            finally:
                zf.close()

        elif archive_path.name.lower().endswith((".tar.gz", ".tgz")) or archive_path.suffix.lower() in (".tar", ".gz"):
            with tarfile.open(archive_path, "r:*") as tf:
                for member in tf.getmembers():
                    if not member.isfile():
                        continue
                    name = member.name
                    if "/__MACOSX" in name or name.startswith("."):
                        continue
                    tf.extract(member, temp_dir)
                    extracted = temp_path / member.name
                    if extracted.is_file():
                        suf = extracted.suffix.lower()
                        if suf in SUPPORTED_DOC_EXTENSIONS and not is_deprecated_path(member.name):
                            doc_files.append(extracted)
        elif archive_path.suffix.lower() == ".rar":
            try:
                import rarfile
            except ImportError:
                shutil.rmtree(temp_dir, ignore_errors=True)
                raise ValueError(
                    "解压 .rar 需要安装 rarfile：pip install rarfile；且系统需安装 UnRAR（如 Windows 下将 UnRAR 加入 PATH，或从 https://www.rarlab.com/rar_add.htm 下载）。也可先将文件改为 .zip 后上传。"
                ) from None
            try:
                with rarfile.RarFile(archive_path) as rf:
                    for name in rf.namelist():
                        if name.endswith("/") or "/__MACOSX" in name or not name.strip():
                            continue
                        base = Path(name).name
                        if not base or base.startswith("."):
                            continue
                        rf.extract(name, temp_dir)
                        extracted = temp_path / name
                        if extracted.is_file():
                            suf = extracted.suffix.lower()
                            if suf in SUPPORTED_DOC_EXTENSIONS and not is_deprecated_path(name):
                                doc_files.append(extracted)
            except Exception as e:
                shutil.rmtree(temp_dir, ignore_errors=True)
                raise ValueError(
                    f"解压 .rar 失败: {e}。请确认已安装 rarfile（pip install rarfile）且系统已安装 UnRAR，或将文件转为 .zip 后上传。"
                ) from e
        else:
            shutil.rmtree(temp_dir, ignore_errors=True)
            raise ValueError("不支持的压缩格式: " + str(archive_path.suffix))

        return temp_dir, doc_files

    except Exception:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise


def load_directory(dir_path) -> List[Document]:
    """加载目录下所有支持格式的文件；路径中含「废弃」的文件或目录跳过。"""
    path = Path(dir_path)
    all_docs = []

    for ext in LOADER_MAP:
        for file_path in path.rglob(f"*{ext}"):
            if is_deprecated_path(file_path):
                continue
            try:
                docs = load_single_file(file_path)
                all_docs.extend(docs)
            except Exception as e:
                print(f"加载文件失败 {file_path}: {e}")

    return all_docs


def split_documents(
    documents: List[Document],
    chunk_size: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
) -> List[Document]:
    """将文档分块"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or settings.chunk_size,
        chunk_overlap=chunk_overlap or settings.chunk_overlap,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )
    return splitter.split_documents(documents)


def load_and_split(file_path) -> List[Document]:
    """加载并分块单个文件"""
    docs = load_single_file(file_path)
    return split_documents(docs)


def load_and_split_directory(dir_path) -> List[Document]:
    """加载并分块目录下所有文件"""
    docs = load_directory(dir_path)
    return split_documents(docs)
