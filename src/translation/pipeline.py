"""
翻译流水线：解析 → 分句 → 调用 LLM 翻译 → 按结构回填 → 写出到新路径（不覆盖原稿）。
支持页眉页脚、sheet 名称、文件名的翻译；DOCX 采用复制原稿再仅替换文字，保留图片、字体与格式。
"""
import re
import shutil
from pathlib import Path
from typing import List, Optional

from .models import TextBlock, SUPPORTED_EXTENSIONS
from .parser import parse_document
from .segment import blocks_to_sentences, apply_translations_to_blocks
from .translator import translate_sentences
from .correction import correct_text, CorrectionStats

_HAS_CJK = re.compile(r"[\u4e00-\u9fff]")


def _has_chinese(s: str) -> bool:
    return bool(s and _HAS_CJK.search(s))


def _translate_stem(
    stem: str,
    collection_name: Optional[str],
    use_kb: bool,
    provider: Optional[str],
    target_lang: str = "en",
    company_overrides: Optional[dict] = None,
    kb_query_extra: Optional[str] = None,
    translation_cache: Optional[dict] = None,
    running_glossary: Optional[dict] = None,
) -> str:
    """翻译文件名/标题等短文本为目标语言，无中文则原样返回。"""
    if not stem or not _has_chinese(stem):
        return stem
    try:
        out = translate_sentences(
            [stem],
            collection_name=collection_name,
            use_kb=use_kb,
            provider=provider,
            target_lang=target_lang or "en",
            company_overrides=company_overrides,
            kb_query_extra=kb_query_extra,
            cache=translation_cache,
            running_glossary=running_glossary,
        )
        return (out[0] or stem).strip() or stem
    except Exception:
        return stem


def _paragraph_has_drawing(para) -> bool:
    """段落内是否含图片/绘制对象（替换全文时会保留）。"""
    try:
        from docx.oxml.ns import qn
        for run in para.runs:
            if run._element.find(qn("w:drawing")) is not None:
                return True
        return False
    except Exception:
        return False


def _replace_paragraph_text_keep_format(para, new_text: str) -> None:
    """
    替换段落文字，尽量保留字体/格式；若段落含图片则只改文字 run，不删图片。
    """
    from docx.oxml.ns import qn
    if _paragraph_has_drawing(para):
        text_runs = [r for r in para.runs if r._element.find(qn("w:drawing")) is None]
        if text_runs:
            text_runs[0].text = new_text
            for r in text_runs[1:]:
                r.text = ""
        return
    if not para.runs:
        para.add_run(new_text)
        return
    try:
        first = para.runs[0]
        font_name = first.font.name
        font_size = first.font.size
        bold = first.bold
        italic = first.italic
    except Exception:
        font_name = font_size = bold = italic = None
    para.clear()
    run = para.add_run(new_text)
    if font_name is not None:
        try:
            run.font.name = font_name
            if font_size is not None:
                run.font.size = font_size
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
        except Exception:
            pass


def _write_docx(source_path: Path, out_path: Path, blocks: List[TextBlock]) -> None:
    """
    复制原稿到输出路径，再仅替换正文/表格/页眉页脚中的文字，保留图片、字体与排版。
    """
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table

    shutil.copy2(source_path, out_path)
    doc = Document(out_path)
    block_idx = 0
    body = doc.element.body
    for child in body:
        if isinstance(child, CT_P):
            if block_idx < len(blocks) and blocks[block_idx].block_type == "paragraph":
                text = blocks[block_idx].translated_text or blocks[block_idx].original_text
                from docx.text.paragraph import Paragraph
                para_obj = Paragraph(child, doc)
                _replace_paragraph_text_keep_format(para_obj, text)
                block_idx += 1
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    if block_idx < len(blocks) and blocks[block_idx].block_type == "table_cell":
                        cell.text = blocks[block_idx].translated_text or blocks[block_idx].original_text
                        block_idx += 1
    for section_idx, section in enumerate(doc.sections):
        try:
            header_blocks = sorted(
                [b for b in blocks if b.block_type == "header" and len(b.path) >= 3 and b.path[0] == section_idx and b.path[1] == "h"],
                key=lambda b: b.path[2],
            )
            for i, b in enumerate(header_blocks):
                if i < len(section.header.paragraphs):
                    _replace_paragraph_text_keep_format(section.header.paragraphs[i], b.translated_text or b.original_text)
            footer_blocks = sorted(
                [b for b in blocks if b.block_type == "footer" and len(b.path) >= 3 and b.path[0] == section_idx and b.path[1] == "f"],
                key=lambda b: b.path[2],
            )
            for i, b in enumerate(footer_blocks):
                if i < len(section.footer.paragraphs):
                    _replace_paragraph_text_keep_format(section.footer.paragraphs[i], b.translated_text or b.original_text)
        except Exception:
            pass
    doc.save(out_path)


def _write_txt(out_path: Path, blocks: List[TextBlock]) -> None:
    lines = [
        (b.translated_text or b.original_text)
        for b in blocks
        if getattr(b, "block_type", "") == "line"
    ]
    out_path.write_text("\n".join(lines), encoding="utf-8")


def _write_xlsx(source_path: Path, out_path: Path, blocks: List[TextBlock]) -> None:
    import openpyxl
    from openpyxl.cell.cell import MergedCell
    from copy import copy

    sheet_name_map = {}
    for b in blocks:
        if b.block_type == "sheet_name" and len(b.path) >= 1:
            orig = b.path[0]
            sheet_name_map[orig] = (b.translated_text or b.original_text or orig).strip()[:31]

    wb = openpyxl.load_workbook(source_path, read_only=False, data_only=False)
    # 按 (sheet, row) 记录该行最大行数（便于设置行高、打印完整）
    row_lines: dict = {}

    for block in blocks:
        if block.block_type != "table_cell" or len(block.path) != 3:
            continue
        sheet_name, row_idx, col_idx = block.path
        ws = wb.worksheets[sheet_name] if isinstance(sheet_name, int) else wb[sheet_name]
        cell = ws.cell(row=row_idx, column=col_idx)
        if isinstance(cell, MergedCell):
            # 合并区域只写左上角，避免跳过块导致后续单元格与译文错位
            try:
                for rng in ws.merged_cells.ranges:
                    if rng.min_row <= row_idx <= rng.max_row and rng.min_col <= col_idx <= rng.max_col:
                        cell = ws.cell(row=rng.min_row, column=rng.min_col)
                        break
                else:
                    continue
            except Exception:
                continue
        text = block.translated_text or block.original_text
        orig = block.original_text or ""
        # 仅在内容实际变化时写入，避免破坏原单元格富文本/上标等特殊格式
        if text != orig:
            cell.value = text
            if "\n" in (text or ""):
                # 保持原有对齐，仅在多行文本时启用自动换行，避免改动整体版式
                try:
                    ali = copy(cell.alignment) if cell.alignment is not None else None
                    if ali is not None:
                        ali.wrap_text = True
                        if not ali.vertical:
                            ali.vertical = "top"
                        cell.alignment = ali
                except Exception:
                    pass
            lines = (text or "").count("\n") + 1
            key = (sheet_name, row_idx)
            row_lines[key] = max(row_lines.get(key, 1), min(lines, 50))

    for (sname, row_idx), lines in row_lines.items():
        ws = wb[sname] if isinstance(sname, str) else wb.worksheets[sname]
        h = max(15.0, 15.0 * lines)
        if ws.row_dimensions[row_idx].height is None or (ws.row_dimensions[row_idx].height or 0) < h:
            ws.row_dimensions[row_idx].height = h

    # 写回页眉页脚（在重命名 sheet 之前，仍用原名定位）
    for b in blocks:
        if b.block_type != "xlsx_header_footer" or len(b.path) != 3:
            continue
        sheet_name, attr, part = b.path
        try:
            ws = wb[sheet_name]
            hf = getattr(ws, attr, None)
            if hf is None:
                continue
            part_obj = getattr(hf, part, None) if part != "centre" else getattr(hf, "center", None)
            if part_obj is not None:
                part_obj.text = (b.translated_text or b.original_text or "").strip()
        except Exception:
            pass

    # 写回翻译后的 sheet 名称（Excel 限制 31 字符，重名时加后缀）
    used = set()
    for ws in wb.worksheets:
        new_title = sheet_name_map.get(ws.title, ws.title)
        if not new_title:
            new_title = ws.title
        if new_title in used:
            base, n = new_title[:28], 1
            while f"{base}_{n}" in used:
                n += 1
            new_title = f"{base}_{n}"
        used.add(new_title)
        ws.title = new_title
    wb.save(out_path)
    wb.close()


def _write_document(source_path: Path, out_path: Path, blocks: List[TextBlock]) -> None:
    suffix = source_path.suffix.lower()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if suffix == ".docx":
        _write_docx(source_path, out_path, blocks)
    elif suffix == ".txt":
        _write_txt(out_path, blocks)
    elif suffix == ".xlsx":
        _write_xlsx(source_path, out_path, blocks)
    else:
        raise ValueError(f"不支持的写出格式: {suffix}")


def translate_file(
    input_path: str,
    output_path: Optional[str] = None,
    output_dir: Optional[str] = None,
    collection_name: Optional[str] = None,
    use_kb: bool = True,
    provider: Optional[str] = None,
    target_lang: str = "en",
    company_overrides: Optional[dict] = None,
    kb_query_extra: Optional[str] = None,
    translation_cache: Optional[dict] = None,
    running_glossary: Optional[dict] = None,
) -> str:
    """
    翻译单个文件（docx/txt/xlsx），正文、页眉页脚、sheet 名、文件名均译为目标语言。
    target_lang: en 英文, de 德文, zh 中文。company_overrides 可固定公司名称/地址/联系人/电话等译法。
    若指定 output_path 则直接使用；否则用 output_dir（或文件所在目录）+ 翻译后的文件名。
    返回实际输出路径。
    """
    path = Path(input_path)
    if not path.is_file():
        raise FileNotFoundError(f"文件不存在: {input_path}")
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持格式: {path.suffix}，支持: {list(SUPPORTED_EXTENSIONS)}")

    if running_glossary is None:
        running_glossary = {}

    target_lang = (target_lang or "en").strip() or "en"
    if target_lang not in ("en", "de", "zh"):
        target_lang = "en"

    translated_stem = _translate_stem(
        path.stem, collection_name, use_kb, provider,
        target_lang=target_lang, company_overrides=company_overrides,
        kb_query_extra=kb_query_extra, translation_cache=translation_cache,
        running_glossary=running_glossary,
    )
    if output_path is None:
        base = Path(output_dir) if output_dir else path.parent
        output_path = base / f"{translated_stem}{path.suffix}"
    else:
        output_path = Path(output_path)

    blocks = parse_document(str(path))
    segment_map, to_translate = blocks_to_sentences(blocks, target_lang=target_lang)
    if to_translate:
        translations = translate_sentences(
            to_translate,
            collection_name=collection_name,
            use_kb=use_kb,
            provider=provider,
            target_lang=target_lang,
            company_overrides=company_overrides,
            kb_query_extra=kb_query_extra,
            cache=translation_cache,
            running_glossary=running_glossary,
        )
        apply_translations_to_blocks(blocks, segment_map, translations, target_lang=target_lang)
    else:
        for b in blocks:
            b.translated_text = b.original_text

    _write_document(path, output_path, blocks)
    return str(output_path)


def translate_path(
    input_path: str,
    output_dir: Optional[str] = None,
    collection_name: Optional[str] = None,
    use_kb: bool = True,
    provider: Optional[str] = None,
    target_lang: str = "en",
    company_overrides: Optional[dict] = None,
    kb_query_extra: Optional[str] = None,
) -> List[str]:
    """
    翻译单文件、目录或 zip。目录递归处理支持格式；zip 解压后按目录处理，结果写出到 output_dir
    （默认：单文件为同目录，目录为 原目录_translated，zip 为 解压目录_translated）。
    返回所有输出文件路径列表。
    """
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"路径不存在: {input_path}")

    files_to_process: List[Path] = []
    temp_dirs: List[str] = []
    base_out: Optional[Path] = None

    if path.is_file():
        if path.suffix.lower() == ".zip":
            from src.core.document_loader import extract_archive
            temp_dir, extracted = extract_archive(path)
            temp_dirs.append(temp_dir)
            for p in extracted:
                if p.suffix.lower() in SUPPORTED_EXTENSIONS:
                    files_to_process.append(p)
            base_out = Path(output_dir) if output_dir else Path(temp_dir).parent / f"{path.stem}_translated"
        elif path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files_to_process = [path]
            base_out = Path(output_dir) if output_dir else path.parent
        else:
            return []
    else:
        for ext in SUPPORTED_EXTENSIONS:
            files_to_process.extend(path.rglob(f"*{ext}"))
        base_out = Path(output_dir) if output_dir else path.parent / f"{path.name}_translated"

    if not base_out:
        base_out = Path(output_dir) if output_dir else path.parent
    base_out = base_out.resolve()
    out_paths: List[str] = []
    # 同一批任务内相同原文复用译法（多文件/zip/目录）
    shared_translation_cache: dict = {}
    shared_running_glossary: dict = {}

    for fp in files_to_process:
        try:
            rel = fp.relative_to(path) if path.is_dir() else fp.name
            if temp_dirs:
                try:
                    rel = fp.relative_to(temp_dirs[0])
                except ValueError:
                    rel = fp.name
            rel_path = Path(rel) if isinstance(rel, str) else rel
            # 不传 output_path，由 translate_file 按翻译后的文件名生成到 base_out 下
            out_file = translate_file(
                str(fp),
                output_dir=str(base_out / rel_path.parent),
                collection_name=collection_name,
                use_kb=use_kb,
                provider=provider,
                target_lang=target_lang or "en",
                company_overrides=company_overrides,
                kb_query_extra=kb_query_extra,
                translation_cache=shared_translation_cache,
                running_glossary=shared_running_glossary,
            )
            out_paths.append(out_file)
        except Exception:
            if temp_dirs:
                import shutil
                for d in temp_dirs:
                    shutil.rmtree(d, ignore_errors=True)
            raise

    for d in temp_dirs:
        try:
            import shutil
            shutil.rmtree(d, ignore_errors=True)
        except Exception:
            pass
    return out_paths


def correct_file(
    input_path: str,
    output_path: Optional[str] = None,
    output_dir: Optional[str] = None,
    target_lang: str = "en",
    collection_name: Optional[str] = None,
    use_kb: bool = True,
    provider: Optional[str] = None,
    kb_query_extra: Optional[str] = None,
    translation_cache: Optional[dict] = None,
    running_glossary: Optional[dict] = None,
) -> tuple:
    """
    校正已翻译文件（docx/txt/xlsx）：同词异译、截断词、数值表达模式等。
    返回：(输出路径, 统计)。
    """
    path = Path(input_path)
    if not path.is_file():
        raise FileNotFoundError(f"文件不存在: {input_path}")
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持格式: {path.suffix}，支持: {list(SUPPORTED_EXTENSIONS)}")

    if output_path is None:
        base = Path(output_dir) if output_dir else path.parent
        output_path = base / f"{path.stem}_corrected{path.suffix}"
    else:
        output_path = Path(output_path)

    blocks = parse_document(str(path))
    stats = CorrectionStats(total_blocks=len(blocks))
    doc_all = "\n".join((b.original_text or "") for b in blocks if getattr(b, "block_type", "") in ("paragraph", "table_cell", "line", "header", "footer"))

    for b in blocks:
        src = b.original_text or ""
        fixed, d = correct_text(src, target_lang=target_lang, doc_all_text=doc_all)
        b.translated_text = fixed
        if fixed != src:
            stats.changed_blocks += 1
        stats.term_unified += d.get("term_unified", 0)
        stats.truncation_fixed += d.get("truncation_fixed", 0)
        stats.numeric_fixed += d.get("numeric_fixed", 0)

    # 第二阶段：针对校正后仍“翻译不完整”的片段做补全翻译
    if running_glossary is None:
        running_glossary = {}
    # 用校正后文本作为新一轮输入，按目标语言判定需翻译片段
    for b in blocks:
        b.original_text = b.translated_text or b.original_text
    seg_map, to_translate = blocks_to_sentences(blocks, target_lang=target_lang)
    if to_translate:
        completed = translate_sentences(
            to_translate,
            collection_name=collection_name,
            use_kb=use_kb,
            cache=translation_cache,
            provider=provider,
            target_lang=target_lang,
            kb_query_extra=kb_query_extra,
            running_glossary=running_glossary,
        )
        apply_translations_to_blocks(blocks, seg_map, completed, target_lang=target_lang)
        # 补全翻译造成的文本变化也计入 changed_blocks
        changed2 = 0
        for b in blocks:
            now = b.translated_text or ""
            before = b.original_text or ""
            if now != before:
                changed2 += 1
        stats.changed_blocks += changed2

    _write_document(path, output_path, blocks)
    return str(output_path), stats


def correct_path(
    input_path: str,
    output_dir: Optional[str] = None,
    target_lang: str = "en",
    collection_name: Optional[str] = None,
    use_kb: bool = True,
    provider: Optional[str] = None,
    kb_query_extra: Optional[str] = None,
) -> tuple:
    """
    批量校正：单文件、目录或 zip。
    返回：(输出路径列表, 汇总统计dict)。
    """
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"路径不存在: {input_path}")

    files_to_process: List[Path] = []
    temp_dirs: List[str] = []
    base_out: Optional[Path] = None

    if path.is_file():
        if path.suffix.lower() == ".zip":
            from src.core.document_loader import extract_archive
            temp_dir, extracted = extract_archive(path)
            temp_dirs.append(temp_dir)
            for p in extracted:
                if p.suffix.lower() in SUPPORTED_EXTENSIONS:
                    files_to_process.append(p)
            base_out = Path(output_dir) if output_dir else Path(temp_dir).parent / f"{path.stem}_corrected"
        elif path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files_to_process = [path]
            base_out = Path(output_dir) if output_dir else path.parent
        else:
            return [], {}
    else:
        for ext in SUPPORTED_EXTENSIONS:
            files_to_process.extend(path.rglob(f"*{ext}"))
        base_out = Path(output_dir) if output_dir else path.parent / f"{path.name}_corrected"

    if not base_out:
        base_out = Path(output_dir) if output_dir else path.parent
    base_out = base_out.resolve()

    out_paths: List[str] = []
    summary = CorrectionStats()
    shared_translation_cache: dict = {}
    shared_running_glossary: dict = {}

    for fp in files_to_process:
        rel = fp.relative_to(path) if path.is_dir() else Path(fp.name)
        if temp_dirs:
            try:
                rel = fp.relative_to(temp_dirs[0])
            except ValueError:
                rel = Path(fp.name)
        out, st = correct_file(
            str(fp),
            output_dir=str(base_out / rel.parent),
            target_lang=target_lang,
            collection_name=collection_name,
            use_kb=use_kb,
            provider=provider,
            kb_query_extra=kb_query_extra,
            translation_cache=shared_translation_cache,
            running_glossary=shared_running_glossary,
        )
        out_paths.append(out)
        summary.total_blocks += st.total_blocks
        summary.changed_blocks += st.changed_blocks
        summary.term_unified += st.term_unified
        summary.truncation_fixed += st.truncation_fixed
        summary.numeric_fixed += st.numeric_fixed

    for d in temp_dirs:
        try:
            shutil.rmtree(d, ignore_errors=True)
        except Exception:
            pass
    return out_paths, {
        "total_blocks": summary.total_blocks,
        "changed_blocks": summary.changed_blocks,
        "term_unified": summary.term_unified,
        "truncation_fixed": summary.truncation_fixed,
        "numeric_fixed": summary.numeric_fixed,
    }
