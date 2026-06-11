"""
带结构的文档解析：提取段落、表格单元格、标题等，供分句与回填使用。
支持 .docx（python-docx）、.txt、.xlsx（openpyxl）。
"""
from pathlib import Path
from typing import List, Tuple
import zipfile
import xml.etree.ElementTree as ET

from .models import TextBlock, SUPPORTED_EXTENSIONS

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = "{%s}" % _W_NS


def parse_document(file_path: str) -> List[TextBlock]:
    """
    按后缀解析单个文件，返回文档顺序的文本块列表。
    不支持格式抛出 ValueError。
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的格式: {suffix}，首期支持: {list(SUPPORTED_EXTENSIONS)}"
        )
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".txt":
        return _parse_txt(path)
    if suffix == ".xlsx":
        return _parse_xlsx(path)
    raise ValueError(f"未实现解析: {suffix}")


def _paragraph_visible_wt_nodes(p_el) -> List:
    """与 _paragraph_text_as_revised 一致：可见 w:t 节点（不含 w:del / w:moveFrom 内文本）。"""
    try:
        from docx.oxml.ns import qn

        nodes: List = []
        for t in p_el.iter(qn("w:t")):
            skip = False
            el = t
            while el is not None:
                tag = el.tag
                if tag in (qn("w:del"), qn("w:moveFrom")):
                    skip = True
                    break
                el = el.getparent()
            if not skip:
                nodes.append(t)
        return nodes
    except Exception:
        return []


def _set_paragraph_visible_text(para, new_text: str) -> None:
    """
    按「接受修订后」视图写回段落：更新全部可见 w:t，并清空 w:del / w:moveFrom 内残留，
    避免 Track Changes 修订区仍留中文而正文已译英。
    """
    try:
        from docx.oxml.ns import qn

        p = para._element
        nodes = _paragraph_visible_wt_nodes(p)
        text = new_text or ""
        if nodes:
            nodes[0].text = text
            for t in nodes[1:]:
                t.text = ""
        else:
            para.add_run(text)
        for del_el in p.iter(qn("w:del")):
            for t in del_el.iter(qn("w:t")):
                t.text = ""
        for mf in p.iter(qn("w:moveFrom")):
            for t in mf.iter(qn("w:t")):
                t.text = ""
    except Exception:
        try:
            para.text = new_text or ""
        except Exception:
            pass


def _set_cell_visible_text(cell, new_text: str) -> None:
    """表格单元格 revision-aware 写回（多段落按换行对齐）。"""
    paras = list(getattr(cell, "paragraphs", None) or [])
    if not paras:
        cell.text = new_text or ""
        return
    lines = (new_text or "").split("\n")
    if len(paras) == 1:
        _set_paragraph_visible_text(paras[0], new_text or "")
        return
    for i, para in enumerate(paras):
        _set_paragraph_visible_text(
            para, lines[i] if i < len(lines) else ""
        )


def _paragraph_text_as_revised(para) -> str:
    """近似「接受修订后」文本：跳过 w:del / w:moveFrom，保留 w:ins 与普通 run。"""
    try:
        from docx.oxml.ns import qn

        pieces: List[str] = []
        for t in para._element.iter(qn("w:t")):
            txt = t.text
            if not txt:
                continue
            el = t
            skip = False
            while el is not None:
                tag = el.tag
                if tag in (qn("w:del"), qn("w:moveFrom")):
                    skip = True
                    break
                el = el.getparent()
            if skip:
                continue
            pieces.append(txt)
        return "".join(pieces).strip()
    except Exception:
        return (para.text or "").strip()


def _paragraphs_in_hf_container(container, doc) -> list:
    """页眉/页脚内全部段落（含表格内嵌段落），顺序与 OOXML 一致。"""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    out = []
    try:
        root = container._element
    except Exception:
        return out
    for p_el in root.iter(qn("w:p")):
        try:
            out.append(Paragraph(p_el, doc))
        except Exception:
            pass
    return out


def _parse_docx_comments(path: Path) -> List[TextBlock]:
    """解析 word/comments.xml 中的批注正文（侧边修订说明等）。"""
    blocks: List[TextBlock] = []
    try:
        with zipfile.ZipFile(path, "r") as zf:
            if "word/comments.xml" not in zf.namelist():
                return blocks
            root = ET.fromstring(zf.read("word/comments.xml"))
    except Exception:
        return blocks
    for comment_el in root.findall(f"{_W}comment"):
        cid = comment_el.get(f"{_W}id") or comment_el.get("id") or ""
        p_idx = 0
        for p_el in comment_el.findall(f".//{_W}p"):
            parts = []
            for t_el in p_el.findall(f".//{_W}t"):
                if t_el.text:
                    parts.append(t_el.text)
            text = "".join(parts).strip()
            blocks.append(
                TextBlock(
                    block_type="comment",
                    path=(str(cid), p_idx),
                    original_text=text,
                )
            )
            p_idx += 1
    return blocks


def _parse_docx(path: Path) -> List[TextBlock]:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    doc = Document(path)
    blocks: List[TextBlock] = []
    body = doc.element.body
    table_idx = 0
    for child in body:
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = _paragraph_text_as_revised(para)
            blocks.append(TextBlock(
                block_type="paragraph",
                path=(len(blocks),),
                original_text=text,
            ))
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if isinstance(cell, _Cell):
                        if cell.paragraphs:
                            chunks = [_paragraph_text_as_revised(p) for p in cell.paragraphs]
                            text = "\n".join(c for c in chunks if c).strip()
                        else:
                            text = (cell.text or "").strip()
                    else:
                        text = ""
                    blocks.append(TextBlock(
                        block_type="table_cell",
                        path=(table_idx, row_idx, col_idx),
                        original_text=text,
                    ))
            table_idx += 1
    # 页眉、页脚：遍历容器内全部 w:p（含表格内），避免仅 top-level paragraphs 漏译
    for section_idx, section in enumerate(doc.sections):
        try:
            for p_idx, para in enumerate(_paragraphs_in_hf_container(section.header, doc)):
                blocks.append(TextBlock(
                    block_type="header",
                    path=(section_idx, "h", p_idx),
                    original_text=_paragraph_text_as_revised(para),
                ))
            for p_idx, para in enumerate(_paragraphs_in_hf_container(section.footer, doc)):
                blocks.append(TextBlock(
                    block_type="footer",
                    path=(section_idx, "f", p_idx),
                    original_text=_paragraph_text_as_revised(para),
                ))
        except Exception:
            pass
    blocks.extend(_parse_docx_comments(path))
    return blocks


def _parse_txt(path: Path) -> List[TextBlock]:
    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = path.read_text(encoding="gbk")
    lines = content.splitlines()
    blocks = []
    for i, line in enumerate(lines):
        blocks.append(TextBlock(
            block_type="line",
            path=(i,),
            original_text=line,
        ))
    return blocks


def _xlsx_header_footer_keys():
    """Excel 页眉页脚键：(属性名, 区域)"""
    return [
        ("oddHeader", "left"), ("oddHeader", "center"), ("oddHeader", "right"),
        ("oddFooter", "left"), ("oddFooter", "center"), ("oddFooter", "right"),
        ("evenHeader", "left"), ("evenHeader", "center"), ("evenHeader", "right"),
        ("evenFooter", "left"), ("evenFooter", "center"), ("evenFooter", "right"),
        ("firstHeader", "left"), ("firstHeader", "center"), ("firstHeader", "right"),
        ("firstFooter", "left"), ("firstFooter", "center"), ("firstFooter", "right"),
    ]


def _xlsx_is_merged_slave(sheet, row_idx: int, col_idx: int) -> bool:
    """若为合并区域内非左上角单元格，返回 True（解析时跳过，避免与主格重复导致写入错位）。"""
    try:
        for rng in sheet.merged_cells.ranges:
            if rng.min_row <= row_idx <= rng.max_row and rng.min_col <= col_idx <= rng.max_col:
                return (row_idx, col_idx) != (rng.min_row, rng.min_col)
    except Exception:
        pass
    return False


def _parse_xlsx(path: Path) -> List[TextBlock]:
    import openpyxl
    # read_only=False 以便读取页眉页脚（openpyxl 在 read_only 下可能不加载）
    wb = openpyxl.load_workbook(path, read_only=False, data_only=True)
    blocks: List[TextBlock] = []
    try:
        for sheet in wb.worksheets:
            blocks.append(TextBlock(
                block_type="sheet_name",
                path=(sheet.title,),
                original_text=sheet.title or "",
            ))
            # 页眉页脚：oddHeader/oddFooter 等下的 left/center/right
            for attr, part in _xlsx_header_footer_keys():
                try:
                    hf = getattr(sheet, attr, None)
                    if hf is None:
                        continue
                    part_obj = getattr(hf, part, None) if part != "centre" else getattr(hf, "center", None)
                    if part_obj is not None:
                        text = (getattr(part_obj, "text", None) or "").strip()
                        blocks.append(TextBlock(
                            block_type="xlsx_header_footer",
                            path=(sheet.title, attr, part),
                            original_text=text,
                        ))
                except Exception:
                    pass
        for sheet in wb.worksheets:
            # 按固定行列遍历，合并区域只保留左上角一格，避免与 iter_rows 重复格导致译文错位
            max_row = sheet.max_row or 0
            max_col = sheet.max_column or 0
            for row_idx in range(1, max_row + 1):
                for col_idx in range(1, max_col + 1):
                    if _xlsx_is_merged_slave(sheet, row_idx, col_idx):
                        continue
                    value = sheet.cell(row=row_idx, column=col_idx).value
                    text = str(value).strip() if value is not None else ""
                    blocks.append(TextBlock(
                        block_type="table_cell",
                        path=(sheet.title, row_idx, col_idx),
                        original_text=text,
                    ))
    finally:
        wb.close()
    return blocks
