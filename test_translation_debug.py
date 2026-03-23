"""调试翻译模块：不调用 LLM，仅测试解析与写回，检查是否有报错。"""
import sys
import tempfile
from pathlib import Path

# 项目根
ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

def test_txt():
    """TXT 解析与写回"""
    from src.translation.parser import parse_document
    from src.translation.segment import blocks_to_sentences, apply_translations_to_blocks
    from src.translation.pipeline import translate_file, _write_document
    from src.translation.models import TextBlock

    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("第一行。\n第二行。\nHello world.\n")
        txt_path = Path(f.name)
    out = txt_path.parent / "test_txt_out.txt"
    try:
        blocks = parse_document(str(txt_path))
        assert len(blocks) == 3, blocks
        segment_map, to_translate = blocks_to_sentences(blocks)
        # 不调 LLM，直接回填原文
        apply_translations_to_blocks(blocks, segment_map, [b.original_text for b in blocks if b.original_text])
        for b in blocks:
            if not b.translated_text:
                b.translated_text = b.original_text
        _write_document(txt_path, out, blocks)
        assert out.exists()
        print("  [OK] TXT 解析与写回")
    finally:
        txt_path.unlink(missing_ok=True)
        if out.exists():
            out.unlink(missing_ok=True)


def test_xlsx_no_merge():
    """XLSX 解析与写回（无合并单元格）"""
    import openpyxl
    from src.translation.parser import parse_document
    from src.translation.segment import blocks_to_sentences, apply_translations_to_blocks
    from src.translation.pipeline import _write_document

    xlsx_path = Path(tempfile.gettempdir()) / "test_trans_debug.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "你好"
    ws["B1"] = "世界"
    ws["A2"] = "测试"
    wb.save(xlsx_path)
    out = xlsx_path.parent / "test_xlsx_out.xlsx"
    try:
        blocks = parse_document(str(xlsx_path))
        segment_map, to_translate = blocks_to_sentences(blocks)
        apply_translations_to_blocks(blocks, segment_map, [b.original_text for b in blocks if b.original_text])
        for b in blocks:
            if not b.translated_text:
                b.translated_text = b.original_text
        _write_document(xlsx_path, out, blocks)
        assert out.exists()
        print("  [OK] XLSX 无合并 解析与写回")
    finally:
        xlsx_path.unlink(missing_ok=True)
        if out.exists():
            out.unlink(missing_ok=True)


def test_xlsx_with_merge():
    """XLSX 解析与写回（含合并单元格）"""
    import openpyxl
    from openpyxl.utils import get_column_letter
    from src.translation.parser import parse_document
    from src.translation.segment import blocks_to_sentences, apply_translations_to_blocks
    from src.translation.pipeline import _write_document

    xlsx_path = Path(tempfile.gettempdir()) / "test_trans_merge.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "标题"
    ws["A2"] = "内容1"
    ws["B2"] = "内容2"
    ws.merge_cells("A1:B1")
    wb.save(xlsx_path)
    out = xlsx_path.parent / "test_xlsx_merge_out.xlsx"
    try:
        blocks = parse_document(str(xlsx_path))
        segment_map, to_translate = blocks_to_sentences(blocks)
        apply_translations_to_blocks(blocks, segment_map, [b.original_text for b in blocks if b.original_text])
        for b in blocks:
            if not b.translated_text:
                b.translated_text = b.original_text
        _write_document(xlsx_path, out, blocks)
        assert out.exists()
        print("  [OK] XLSX 含合并 解析与写回")
    finally:
        xlsx_path.unlink(missing_ok=True)
        if out.exists():
            out.unlink(missing_ok=True)


def test_docx():
    """DOCX 解析与写回"""
    from docx import Document
    from src.translation.parser import parse_document
    from src.translation.segment import blocks_to_sentences, apply_translations_to_blocks
    from src.translation.pipeline import _write_document

    docx_path = Path(tempfile.gettempdir()) / "test_trans_debug.docx"
    doc = Document()
    doc.add_paragraph("段落一。")
    doc.add_paragraph("段落二。")
    doc.save(docx_path)
    out = docx_path.parent / "test_docx_out.docx"
    try:
        blocks = parse_document(str(docx_path))
        segment_map, to_translate = blocks_to_sentences(blocks)
        apply_translations_to_blocks(blocks, segment_map, [b.original_text for b in blocks if b.original_text])
        for b in blocks:
            if not b.translated_text:
                b.translated_text = b.original_text
        _write_document(docx_path, out, blocks)
        assert out.exists()
        print("  [OK] DOCX 解析与写回")
    finally:
        docx_path.unlink(missing_ok=True)
        if out.exists():
            out.unlink(missing_ok=True)


def test_docx_with_table():
    """DOCX 含表格 解析与写回"""
    from docx import Document
    from src.translation.parser import parse_document
    from src.translation.segment import blocks_to_sentences, apply_translations_to_blocks
    from src.translation.pipeline import _write_document

    docx_path = Path(tempfile.gettempdir()) / "test_trans_table.docx"
    doc = Document()
    doc.add_paragraph("表前段落。")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A1"
    t.cell(0, 1).text = "B1"
    t.cell(1, 0).text = "A2"
    t.cell(1, 1).text = "B2"
    doc.add_paragraph("表后段落。")
    doc.save(docx_path)
    out = docx_path.parent / "test_docx_table_out.docx"
    try:
        blocks = parse_document(str(docx_path))
        segment_map, to_translate = blocks_to_sentences(blocks)
        apply_translations_to_blocks(blocks, segment_map, [b.original_text for b in blocks if b.original_text])
        for b in blocks:
            if not b.translated_text:
                b.translated_text = b.original_text
        _write_document(docx_path, out, blocks)
        assert out.exists()
        print("  [OK] DOCX 含表格 解析与写回")
    finally:
        docx_path.unlink(missing_ok=True)
        if out.exists():
            out.unlink(missing_ok=True)


def test_translate_path_single_file():
    """translate_path 单文件（不调 LLM，用无中文文件避免实际翻译）"""
    from src.translation.pipeline import translate_path
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("Hello.\nWorld.\n")
        txt_path = Path(f.name)
    out_dir = Path(tempfile.mkdtemp())
    try:
        # 无中文时 to_translate 为空，不会调 LLM
        out_paths = translate_path(
            str(txt_path),
            output_dir=str(out_dir),
            use_kb=False,
        )
        assert len(out_paths) == 1
        assert Path(out_paths[0]).exists()
        print("  [OK] translate_path 单文件")
    finally:
        txt_path.unlink(missing_ok=True)
        import shutil
        shutil.rmtree(out_dir, ignore_errors=True)


if __name__ == "__main__":
    print("调试翻译模块（不调用 LLM）...")
    test_txt()
    test_xlsx_no_merge()
    test_xlsx_with_merge()
    test_docx()
    test_docx_with_table()
    test_translate_path_single_file()
    print("全部通过。")
